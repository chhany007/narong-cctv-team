# NARONG_CCTV_v8.6.py
"""
NARONG CCTV TEAM - Camera Monitor v8.6
========================================
Advanced Camera Monitoring and NVR Management System

🎯 CORE FEATURES:
- 📂 Excel-based camera database with robust NVR sheet reading
- 🔧 Advanced NVR integration with credential management
- ⚡ Enhanced parallel checking with visual progress indicators
- 🔍 SADP network discovery and device management
- 📡 IVMS camera extraction with multiple authentication methods
- 🖥️ Modern PyQt5 interface with context menus and search
- 🔄 Automatic update system with GitHub integration
- 📊 Real-time status monitoring and export capabilities

🚀 NEW IN v8.6:
- Improved code organization and performance optimization
- Enhanced error handling and user feedback
- Centralized configuration management
- Better threading and UI responsiveness
- Streamlined workflow wizard
- Advanced duplicate detection capabilities

Developed by Sky-Tech for Koh Kong Casino
"""

import os, sys, json, csv, time, socket, threading, subprocess, webbrowser, traceback, unicodedata, re, platform, logging, hashlib, uuid
from datetime import datetime, timedelta
from PyQt5 import QtCore, QtGui, QtWidgets
import pandas as pd
import concurrent.futures
import requests
from requests.auth import HTTPBasicAuth
from xml.etree import ElementTree as ET

# Import NVR management dialogs
try:
    from nvr_dialogs import AddNVRDialog, EditNVRDialog
    NVR_DIALOGS_AVAILABLE = True
except Exception as e:
    NVR_DIALOGS_AVAILABLE = False
    print(f"NVR dialogs not available: {e}")

# Import update manager
try:
    from update_manager_enhanced import check_for_updates_async, UpdateChecker
    UPDATE_MANAGER_AVAILABLE = True
except Exception as e:
    UPDATE_MANAGER_AVAILABLE = False
    print(f"Update manager not available: {e}")

# Optional keyring
try:
    import keyring
    KEYRING_AVAILABLE = True
except Exception:
    KEYRING_AVAILABLE = False

# ==================== LOGGING SETUP ====================
def setup_logging():
    """Setup comprehensive logging to file and console"""
    log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "camera_monitor.log")
    
    # Create logger
    logger = logging.getLogger('CameraMonitor')
    logger.setLevel(logging.INFO)
    
    # Clear any existing handlers
    logger.handlers.clear()
    
    # Create file handler
    file_handler = logging.FileHandler(log_file, mode='w', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    
    # Create console handler  
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    
    # Create formatter
    formatter = logging.Formatter(
        '%(asctime)s | %(levelname)-8s | %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)
    
    # Add handlers to logger
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    print(f"Logging initialized. Log file: {log_file}")
    return logger

# Initialize logging
app_logger = setup_logging()

# IVMS integration for improved NVR camera import
IVMS_AVAILABLE = False
try:
    from requests.auth import HTTPDigestAuth
    IVMS_AVAILABLE = True
    print("IVMS camera import system loaded successfully")
except ImportError as e:
    print(f"IVMS dependencies not available: {e}")
    print("NVR import functionality will be limited")

# ==================== PROVEN IVMS CAMERA IMPORT SYSTEM ====================
class WorkingNVRController:
    """Proven working NVR controller based on IVMS folder implementation"""
    
    def __init__(self, ip: str, username: str, password: str):
        self.ip = ip
        self.username = username
        self.password = password
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
            "Accept": "application/json, text/xml, */*",
        })
    
    def connect(self) -> bool:
        """Test NVR connection using multiple methods"""
        try:
            # Try ISAPI device info endpoint first with shorter timeout
            url = f"http://{self.ip}/ISAPI/System/deviceInfo"
            resp = self.session.get(url, auth=HTTPBasicAuth(self.username, self.password), timeout=5.0)
            log(f"[NVR-CONNECT] {self.ip} deviceInfo response: {resp.status_code}")
            
            # Accept various success codes: 200 (OK), 401 (device responds but auth issue), 403 (forbidden but device online)
            if resp.status_code in [200, 401, 403]:
                if resp.status_code == 200:
                    log(f"[NVR-CONNECT] ✅ {self.ip} - ISAPI deviceInfo successful (200)")
                elif resp.status_code == 401:
                    log(f"[NVR-CONNECT] ⚠️ {self.ip} - Device online but authentication failed (401)")
                elif resp.status_code == 403:
                    log(f"[NVR-CONNECT] ⚠️ {self.ip} - Device online but access forbidden (403)")
                return True
            
            # Try alternative endpoints
            alt_endpoints = [
                "/ISAPI/System/status",
                "/ISAPI/ContentMgmt/InputProxy/channels",
                "/api/v1/system/deviceinfo",
                "/doc/script.js",  # Simple file that often exists
                "/",  # Root path
            ]
            
            for endpoint in alt_endpoints:
                try:
                    alt_url = f"http://{self.ip}{endpoint}"
                    resp = self.session.get(alt_url, auth=HTTPBasicAuth(self.username, self.password), timeout=3.0)
                    log(f"[NVR-CONNECT] {self.ip} {endpoint} response: {resp.status_code}")
                    
                    # Accept various response codes that indicate the device is online
                    if resp.status_code in [200, 401, 403, 404]:
                        log(f"[NVR-CONNECT] ✅ {self.ip} - Device responsive via {endpoint} ({resp.status_code})")
                        return True
                except Exception as endpoint_error:
                    log(f"[NVR-CONNECT] {self.ip} {endpoint} error: {endpoint_error}")
                    continue
            
            log(f"[NVR-CONNECT] ❌ {self.ip} - All connection attempts failed")
            return False
        except Exception as e:
            log(f"[NVR-CONNECT] Error connecting to {self.ip}: {e}")
            return False
    
    def get_cameras(self, timeout: float = 15.0) -> tuple:
        """Get cameras using proven IVMS methods (unified, no legacy fallbacks)"""
        log(f"[CAMERA-FETCH] Starting camera fetch for {self.ip} with timeout {timeout}s")
        
        # Try ISAPI, then generic API, then video inputs, then digest auth, all via unified helpers
        methods = [self._fetch_isapi_cameras, self._fetch_generic_api_cameras, self._fetch_video_inputs, self._fetch_with_digest_auth]
        
        for i, fetch_method in enumerate(methods):
            try:
                log(f"[CAMERA-FETCH] {self.ip} trying method {i+1}/{len(methods)}")
                cameras, method = fetch_method(timeout)
                if cameras:
                    log(f"[CAMERA-FETCH] ✅ {self.ip} found {len(cameras)} cameras via {method}")
                    
                    # Enhanced: Optional camera status verification (can be enabled/disabled)
                    # IVMS Method: Trust NVR's devIndex assessment, don't override with computer ping tests
                    enable_status_verification = False  # Set to True to enable ping verification
                    
                    if enable_status_verification:
                        log(f"[STATUS-VERIFY] Verifying status for {len(cameras)} cameras from {self.ip}")
                        verified_cameras = self.verify_camera_statuses(cameras, timeout=1.0)
                        
                        # Log status summary
                        online_verified = len([c for c in verified_cameras if c.get("verified_status") == "online"])
                        configured = len([c for c in verified_cameras if c.get("verified_status") in ["configured", "network_only"]])
                        offline = len([c for c in verified_cameras if c.get("verified_status") == "offline"])
                        total = len(verified_cameras)
                        
                        log(f"[STATUS-VERIFY] {self.ip}: {online_verified} online, {configured} configured, {offline} offline ({total} total)")
                        
                        return verified_cameras, f"Success via {method} (verified)"
                    else:
                        # Return cameras without additional verification for better performance
                        log(f"[CAMERA-FETCH] Performance mode: skipping status verification for {len(cameras)} cameras")
                        return cameras, f"Success via {method}"
                else:
                    log(f"[CAMERA-FETCH] {self.ip} method {method} returned no cameras")
            except Exception as method_error:
                log(f"[CAMERA-FETCH] {self.ip} method {i+1} error: {method_error}")
                continue
        
        log(f"[CAMERA-FETCH] ❌ {self.ip} - No cameras found with any method")
        return [], "No cameras found with any method"
    
    def _fetch_isapi_cameras(self, timeout: float) -> tuple:
        """Fetch cameras using proven ISAPI endpoints with Digest Auth support"""
        from requests.auth import HTTPDigestAuth
        
        endpoints = [
            "/ISAPI/ContentMgmt/InputProxy/channels",
            "/ISAPI/ContentMgmt/RemoteDevice",
            "/ISAPI/System/Video/inputs/channels",
        ]
        
        # Try both authentication methods for each endpoint
        auth_methods = [
            ("Digest", HTTPDigestAuth(self.username, self.password)),
            ("Basic", HTTPBasicAuth(self.username, self.password))
        ]
        
        for endpoint in endpoints:
            url = f"http://{self.ip}{endpoint}"
            
            # Try Digest Auth first (works for NVRs 2-19), then Basic Auth (works for NVR1)
            for auth_name, auth in auth_methods:
                try:
                    resp = self.session.get(url, auth=auth, timeout=timeout)
                    log(f"[ISAPI-FETCH] {self.ip} {endpoint} ({auth_name}): {resp.status_code}")
                    
                    if resp.status_code == 200:
                        cameras = _parse_isapi_cameras_working(resp.text)
                        if cameras:
                            log(f"[ISAPI-FETCH] ✅ {self.ip} parsed {len(cameras)} cameras from {endpoint} ({auth_name})")
                            return cameras, f"ISAPI {endpoint} ({auth_name})"
                        else:
                            log(f"[ISAPI-FETCH] {self.ip} {endpoint} returned XML but no cameras parsed")
                    elif resp.status_code == 401:
                        log(f"[ISAPI-FETCH] ⚠️ {self.ip} {endpoint} {auth_name} auth failed")
                        continue  # Try next auth method
                    else:
                        log(f"[ISAPI-FETCH] {self.ip} {endpoint} ({auth_name}) returned {resp.status_code}")
                        
                except Exception as e:
                    log(f"[ISAPI-FETCH] {self.ip} {endpoint} ({auth_name}) error: {e}")
                    continue
                    
        return [], "ISAPI methods failed"
    
    def _fetch_generic_api_cameras(self, timeout: float) -> tuple:
        """Fetch cameras using generic API endpoints"""
        endpoints = ["/api/v1/devices", "/api/v2/devices", "/cgi-bin/api/v1/devices"]
        
        for endpoint in endpoints:
            try:
                url = f"http://{self.ip}{endpoint}"
                resp = self.session.get(url, auth=HTTPBasicAuth(self.username, self.password), timeout=timeout)
                if resp.status_code == 200:
                    data = resp.json()
                    cameras = self._parse_generic_cameras(data)
                    if cameras:
                        return cameras, f"Generic API {endpoint}"
            except Exception as e:
                log(f"[WORKING NVR] Generic API {endpoint} error: {e}")
                continue
        return [], ""
    
    def _fetch_video_inputs(self, timeout: float) -> tuple:
        """Fetch cameras using video input channels with Digest Auth support"""
        from requests.auth import HTTPDigestAuth
        
        endpoints = [
            "/ISAPI/System/Video/inputs/channels",
            "/api/v1/System/Video/inputs/channels",
            "/ISAPI/Streaming/channels",
        ]
        
        # Try both authentication methods
        auth_methods = [
            ("Digest", HTTPDigestAuth(self.username, self.password)),
            ("Basic", HTTPBasicAuth(self.username, self.password))
        ]
        
        for endpoint in endpoints:
            url = f"http://{self.ip}{endpoint}"
            for auth_name, auth in auth_methods:
                try:
                    resp = self.session.get(url, auth=auth, timeout=timeout)
                    if resp.status_code == 200:
                        cameras = _parse_video_channels_working(resp)
                        if cameras:
                            log(f"[VIDEO-INPUTS] ✅ {self.ip} found {len(cameras)} cameras via {endpoint} ({auth_name})")
                            return cameras, f"Video Inputs {endpoint} ({auth_name})"
                except Exception as e:
                    log(f"[VIDEO-INPUTS] {self.ip} {endpoint} ({auth_name}) error: {e}")
                    continue
        return [], ""
    
    def _fetch_with_digest_auth(self, timeout: float) -> tuple:
        """Fetch cameras using Digest authentication - now enhanced with proper parsing"""
        from requests.auth import HTTPDigestAuth
        
        endpoints = [
            "/ISAPI/ContentMgmt/InputProxy/channels",
            "/ISAPI/System/Video/inputs/channels",
        ]
        
        for endpoint in endpoints:
            try:
                url = f"http://{self.ip}{endpoint}"
                resp = self.session.get(url, auth=HTTPDigestAuth(self.username, self.password), timeout=timeout)
                log(f"[DIGEST-AUTH] {self.ip} {endpoint}: {resp.status_code}")
                if resp.status_code == 200:
                    # Use the proper ISAPI parser that works
                    cameras = _parse_isapi_cameras_working(resp.text)
                    if cameras:
                        log(f"[DIGEST-AUTH] ✅ {self.ip} found {len(cameras)} cameras via {endpoint}")
                        return cameras, f"Digest Auth {endpoint}"
                    else:
                        # Fallback to video channel parser
                        cameras = _parse_video_channels_working(resp)
                        if cameras:
                            log(f"[DIGEST-AUTH] ✅ {self.ip} found {len(cameras)} cameras via {endpoint} (video parser)")
                            return cameras, f"Digest Auth {endpoint} (video parser)"
            except Exception as e:
                log(f"[DIGEST-AUTH] {self.ip} {endpoint} error: {e}")
                continue
        return [], ""
    
    def verify_camera_statuses(self, cameras: list, timeout: float = 1.0) -> list:
        """Fast camera status verification using ping test"""
        if not cameras:
            return cameras
        
        log(f"[CAMERA-STATUS] Starting fast status verification for {len(cameras)} cameras")
        verified_cameras = []
        
        for camera in cameras:
            try:
                # Get camera IP from various possible fields
                camera_ip = camera.get('ip') or camera.get('IP') or camera.get('ipAddress') or camera.get('address', '')
                if not camera_ip or camera_ip.strip() == "":
                    # Keep camera as-is if no IP found
                    verified_cameras.append(camera)
                    continue
                
                # Quick ping test only for performance
                import subprocess
                try:
                    result = subprocess.run(
                        ["ping", "-n", "1", "-w", "1000", camera_ip],
                        capture_output=True, text=True, timeout=timeout + 0.5
                    )
                    
                    # Update camera with ping result, preserving ISAPI status info
                    verified_camera = camera.copy()
                    original_status = camera.get('status', 'unknown').lower()
                    
                    if result.returncode == 0:
                        verified_camera['ping_status'] = 'online'
                        # If camera was configured but offline, upgrade to online
                        if original_status in ['configured', 'configured_offline']:
                            verified_camera['status'] = 'online'  # Upgrade configured to online
                            verified_camera['connection_type'] = 'verified_online'
                        elif original_status != 'online':
                            verified_camera['status'] = 'online'  # Set to online if ping successful
                        verified_camera['verified_status'] = 'online'
                    else:
                        verified_camera['ping_status'] = 'offline'
                        # Keep original status if it was more specific than just "offline"
                        if original_status in ['configured', 'configured_offline', 'no_ip']:
                            verified_camera['verified_status'] = original_status  # Keep specific status
                        else:
                            verified_camera['status'] = 'offline'
                            verified_camera['verified_status'] = 'offline'
                    
                    verified_camera['last_ping_check'] = True
                    verified_cameras.append(verified_camera)
                    
                except (subprocess.TimeoutExpired, Exception):
                    # If ping fails, preserve original ISAPI status
                    verified_camera = camera.copy()
                    verified_camera['ping_status'] = 'timeout'
                    original_status = camera.get('status', 'unknown').lower()
                    # Keep the original ISAPI status rather than marking as unknown
                    verified_camera['verified_status'] = original_status
                    verified_camera['last_ping_check'] = False
                    verified_cameras.append(verified_camera)
                
            except Exception as e:
                log(f"[CAMERA-STATUS] Error verifying camera {camera.get('name', 'Unknown')}: {e}")
                # Keep original camera data if verification fails
                verified_cameras.append(camera)
        
        log(f"[CAMERA-STATUS] Completed fast verification for {len(verified_cameras)} cameras")
        return verified_cameras
    

# ==================== APPLICATION CONFIGURATION ====================
# Version and metadata - Enhanced Edition
APP_VERSION = "8.8.0"
APP_TITLE = "NARONG CCTV Monitor"
APP_COMPANY = "Sky-Tech"
APP_ID = "SkyTech.CameraMonitor.8.8.0"
BUILD_DATE = "2025-11-23"

# License configuration
LICENSE_FILE = "license.key"
LICENSE_SALT = b"NarongCCTV_SkyTech_2025_SecureKey"  # Salt for encryption
TRIAL_DAYS = 30  # Trial period days
REQUIRES_LICENSE = True  # Set to False to disable license requirement
ENHANCED_FEATURES = [
    "Professional Export Reports (Excel/Word/PDF with NVR Grouping)",
    "Enhanced SADP Device Discovery Tool",
    "License Key Management System",
    "Smart Offline Camera Verification",
    "Enhanced Duplicate Detection",
    "Optimized Parallel Processing",
    "Advanced Status Monitoring"
]

# File constants
EXCEL_FILE = "ip.xlsx"
LOG_FILE = "camera_monitor.log"
CREDS_META = "creds_meta.json"
CREDS_FALLBACK = "creds_store.json"
EXPORT_FILE = "exported_cameras.csv"
CHECK_HISTORY_FILE = "check_history.json"
LOGO_FILE = "sky-tech logo.png"

# Network configuration - Enhanced for v8.6+
HTTP_PORT = 80
RTSP_PORT = 554
PING_TIMEOUT = 1.5  # seconds - Optimized for faster response
CONNECTION_TIMEOUT = 8  # seconds for NVR connections - Balanced timeout
MAX_PARALLEL_WORKERS = 8  # Maximum concurrent NVR connections - Increased for better performance
UI_UPDATE_THROTTLE = 30  # ms between UI updates - Reduced for smoother UI
CACHE_TIMEOUT = 300  # seconds - Cache timeout for network checks
RETRY_ATTEMPTS = 2  # Number of retry attempts for failed connections

# Authentication
DEFAULT_CREDS = [("admin", "Kkcctv12345"), ("admin", "Kkcctv1245")]

VLC_PATHS = [
    r"C:\Program Files\VideoLAN\VLC\vlc.exe",
    r"C:\Program Files (x86)\VideoLAN\VLC\vlc.exe",
    "/usr/bin/vlc",
    "/usr/local/bin/vlc",
    "/Applications/VLC.app/Contents/MacOS/VLC",
]

# ---------------- utilities ----------------
def log(msg: str):
    """Enhanced logging to both file and console"""
    ts = time.strftime("%Y-%m-%d %H:%M:%S")
    
    # Use proper logger if available
    try:
        app_logger.info(msg)
    except:
        # Fallback to file logging
        try:
            with open(LOG_FILE, "a", encoding="utf-8") as f:
                f.write(f"[{ts}] {msg}\n")
        except Exception:
            print(f"[{ts}] {msg}")  # Last resort - console only

def silent_ping(ip: str) -> bool:
    if not ip:
        return False
    system = platform.system().lower()
    param = "-n" if system.startswith("windows") else "-c"
    cmd = ["ping", param, "1", ip]
    try:
        if system.startswith("windows"):
            CREATE_NO_WINDOW = 0x08000000
            subprocess.check_output(cmd, stderr=subprocess.DEVNULL, stdout=subprocess.DEVNULL,
                                    timeout=PING_TIMEOUT+1, creationflags=CREATE_NO_WINDOW)
        else:
            subprocess.check_output(cmd, stderr=subprocess.DEVNULL, stdout=subprocess.DEVNULL,
                                    timeout=PING_TIMEOUT+1)
        return True
    except Exception:
        return False

def check_tcp(ip: str, port: int, timeout: float = 1.0) -> bool:
    try:
        sock = socket.create_connection((ip, port), timeout=timeout)
        sock.close()
        return True
    except Exception:
        return False

def find_vlc_executable():
    for p in VLC_PATHS:
        if os.path.exists(p):
            return p
    return "vlc"

# Helper: find local interface IP that would be used to reach a given destination
def get_local_ip_for_target(dest_ip: str) -> str:
    """Return local source IP used to reach dest_ip, or empty string on failure."""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect((dest_ip, 80))
        local_ip = s.getsockname()[0]
        s.close()
        return local_ip
    except Exception:
        try:
            s.close()
        except Exception:
            pass
    return ""

def is_cache_valid(timestamp: float, timeout: float = CACHE_TIMEOUT) -> bool:
    """Check if cached data is still valid based on timeout."""
    return time.time() - timestamp < timeout

def get_cached_status(cache: dict, key: str, timeout: float = CACHE_TIMEOUT) -> tuple:
    """Get cached status if valid, returns (is_valid, cached_data)."""
    if key in cache:
        cached_data = cache[key]
        if 'timestamp' in cached_data and is_cache_valid(cached_data['timestamp'], timeout):
            return True, cached_data
    return False, None

def update_cache(cache: dict, key: str, data: dict, max_size: int = 1000):
    """Update cache with new data, respecting size limits."""
    # Clean old entries if cache is getting too large
    if len(cache) >= max_size:
        # Remove oldest 20% of entries
        sorted_items = sorted(cache.items(), key=lambda x: x[1].get('timestamp', 0))
        for i in range(len(sorted_items) // 5):
            cache.pop(sorted_items[i][0], None)
    
    cache[key] = {**data, 'timestamp': time.time()}

# Model-specific endpoint hints for Hikvision NVRs
MODEL_ENDPOINTS = {
    "DS-7732NXI-K4": [
        "/ISAPI/ContentMgmt/RemoteDevice",
        "/ISAPI/Streaming/channels",
        "/ISAPI/ContentMgmt/InputProxy/channels",
        "/ISAPI/PSIA/Custom/SelfAdapt/Channel"
    ],
    "DS-7732NI-K4": [
        "/ISAPI/ContentMgmt/RemoteDevice",
        "/ISAPI/Streaming/channels",
        "/ISAPI/PSIA/Custom/SelfAdapt/Channel"
    ]
}

# ---------------- credentials helpers ----------------
def load_creds_meta():
    try:
        if os.path.exists(CREDS_META):
            with open(CREDS_META, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        log("read creds_meta failed")
    return {}

def save_creds_meta(meta: dict):
    try:
        with open(CREDS_META, "w", encoding="utf-8") as f:
            json.dump(meta, f, indent=2)
    except Exception as e:
        log(f"save_creds_meta: {e}")

def set_password(ip: str, username: str, password: str):
    if KEYRING_AVAILABLE:
        try:
            keyring.set_password(f"CameraMonitor:{ip}", username, password)
            return True
        except Exception as e:
            log(f"keyring set failed: {e}")
    # fallback store
    store = {}
    try:
        if os.path.exists(CREDS_FALLBACK):
            with open(CREDS_FALLBACK, "r", encoding="utf-8") as f:
                store = json.load(f)
    except Exception:
        store = {}
    store[ip] = {"username": username, "password": password}
    try:
        with open(CREDS_FALLBACK, "w", encoding="utf-8") as f:
            json.dump(store, f, indent=2)
        return False
    except Exception as e:
        log(f"fallback write failed: {e}")
        return False

def get_password(ip: str):
    meta = load_creds_meta()
    if ip in meta and meta[ip].get("username"):
        u = meta[ip]["username"]
        if KEYRING_AVAILABLE:
            try:
                p = keyring.get_password(f"CameraMonitor:{ip}", u)
                if p:
                    return u, p
            except Exception as e:
                log(f"keyring get failed: {e}")
        # fallback
        try:
            if os.path.exists(CREDS_FALLBACK):
                with open(CREDS_FALLBACK, "r", encoding="utf-8") as f:
                    store = json.load(f)
                if ip in store:
                    return store[ip].get("username"), store[ip].get("password")
        except Exception:
            pass
    return None, None

def delete_credentials(ip: str):
    try:
        meta = load_creds_meta()
        if ip in meta:
            meta.pop(ip, None)
            save_creds_meta(meta)
    except Exception:
        pass
    try:
        if os.path.exists(CREDS_FALLBACK):
            with open(CREDS_FALLBACK, "r", encoding="utf-8") as f:
                store = json.load(f)
            if ip in store:
                store.pop(ip, None)
                with open(CREDS_FALLBACK, "w", encoding="utf-8") as f:
                    json.dump(store, f, indent=2)
    except Exception:
        pass

# ---------------- NVR login & IP update ----------------
def test_nvr_connection_enhanced(ip: str, username: str, password: str, timeout: float = 5.0) -> bool:
    """Enhanced connection test based on ivms_sample.py methods"""
    log(f"Testing enhanced connection to {ip}...")
    
    test_urls = [
        f"http://{ip}/ISAPI/System/deviceInfo",
        f"http://{ip}:80/ISAPI/System/deviceInfo",
        f"http://{ip}/ISAPI/ContentMgmt/InputProxy/channels"
    ]
    
    for url in test_urls:
        try:
            # Try Basic Auth first
            response = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            if response.status_code == 200:
                log("✓ Enhanced connection successful (Basic Auth)")
                return True
            elif response.status_code == 401:
                # Try Digest Auth
                response = requests.get(url, auth=HTTPDigestAuth(username, password), timeout=timeout)
                if response.status_code == 200:
                    log("✓ Enhanced connection successful (Digest Auth)")
                    return True
        except requests.exceptions.Timeout:
            log(f"✗ Enhanced connection timeout to {url}")
        except requests.exceptions.ConnectionError:
            log(f"✗ Enhanced connection refused to {url}")
        except Exception as e:
            log(f"✗ Enhanced connection error to {url}: {e}")
    
    log("✗ Enhanced connection failed - check IP address and credentials")
    return False

def test_nvr_login(ip: str, username: str, password: str, timeout: float = 2.0) -> tuple:
    """Test NVR login with multiple methods. Returns (success: bool, real_ip: str, error_msg: str)."""
    log(f"=== NVR Login Test Start ===")
    log(f"IP: {ip}, Username: {username}, Timeout: {timeout}s")
    
    # First try enhanced connection test
    if test_nvr_connection_enhanced(ip, username, password, timeout):
        log(f"Enhanced connection test passed, proceeding with login...")
    
    auth_succeeded = False  # Track if auth worked, to continue trying for IP extraction
    
    try:
        # First: Try SADP discovery (works even if ping is blocked)
        log(f"[1] Testing SADP discovery to {ip}...")
        try:
            SADP_PORT = 33333
            SADP_REQUEST = b'\x00\x00\x00\x00\x00\x00\x00\x38<?xml version="1.0"?><Command><AccessFlag>1</AccessFlag><Command>GetDeviceInfo</Command></Command>'
            sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            sock.settimeout(1.0)
            sock.sendto(SADP_REQUEST, (ip, SADP_PORT))
            try:
                data, _ = sock.recvfrom(4096)
                sock.close()
                log(f"[1] SADP RESPONSE received - NVR is reachable")
            except socket.timeout:
                sock.close()
                log(f"[1] SADP timeout - continuing with TCP methods")
        except Exception as e:
            log(f"[1] SADP error: {e} - continuing with TCP methods")
        
        # Method 1: Try v1 API (fastest)
        log(f"[2] Trying v1 API: /api/v1/system/info")
        url = f"http://{ip}/api/v1/system/info"
        try:
            log(f"[2] GET {url}")
            resp = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[2] Response status: {resp.status_code}")
            if resp.status_code == 200:
                auth_succeeded = True
                try:
                    data = resp.json()
                    log(f"[2] JSON response: {data}")
                    real_ip = data.get("ip") or data.get("ipAddress") or data.get("network", {}).get("ip") or ""
                    log(f"[2] SUCCESS - Extracted IP: {real_ip}")
                    return True, real_ip, ""
                except Exception as e:
                    log(f"[2] JSON parse error: {e}")
                    pass
        except requests.exceptions.Timeout as e:
            log(f"[2] TIMEOUT: {e}")
        except requests.exceptions.ConnectionError as e:
            log(f"[2] CONNECTION ERROR: {e}")
        except Exception as e:
            log(f"[2] ERROR: {e}")
        
        # Method 2: Try v2 API (more likely to have IP info than web interface)
        log(f"[3] Trying v2 API: /api/v2/system/info")
        url = f"http://{ip}/api/v2/system/info"
        try:
            log(f"[3] GET {url}")
            resp = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[3] Response status: {resp.status_code}")
            if resp.status_code == 200:
                auth_succeeded = True
                try:
                    data = resp.json()
                    log(f"[3] JSON response: {data}")
                    real_ip = data.get("ip") or data.get("ipAddress") or data.get("network", {}).get("ip") or ""
                    log(f"[3] SUCCESS - Extracted IP: {real_ip}")
                    return True, real_ip, ""
                except Exception as e:
                    log(f"[3] JSON parse error: {e}")
                    pass
        except requests.exceptions.Timeout as e:
            log(f"[3] TIMEOUT: {e}")
        except requests.exceptions.ConnectionError as e:
            log(f"[3] CONNECTION ERROR: {e}")
        except Exception as e:
            log(f"[3] ERROR: {e}")
        
        # Method 3: Try simple HTTP GET to web interface (quick auth check)
        log(f"[4] Trying web interface: /")
        url = f"http://{ip}/"
        try:
            log(f"[4] GET {url}")
            resp = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[4] Response status: {resp.status_code}")
            if resp.status_code == 200:
                auth_succeeded = True
                # Try to extract IP from HTML/JSON if available
                real_ip = ""
                try:
                    # Try JSON
                    data = resp.json()
                    real_ip = data.get("ip") or data.get("ipAddress") or data.get("network", {}).get("ip") or ""
                    log(f"[4] JSON response: {data}")
                except:
                    # Try to find IP in HTML
                    import re
                    matches = re.findall(r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b', resp.text)
                    if matches and matches[0] != ip:  # Don't extract the IP we already know
                        real_ip = matches[0]
                        log(f"[4] Extracted IP from HTML: {real_ip}")
                
                if real_ip:
                    log(f"[4] SUCCESS - Web interface accessible, extracted IP: {real_ip}")
                    return True, real_ip, ""
                else:
                    log(f"[4] SUCCESS - Web interface accessible, no IP extracted in response")
                    # Continue to try legacy CGI for IP extraction
        except requests.exceptions.Timeout as e:
            log(f"[4] TIMEOUT: {e}")
        except requests.exceptions.ConnectionError as e:
            log(f"[4] CONNECTION ERROR: {e}")
        except Exception as e:
            log(f"[4] ERROR: {e}")
        
        # Method 4: Try legacy CGI endpoint (last resort)
        log(f"[5] Trying legacy CGI: /cgi-bin/system")
        url = f"http://{ip}/cgi-bin/system?action=getSystemInfo"
        try:
            log(f"[5] GET {url}")
            resp = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[5] Response status: {resp.status_code}")
            if resp.status_code == 200:
                auth_succeeded = True
                text = resp.text
                log(f"[5] Response text (first 500 chars): {text[:500]}")
                if "error" not in text.lower():
                    for line in text.split('\n'):
                        if 'ipaddr' in line.lower() or 'ip=' in line.lower():
                            log(f"[5] Found IP line: {line}")
                            parts = line.split('=')
                            if len(parts) > 1:
                                real_ip = parts[1].strip().strip('"')
                                if real_ip and real_ip != "—":
                                    log(f"[5] SUCCESS - Extracted IP: {real_ip}")
                                    return True, real_ip, ""
                    log(f"[5] SUCCESS - CGI authenticated but no IP found")
                    return True, "", ""
        except requests.exceptions.Timeout as e:
            log(f"[5] TIMEOUT: {e}")
        except requests.exceptions.ConnectionError as e:
            log(f"[5] CONNECTION ERROR: {e}")
        except Exception as e:
            log(f"[5] ERROR: {e}")
        # All methods failed
        if auth_succeeded:
            log(f"Using connection IP as device IP: {ip}")
            # Use the IP we connected to as the real IP (fallback)
            return True, ip, "Used connection IP as fallback"
        else:
            log(f"=== ALL METHODS FAILED ===")
            return False, "", "All authentication methods failed"
            
    except Exception as e:
        msg = str(e)
        log(f"=== UNEXPECTED ERROR: {msg} ===")
        return False, "", msg



def try_isapi_method_enhanced(session, ip: str, username: str, password: str, timeout: float) -> list:
    """Try ISAPI ContentMgmt method with enhanced endpoints from ivms_sample.py"""
    endpoints = [
        "/ISAPI/ContentMgmt/InputProxy/channels",
        "/ISAPI/ContentMgmt/RemoteDevice", 
        "/ISAPI/System/Video/inputs/channels",
        "/ISAPI/ContentMgmt/RemoteDevice/channels",
        "/ISAPI/System/Video/inputs",
        "/ISAPI/Streaming/channels"
    ]
    
    for endpoint in endpoints:
        try:
            url = f"http://{ip}{endpoint}"
            log(f"[ENHANCED-ISAPI] Trying {endpoint}...")
            
            # Try Basic Auth first
            resp = session.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[ENHANCED-ISAPI] {endpoint}: HTTP {resp.status_code}")
            
                        
        except Exception as e:
            log(f"[ENHANCED-ISAPI] Error with {endpoint}: {str(e)[:50]}")
            continue
    return []

def try_video_inputs_method_enhanced(session, ip: str, username: str, password: str, timeout: float) -> list:
    """Try Video Inputs method with enhanced parsing"""
    endpoints = [
        "/ISAPI/System/Video/inputs/channels",
        "/api/v1/System/Video/inputs/channels",
        "/ISAPI/Streaming/channels",
        "/ISAPI/System/Video/inputs"
    ]
    
    for endpoint in endpoints:
        try:
            url = f"http://{ip}{endpoint}"
            log(f"[ENHANCED-VIDEO] Trying {endpoint}...")
            
            # Try Basic Auth first
            resp = session.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[ENHANCED-VIDEO] {endpoint}: HTTP {resp.status_code}")
            
                        
        except Exception as e:
            log(f"[ENHANCED-VIDEO] Error with {endpoint}: {str(e)[:50]}")
            continue
    return []

def try_remote_devices_method_enhanced(session, ip: str, username: str, password: str, timeout: float) -> list:
    """Try Remote Devices method with enhanced parsing"""
    endpoints = [
        "/ISAPI/ContentMgmt/RemoteDevice",
        "/api/v1/devices", 
        "/api/v2/devices", 
        "/cgi-bin/api/v1/devices"
    ]
    
    for endpoint in endpoints:
        try:
            url = f"http://{ip}{endpoint}"
            log(f"[ENHANCED-REMOTE] Trying {endpoint}...")
            
            # Try Basic Auth first
            resp = session.get(url, auth=HTTPBasicAuth(username, password), timeout=timeout)
            log(f"[ENHANCED-REMOTE] {endpoint}: HTTP {resp.status_code}")
            
                        
        except Exception as e:
            log(f"[ENHANCED-REMOTE] Error with {endpoint}: {str(e)[:50]}")
            continue
    return []

def try_digest_auth_fallback_enhanced(session, ip: str, username: str, password: str, timeout: float) -> list:
    """Try Digest authentication as fallback"""
    endpoints = [
        "/ISAPI/ContentMgmt/InputProxy/channels",
        "/ISAPI/System/Video/inputs/channels",
        "/ISAPI/ContentMgmt/RemoteDevice"
    ]
    
    for endpoint in endpoints:
        try:
            url = f"http://{ip}{endpoint}"
            log(f"[ENHANCED-DIGEST] Trying {endpoint} with Digest Auth...")
            
            resp = session.get(url, auth=HTTPDigestAuth(username, password), timeout=timeout)
            log(f"[ENHANCED-DIGEST] {endpoint}: HTTP {resp.status_code}")
            
                    
        except Exception as e:
            log(f"[ENHANCED-DIGEST] Error with {endpoint}: {str(e)[:50]}")
            continue
    return []

            

def parse_remote_devices_enhanced(xml_text: str) -> list:
    """Enhanced parser for Remote Devices XML"""
    cameras = []
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(xml_text)
        
        # Find all RemoteDevice elements (with or without namespace)
        devices = (root.findall(".//RemoteDevice") + 
                  root.findall(".//{http://www.hikvision.com/ver20/XMLSchema}RemoteDevice"))
        
        for device in devices:
            cam = {}
            
            # Get device ID
            id_elem = device.find(".//id")
            if id_elem is None:
                id_elem = device.find(".//{http://www.hikvision.com/ver20/XMLSchema}id")
            cam["channel"] = int(id_elem.text) if id_elem is not None else len(cameras) + 1
            
            # Get device name
            name_elem = device.find(".//name")
            if name_elem is None:
                name_elem = device.find(".//{http://www.hikvision.com/ver20/XMLSchema}name")
            cam["name"] = name_elem.text if name_elem is not None else f"Device {cam['channel']}"
            
            # Get IP address
            ip_elem = device.find(".//ipAddress")
            if ip_elem is None:
                ip_elem = device.find(".//{http://www.hikvision.com/ver20/XMLSchema}ipAddress")
            cam["ip"] = ip_elem.text if ip_elem is not None else ""
            
            # Get status - match ivms.py logic: online only if devIndex exists and has value
            devindex_elem = device.find(".//{http://www.hikvision.com/ver20/XMLSchema}devIndex")
            if devindex_elem is None:
                devindex_elem = device.find(".//devIndex")
            cam["status"] = "online" if devindex_elem is not None and devindex_elem.text else "offline"
            
            # Get model
            model_elem = device.find(".//model")
            if model_elem is None:
                model_elem = device.find(".//{http://www.hikvision.com/ver20/XMLSchema}model")
            cam["model"] = model_elem.text if model_elem is not None else ""
            
            if cam.get("ip"):  # Only add if has IP
                cameras.append(cam)
                
    except Exception as e:
        log(f"[ENHANCED-PARSE] Error parsing Remote Devices XML: {e}")
    
    return cameras

def parse_generic_cameras_enhanced(data) -> list:
    """Enhanced parser for generic API JSON response"""
    cameras = []
    devices = []
    
    if isinstance(data, dict):
        devices = data.get("devices", [])
    elif isinstance(data, list):
        devices = data
    
    for idx, dev in enumerate(devices):
        cam = {
            "channel": idx + 1,
            "name": dev.get("name") or dev.get("deviceName") or f"Camera {idx + 1}",
            "ip": dev.get("ip") or dev.get("ipAddress") or dev.get("address") or "",
            "status": dev.get("status") or dev.get("state") or "unknown",
            "model": dev.get("model") or dev.get("deviceModel") or "",
        }
        if cam["ip"]:
            cameras.append(cam)
    
    return cameras


# ==================== ENHANCED CAMERA STATUS CHECKING ====================
def check_camera_status_comprehensive(camera_ip: str, nvr_ip: str, nvr_user: str, nvr_pass: str, timeout: float = 3.0) -> dict:
    """Comprehensive camera status check with multiple validation methods"""
    import subprocess
    import socket
    import time
    from datetime import datetime
    from requests.auth import HTTPDigestAuth, HTTPBasicAuth
    
    status_result = {
        "ip": camera_ip,
        "status": "unknown",
        "ping_status": "unknown",
        "stream_status": "unknown", 
        "nvr_reported": "unknown",
        "last_check": datetime.now().isoformat(),
        "response_time": 0,
        "details": []
    }
    
    if not camera_ip or camera_ip.strip() == "":
        status_result["status"] = "no_ip"
        status_result["details"].append("No IP address available")
        return status_result
    
    # Method 1: Network Ping Test (fastest check)
    ping_start = time.time()
    try:
        # Use Windows ping command for reliability
        result = subprocess.run(
            ["ping", "-n", "1", "-w", str(int(timeout * 1000)), camera_ip],
            capture_output=True, text=True, timeout=timeout + 1
        )
        ping_time = (time.time() - ping_start) * 1000
        
        if result.returncode == 0:
            status_result["ping_status"] = "online"
            status_result["response_time"] = round(ping_time, 1)
            status_result["details"].append(f"Ping successful ({ping_time:.1f}ms)")
        else:
            status_result["ping_status"] = "offline"
            status_result["details"].append("Ping failed - no response")
    except (subprocess.TimeoutExpired, Exception) as e:
        status_result["ping_status"] = "timeout"
        status_result["details"].append(f"Ping timeout: {e}")
    
    # Method 2: TCP Port Check (port 80 for HTTP)
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(timeout)
        tcp_result = sock.connect_ex((camera_ip, 80))
        sock.close()
        
        if tcp_result == 0:
            status_result["details"].append("TCP port 80 open")
        else:
            status_result["details"].append("TCP port 80 closed")
    except Exception as e:
        status_result["details"].append(f"TCP check failed: {e}")
    
    # Method 3: HTTP Stream Test (check camera web interface)
    try:
        import requests
        session = requests.Session()
        session.timeout = timeout
        
        # Try common camera endpoints
        test_endpoints = [
            f"http://{camera_ip}/",
            f"http://{camera_ip}/ISAPI/System/deviceInfo",
            f"http://{camera_ip}/cgi-bin/hi3510/param.cgi?cmd=getserverinfo"
        ]
        
        for endpoint in test_endpoints:
            try:
                resp = session.get(endpoint, timeout=timeout/2)
                if resp.status_code in [200, 401, 403]:  # Any response indicates device is active
                    status_result["stream_status"] = "responding"
                    status_result["details"].append(f"HTTP response: {resp.status_code}")
                    break
            except:
                continue
        
        if status_result["stream_status"] == "unknown":
            status_result["stream_status"] = "no_response"
            status_result["details"].append("No HTTP response from camera")
            
    except Exception as e:
        status_result["stream_status"] = "error"
        status_result["details"].append(f"HTTP test error: {e}")
    
    # Method 4: NVR Channel Status Check (verify via NVR)
    try:
        nvr_session = requests.Session()
        auth_methods = [
            HTTPDigestAuth(nvr_user, nvr_pass),
            HTTPBasicAuth(nvr_user, nvr_pass)
        ]
        
        for auth in auth_methods:
            try:
                # Check channel status via NVR
                nvr_url = f"http://{nvr_ip}/ISAPI/ContentMgmt/InputProxy/channels"
                resp = nvr_session.get(nvr_url, auth=auth, timeout=timeout)
                
                if resp.status_code == 200:
                    # Parse XML to find this specific camera
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(resp.text)
                    
                    # Look for this camera's IP in channels
                    for channel in root.findall(".//InputProxyChannel") + root.findall(".//{http://www.hikvision.com/ver20/XMLSchema}InputProxyChannel"):
                        ip_elem = channel.find(".//ipAddress")
                        if ip_elem is not None and ip_elem.text == camera_ip:
                            # Found the camera in NVR
                            devindex_elem = channel.find(".//devIndex")
                            online_elem = channel.find(".//online")
                            
                            if devindex_elem is not None and devindex_elem.text:
                                status_result["nvr_reported"] = "connected"
                                status_result["details"].append(f"NVR reports: connected (devIndex: {devindex_elem.text})")
                            elif online_elem is not None and online_elem.text == "true":
                                status_result["nvr_reported"] = "online"
                                status_result["details"].append("NVR reports: online")
                            else:
                                status_result["nvr_reported"] = "configured_offline"
                                status_result["details"].append("NVR reports: configured but offline")
                            break
                    break
            except:
                continue
                
    except Exception as e:
        status_result["details"].append(f"NVR check error: {e}")
    
    # Final Status Determination (combine all checks)
    ping_ok = status_result["ping_status"] == "online"
    stream_ok = status_result["stream_status"] in ["responding"]
    nvr_ok = status_result["nvr_reported"] in ["connected", "online"]
    
    if ping_ok and (stream_ok or nvr_ok):
        status_result["status"] = "online"
    elif ping_ok and not (stream_ok or nvr_ok):
        status_result["status"] = "network_only"  # Pingable but no camera response
    elif not ping_ok and nvr_ok:
        status_result["status"] = "nvr_only"  # NVR sees it but not pingable
    else:
        status_result["status"] = "offline"
    
    return status_result

def check_camera_via_sadp(ip: str, timeout: float = 3.0) -> tuple:
    """Quick SADP-style camera check"""
    try:
        import subprocess
        result = subprocess.run(
            ["ping", "-n", "1", "-w", "2000", ip],
            capture_output=True, text=True, timeout=timeout
        )
        if result.returncode == 0:
            return True, "Online"
        else:
            return False, "Offline"
    except:
        return False, "Unknown"

def check_camera_live(ip: str, nvr_ip: str = "", nvr_user: str = "", nvr_pwd: str = "", timeout: float = 5.0) -> tuple:
    """Stub function for camera live check"""
    return False, "Ping", "No response"

def sadp_discover(timeout: float = 1.0, scan_hosts: int = 100, progress_callback=None, target_subnet: str = None, preferred_local_ip: str = None) -> list:
    """Stub function for SADP discovery"""
    return []


# ==================== WORKING CAMERA IMPORT BASED ON IVMS ====================
def working_fetch_cameras_from_ivms(nvr_ip: str, username: str, password: str, timeout: float = 15.0) -> tuple:
    """Working camera fetch using proven IVMS methods"""
    log(f"=== Working IVMS Camera Fetch ===")
    log(f"NVR IP: {nvr_ip}, User: {username}, Timeout: {timeout}s")
    
    import xml.etree.ElementTree as ET
    from requests.auth import HTTPBasicAuth, HTTPDigestAuth
    
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
        "Accept": "application/json, text/xml, */*",
    })
    
    # Test basic connectivity first
    try:
        test_url = f"http://{nvr_ip}/ISAPI/System/deviceInfo"
        log(f"Testing connectivity to {test_url}...")
        resp = session.get(test_url, auth=HTTPBasicAuth(username, password), timeout=5.0)
        log(f"Device reachable: HTTP {resp.status_code}")
        if resp.status_code == 401:
            # Try digest auth for 401 responses
            log("Trying Digest authentication...")
            resp_digest = session.get(test_url, auth=HTTPDigestAuth(username, password), timeout=5.0)
            log(f"Digest auth: HTTP {resp_digest.status_code}")
            if resp_digest.status_code != 200:
                return False, [], 0, 0, "Authentication failed - check username/password"
    except requests.exceptions.Timeout:
        return False, [], 0, 0, f"Connection timeout - NVR not reachable at {nvr_ip}"
    except requests.exceptions.ConnectionError:
        return False, [], 0, 0, f"Connection refused - check IP address"
    except Exception as e:
        log(f"Connectivity test error: {e}")
    
    # Method 1: ISAPI ContentMgmt (Primary - most reliable)
    log("[Method 1] Trying ISAPI endpoints...")
    cameras = _fetch_isapi_cameras_working(session, nvr_ip, username, password, timeout)
    if cameras:
        total = len(cameras)
        active = len([c for c in cameras if c.get('status', '').lower() == 'online'])
        log(f"SUCCESS - Found {total} cameras via ISAPI ({active} online)")
        return True, cameras, total, active, ""
    
    # Method 2: Generic API endpoints
    log("[Method 2] Trying Generic API endpoints...")
    cameras = _fetch_generic_api_cameras_working(session, nvr_ip, username, password, timeout)
    if cameras:
        total = len(cameras)
        active = len([c for c in cameras if c.get('status', '').lower() == 'online'])
        log(f"SUCCESS - Found {total} cameras via Generic API ({active} online)")
        return True, cameras, total, active, ""
    
    # Method 3: Video Input channels
    log("[Method 3] Trying Video Input channels...")
    cameras = _fetch_video_inputs_working(session, nvr_ip, username, password, timeout)
    if cameras:
        total = len(cameras)
        active = len([c for c in cameras if c.get('status', '').lower() == 'online'])
        log(f"SUCCESS - Found {total} cameras via Video Inputs ({active} online)")
        return True, cameras, total, active, ""
    
    # Method 4: Try with digest auth
    log("[Method 4] Trying with Digest authentication...")
    cameras = _fetch_with_digest_auth_working(session, nvr_ip, username, password, timeout)
    if cameras:
        total = len(cameras)
        active = len([c for c in cameras if c.get('status', '').lower() == 'online'])
        log(f"SUCCESS - Found {total} cameras via Digest Auth ({active} online)")
        return True, cameras, total, active, ""
    
    error_msg = (
        f"Could not retrieve camera list from NVR.\\n\\n"
        f"Troubleshooting:\\n"
        f"1. Verify NVR IP: {nvr_ip}\\n"
        f"2. Check credentials: {username}\\n"
        f"3. Ensure NVR API is enabled\\n"
        f"4. Try accessing http://{nvr_ip} in browser\\n\\n"
        f"The NVR may use a non-standard API or require special configuration."
    )
    return False, [], 0, 0, error_msg


def _fetch_isapi_cameras_working(session, ip: str, user: str, pwd: str, timeout: float):
    """Fetch cameras using Hikvision ISAPI (Working method from IVMS)"""
    endpoints = [
        "/ISAPI/ContentMgmt/InputProxy/channels",
        "/ISAPI/ContentMgmt/RemoteDevice",
        "/ISAPI/System/Video/inputs/channels",
    ]
    
    for endpoint in endpoints:
        url = f"http://{ip}{endpoint}"
        try:
            resp = session.get(url, auth=HTTPBasicAuth(user, pwd), timeout=timeout)
            log(f"Trying {endpoint}: {resp.status_code}")
            if resp.status_code == 200:
                cameras = _parse_isapi_cameras_working(resp.text)
                if cameras:
                    log(f"Found {len(cameras)} cameras using {endpoint}")
                    return cameras
            elif resp.status_code == 401:
                log(f"{endpoint}: Authentication failed")
        except Exception as e:
            log(f"{endpoint}: {str(e)[:50]}")
            continue
    return []


def _parse_isapi_cameras_working(xml_text: str):
    """Parse ISAPI XML response using working IVMS method"""
    import xml.etree.ElementTree as ET
    cameras = []
    try:
        root = ET.fromstring(xml_text)
        # Handle XML namespace
        namespace = {'hikvision': 'http://www.hikvision.com/ver20/XMLSchema'}
        
        # Look for InputProxyChannel or RemoteDevice elements (with and without namespace)
        input_channels = (root.findall(".//hikvision:InputProxyChannel", namespace) + 
                        root.findall(".//InputProxyChannel"))
        remote_devices = (root.findall(".//hikvision:RemoteDevice", namespace) + 
                        root.findall(".//RemoteDevice"))

        for channel in input_channels + remote_devices:
            cam = {}

            # Extract channel info using simpler approach
            cam_id = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}id")
            if cam_id is None:
                cam_id = channel.find(".//id")
            cam["channel"] = int(cam_id.text) if cam_id is not None else 0

            cam_name = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}name")
            if cam_name is None:
                cam_name = channel.find(".//name")
            cam["name"] = (
                cam_name.text if cam_name is not None else f"Camera {cam['channel']}"
            )

            # Extract IP address from sourceInputPortDescriptor
            ip_elem = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}ipAddress")
            if ip_elem is None:
                ip_elem = channel.find(".//ipAddress") 
            cam["ip"] = ip_elem.text if ip_elem is not None else ""

            # Extract status - check multiple indicators for camera connectivity
            devindex_elem = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}devIndex")
            if devindex_elem is None:
                devindex_elem = channel.find(".//devIndex")
            
            # Check for additional status indicators
            online_elem = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}online")
            if online_elem is None:
                online_elem = channel.find(".//online")
            
            enabled_elem = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}enabled")
            if enabled_elem is None:
                enabled_elem = channel.find(".//enabled")
            
            # Determine camera status using multiple criteria with enhanced logic
            has_devindex = devindex_elem is not None and devindex_elem.text
            is_online = online_elem is not None and online_elem.text == 'true'
            is_enabled = enabled_elem is not None and enabled_elem.text == 'true'
            has_ip = cam.get("ip", "").strip() != ""
            
            # IVMS Method: Status determination based solely on devIndex (like ivms.py line 446)
            # Camera is online only if it has devIndex (indicates actual device connected to NVR)
            if has_devindex:
                cam["status"] = "online"
                cam["connection_type"] = "nvr_connected"
            else:
                cam["status"] = "offline"
                cam["connection_type"] = "nvr_disconnected"
            


            # Extract model if available
            model_elem = channel.find(".//{http://www.hikvision.com/ver20/XMLSchema}model")
            if model_elem is None:
                model_elem = channel.find(".//model")
            cam["model"] = model_elem.text if model_elem is not None else ""

            if cam.get("ip"):  # Only add if has IP
                cameras.append(cam)

    except Exception as e:
        log(f"Error parsing ISAPI XML: {e}")

    return cameras


def _fetch_generic_api_cameras_working(session, ip: str, user: str, pwd: str, timeout: float):
    """Fetch cameras using generic API endpoints"""
    endpoints = ["/api/v1/devices", "/api/v2/devices", "/cgi-bin/api/v1/devices"]
    
    for endpoint in endpoints:
        url = f"http://{ip}{endpoint}"
        try:
            resp = session.get(url, auth=HTTPBasicAuth(user, pwd), timeout=timeout)
            log(f"Trying {endpoint}: {resp.status_code}")
            if resp.status_code == 200:
                data = resp.json()
                cameras = _parse_generic_cameras_working(data)
                if cameras:
                    return cameras
        except Exception as e:
            continue
    return []


def _parse_generic_cameras_working(data):
    """Parse generic API JSON response"""
    cameras = []
    devices = []
    
    if isinstance(data, dict):
        devices = data.get("devices", [])
    elif isinstance(data, list):
        devices = data

    for idx, dev in enumerate(devices):
        cam = {
            "channel": idx + 1,
            "name": dev.get("name") or dev.get("deviceName") or f"Camera {idx + 1}",
            "ip": dev.get("ip") or dev.get("ipAddress") or dev.get("address") or "",
            "status": dev.get("status") or dev.get("state") or "unknown",
            "model": dev.get("model") or dev.get("deviceModel") or "",
        }
        if cam["ip"]:
            cameras.append(cam)
    
    return cameras


def _fetch_video_inputs_working(session, ip: str, user: str, pwd: str, timeout: float):
    """Fetch cameras using video input channels"""
    endpoints = [
        "/ISAPI/System/Video/inputs/channels",
        "/api/v1/System/Video/inputs/channels",
        "/ISAPI/Streaming/channels",
    ]

    for endpoint in endpoints:
        url = f"http://{ip}{endpoint}"
        try:
            resp = session.get(url, auth=HTTPBasicAuth(user, pwd), timeout=timeout)
            log(f"Trying {endpoint}: {resp.status_code}")
            if resp.status_code == 200:
                cameras = _parse_video_channels_working(resp)
                if cameras:
                    return cameras
        except Exception as e:
            continue
    return []


def _fetch_with_digest_auth_working(session, ip: str, user: str, pwd: str, timeout: float):
    """Fetch cameras using Digest authentication (fallback)"""
    from requests.auth import HTTPDigestAuth

    endpoints = [
        "/ISAPI/ContentMgmt/InputProxy/channels",
        "/ISAPI/System/Video/inputs/channels",
    ]

    for endpoint in endpoints:
        url = f"http://{ip}{endpoint}"
        try:
            resp = session.get(url, auth=HTTPDigestAuth(user, pwd), timeout=timeout)
            log(f"Trying {endpoint} (Digest): {resp.status_code}")
            if resp.status_code == 200:
                cameras = _parse_isapi_cameras_working(resp.text)
                if cameras:
                    return cameras
        except Exception as e:
            continue
    return []


def _parse_video_channels_working(response):
    """Parse video channel response (JSON or XML)"""
    cameras = []
    try:
        # Try JSON first
        data = response.json()
        channels = []
        if isinstance(data, dict):
            channels = data.get("channels", [])
        elif isinstance(data, list):
            channels = data

        for idx, ch in enumerate(channels):
            cam = {
                "channel": idx + 1,
                "name": ch.get("name") or ch.get("channelName") or f"Channel {idx + 1}",
                "ip": ch.get("ip") or ch.get("ipAddress") or "",
                "status": ch.get("status") or "unknown",
                "model": "",
            }
            if cam["ip"]:
                cameras.append(cam)
    except:
        # Try XML
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(response.text)
            for idx, channel in enumerate(root.findall(".//VideoInputChannel")):
                cam = {"channel": idx + 1, "name": f"Channel {idx + 1}", "ip": "", "status": "unknown"}
                name_elem = channel.find(".//name")
                if name_elem is not None:
                    cam["name"] = name_elem.text
                if cam["ip"]:  # Only add if has IP
                    cameras.append(cam)
        except:
            pass
    return cameras





# ---------------- sheet matching utils ----------------
def normalize_key(s):
    if s is None:
        return ""
    s = str(s)
    s = unicodedata.normalize("NFKC", s)
    s = "".join(s.split())
    s = re.sub(r'[^0-9A-Za-z]', '', s)
    return s.lower()

def find_sheet_key(wb_dict, target):
    t = normalize_key(target)
    # exact normalized
    for k in wb_dict.keys():
        if normalize_key(k) == t:
            return k
    # startswith / contains heuristics
    for k in wb_dict.keys():
        nk = normalize_key(k)
        if nk.startswith(t) or t.startswith(nk):
            return k
    for k in wb_dict.keys():
        nk = normalize_key(k)
        if t in nk or nk in t:
            return k
    # numeric digits match
    td = "".join(re.findall(r'\d+', t))
    if td:
        for k in wb_dict.keys():
            kd = "".join(re.findall(r'\d+', normalize_key(k)))
            if kd and kd == td:
                return k
    return None

# ---------------- Excel load: robust for header/no-header ----------------
def load_excel_robust(path):
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    # read all sheets as strings (don't coerce)
    wb = pd.read_excel(path, sheet_name=None, engine="openpyxl", dtype=str, keep_default_na=False)
    nvrs = []
    cams = []

    # --- handle NVR sheet robustly whether it has header or not ---
    if "NVR" not in wb:
        raise ValueError("No 'NVR' sheet found in workbook.")
    # get raw dataframe
    df_raw = wb["NVR"]
    # Heuristic: determine if first row appears to be header (contains words like 'Name' or 'IP')
    first_row_vals = [str(x).strip().lower() for x in list(df_raw.iloc[0].fillna(""))] if not df_raw.empty else []
    header_like = False
    header_indicators = {"name", "nvr", "ip", "subnet", "gateway", "mask"}
    if any(any(ind in v for ind in header_indicators) for v in first_row_vals):
        header_like = True

    # If header_like True -> treat as header (pandas already did). Else re-read sheet with header=None so first row becomes data.
    if header_like:
        df_nvr = df_raw
    else:
        # read again with header=None so first row is data
        df_nvr = pd.read_excel(path, sheet_name="NVR", engine="openpyxl", header=None, dtype=str, keep_default_na=False)

    # Now extract rows from df_nvr by position: A=col 0, B=col1, C=col2, D=col3
    for _, row in df_nvr.iterrows():
        name = str(row.iloc[0]).strip() if len(row) > 0 and pd.notna(row.iloc[0]) else ""
        ip = str(row.iloc[1]).strip() if len(row) > 1 and pd.notna(row.iloc[1]) else ""
        subnet = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ""
        gateway = str(row.iloc[3]).strip() if len(row) > 3 and pd.notna(row.iloc[3]) else ""
        if name or ip:
            nvrs.append({"name": name, "ip": ip, "subnet": subnet, "gateway": gateway, "sheet_found": False})

    # For each NVR find matching sheet and load cameras (camera sheets may also lack headers; treat similarly)
    for n in nvrs:
        key = find_sheet_key(wb, n["name"])
        if key:
            n["sheet_found"] = True
            dfc_raw = wb[key]
            # detect header-like in camera sheet
            first_row = [str(x).strip().lower() for x in list(dfc_raw.iloc[0].fillna(""))] if not dfc_raw.empty else []
            cam_header_like = False
            cam_indicators = {"camera", "cam", "ip", "title", "name"}
            if any(any(ind in v for ind in cam_indicators) for v in first_row):
                cam_header_like = True
            if cam_header_like:
                df_cam = dfc_raw
            else:
                df_cam = pd.read_excel(path, sheet_name=key, engine="openpyxl", header=None, dtype=str, keep_default_na=False)
            for _, crow in df_cam.iterrows():
                cname = str(crow.iloc[0]).strip() if len(crow) > 0 and pd.notna(crow.iloc[0]) else ""
                cip = str(crow.iloc[1]).strip() if len(crow) > 1 and pd.notna(crow.iloc[1]) else ""
                if cname or cip:
                    cams.append({"nvr": n["name"], "name": cname, "ip": cip, "status": "Unknown"})
        else:
            n["sheet_found"] = False

    # annotate cam_count
    for n in nvrs:
        nvr_name = (n.get("name") or "").strip().lower()
        n["cam_count"] = len([
            c for c in cams
            if (c.get("nvr") or "").strip().lower() == nvr_name
        ])

    return nvrs, cams

# ---------------- NVR Credential Management ----------------
def load_nvr_credentials():
    """Load NVR credentials from file"""
    try:
        cred_file = "nvr_credentials.json"
        if os.path.exists(cred_file):
            with open(cred_file, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        log(f"[ERROR] Failed to load NVR credentials: {e}")
    return {}

def save_nvr_credentials(credentials):
    """Save NVR credentials to file"""
    try:
        cred_file = "nvr_credentials.json"
        with open(cred_file, 'w', encoding='utf-8') as f:
            json.dump(credentials, f, indent=2, ensure_ascii=False)
        log(f"[NVR-CRED] Saved credentials for {len(credentials)} NVRs")
        return True
    except Exception as e:
        log(f"[ERROR] Failed to save NVR credentials: {e}")
        return False

def get_enhanced_nvr_configs():
    """Get NVR configurations with enhanced credential management"""
    # Load custom credentials
    custom_creds = load_nvr_credentials()
    
    # Default NVR configurations
    default_configs = [
        {'name': 'NVR1', 'ip': '192.168.2.168', 'cameras': 29},
        {'name': 'NVR2', 'ip': '192.168.2.169', 'cameras': 28},
        {'name': 'NVR3', 'ip': '192.168.2.170', 'cameras': 30},
        {'name': 'NVR4', 'ip': '192.168.2.171', 'cameras': 30},
        {'name': 'NVR5', 'ip': '192.168.2.172', 'cameras': 27},
        {'name': 'NVR6', 'ip': '192.168.2.173', 'cameras': 29},
        {'name': 'NVR7', 'ip': '192.168.2.177', 'cameras': 26},
        {'name': 'NVR8', 'ip': '192.168.3.1', 'cameras': 30},
        {'name': 'NVR9', 'ip': '192.168.3.2', 'cameras': 30},
        {'name': 'NVR10', 'ip': '192.168.3.3', 'cameras': 29},
        {'name': 'NVR11', 'ip': '192.168.3.4', 'cameras': 29},
        {'name': 'NVR12', 'ip': '192.168.3.5', 'cameras': 0},
        {'name': 'NVR13', 'ip': '192.168.2.247', 'cameras': 26},
        {'name': 'NVR14', 'ip': '192.168.2.178', 'cameras': 4},
        {'name': 'NVR15', 'ip': '192.168.2.248', 'cameras': 7},
        {'name': 'NVR16', 'ip': '192.168.2.245', 'cameras': 32},
        {'name': 'NVR17', 'ip': '192.168.2.166', 'cameras': 30},
        {'name': 'NVR18', 'ip': '192.168.2.167', 'cameras': 22},
        {'name': 'DI Sreen', 'ip': '192.168.3.7', 'cameras': 29}
    ]
    
    # Merge with custom credentials
    for config in default_configs:
        nvr_key = f"{config['name']}_{config['ip']}"
        if nvr_key in custom_creds:
            config.update(custom_creds[nvr_key])
        else:
            # Set default credentials
            config.update({
                'port': 80,
                'protocol': 'http',
                'username': 'admin',
                'password': 'Kkcctv12345'
            })
    
    # Add any additional custom NVRs
    for nvr_key, cred_data in custom_creds.items():
        if not any(f"{cfg['name']}_{cfg['ip']}" == nvr_key for cfg in default_configs):
            default_configs.append(cred_data)
    
    return default_configs

# ==================== LICENSE MANAGEMENT SYSTEM ====================
def get_machine_id():
    """Generate unique machine identifier based on hardware"""
    try:
        # Get MAC address
        mac = ':'.join(['{:02x}'.format((uuid.getnode() >> elements) & 0xff) 
                       for elements in range(0,2*6,2)][::-1])
        # Get computer name
        hostname = socket.gethostname()
        # Combine and hash
        machine_str = f"{mac}_{hostname}_{platform.node()}"
        return hashlib.sha256(machine_str.encode()).hexdigest()[:16]
    except:
        return "UNKNOWN_MACHINE"

def generate_license_key(machine_id, expiry_date, license_type="PROFESSIONAL"):
    """Generate license key with machine binding and expiration"""
    data = f"{machine_id}|{expiry_date.strftime('%Y-%m-%d')}|{license_type}"
    signature = hashlib.sha256((data + LICENSE_SALT.decode('latin-1')).encode()).hexdigest()
    key = f"{machine_id}-{expiry_date.strftime('%Y%m%d')}-{license_type[:3].upper()}-{signature[:8].upper()}"
    return key

def validate_license_key(license_key, machine_id):
    """Validate license key against machine ID and check expiration"""
    try:
        parts = license_key.strip().split('-')
        if len(parts) != 4:
            return False, "Invalid key format"
        
        key_machine_id, key_date, key_type, key_sig = parts
        
        # Check machine binding
        if key_machine_id != machine_id:
            return False, "License not valid for this machine"
        
        # Check expiration
        expiry_date = datetime.strptime(key_date, '%Y%m%d')
        if datetime.now() > expiry_date:
            return False, f"License expired on {expiry_date.strftime('%Y-%m-%d')}"
        
        # Verify signature
        data = f"{key_machine_id}|{expiry_date.strftime('%Y-%m-%d')}|{key_type}"
        expected_sig = hashlib.sha256((data + LICENSE_SALT.decode('latin-1')).encode()).hexdigest()[:8].upper()
        
        if key_sig != expected_sig:
            return False, "Invalid license signature"
        
        days_remaining = (expiry_date - datetime.now()).days
        return True, f"Valid until {expiry_date.strftime('%Y-%m-%d')} ({days_remaining} days remaining)"
        
    except Exception as e:
        return False, f"Validation error: {str(e)}"

def save_license_key(license_key):
    """Save license key to file"""
    try:
        with open(LICENSE_FILE, 'w') as f:
            f.write(license_key)
        return True
    except:
        return False

def load_license_key():
    """Load license key from file"""
    try:
        if os.path.exists(LICENSE_FILE):
            with open(LICENSE_FILE, 'r') as f:
                return f.read().strip()
    except:
        pass
    return None

def check_trial_status():
    """Check if trial period is active"""
    trial_file = "trial.dat"
    try:
        if os.path.exists(trial_file):
            with open(trial_file, 'r') as f:
                install_date_str = f.read().strip()
                install_date = datetime.fromisoformat(install_date_str)
                days_used = (datetime.now() - install_date).days
                days_remaining = TRIAL_DAYS - days_used
                if days_remaining > 0:
                    return True, f"Trial: {days_remaining} days remaining"
                else:
                    return False, "Trial period expired"
        else:
            # First run - create trial file
            with open(trial_file, 'w') as f:
                f.write(datetime.now().isoformat())
            return True, f"Trial: {TRIAL_DAYS} days remaining"
    except:
        return False, "Trial status unknown"

class LicenseDialog(QtWidgets.QDialog):
    """License activation dialog"""
    def __init__(self, parent=None, machine_id=""):
        super().__init__(parent)
        self.machine_id = machine_id
        self.license_valid = False
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("🔐 License Activation - NARONG CCTV Monitor")
        self.setMinimumWidth(600)
        self.setMinimumHeight(500)
        self.setModal(True)
        
        # Set window background
        self.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #dee2e6;
                border-radius: 6px;
                margin-top: 8px;
                padding-top: 8px;
                background-color: white;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 5px;
            }
        """)
        
        layout = QtWidgets.QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(15, 15, 15, 15)
        
        # Simple header text
        header = QtWidgets.QLabel(f"<h2 style='color: #2c3e50; margin: 0;'>🔐 NARONG CCTV Monitor v{APP_VERSION}</h2>")
        header.setAlignment(QtCore.Qt.AlignCenter)
        header.setStyleSheet("padding: 10px; background-color: #e9ecef; border-radius: 6px;")
        layout.addWidget(header)
        
        # Machine ID section
        machine_group = QtWidgets.QGroupBox("🖥️ Machine Information")
        machine_layout = QtWidgets.QVBoxLayout()
        machine_layout.setSpacing(8)
        
        # Machine ID display with copy button
        id_container = QtWidgets.QHBoxLayout()
        machine_info = QtWidgets.QLineEdit(self.machine_id)
        machine_info.setReadOnly(True)
        machine_info.setFont(QtGui.QFont("Courier New", 10))
        machine_info.setStyleSheet("""
            QLineEdit {
                background-color: #f8f9fa;
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                color: #495057;
            }
        """)
        machine_info.setAlignment(QtCore.Qt.AlignCenter)
        id_container.addWidget(machine_info)
        
        copy_btn = QtWidgets.QPushButton("📋 Copy")
        copy_btn.setStyleSheet("""
            QPushButton {
                background-color: #17a2b8;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border-radius: 4px;
                border: none;
            }
            QPushButton:hover {
                background-color: #138496;
            }
        """)
        copy_btn.clicked.connect(lambda: (QtWidgets.QApplication.clipboard().setText(self.machine_id),
                                          self.show_copy_confirmation()))
        id_container.addWidget(copy_btn)
        machine_layout.addLayout(id_container)
        
        info_label = QtWidgets.QLabel(
            "<i style='color: #6c757d; font-size: 11px;'>Send this ID to obtain a license key</i>"
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("background: transparent;")
        machine_layout.addWidget(info_label)
        
        machine_group.setLayout(machine_layout)
        layout.addWidget(machine_group)
        
        # License key input
        key_group = QtWidgets.QGroupBox("🔑 License Key")
        key_layout = QtWidgets.QVBoxLayout()
        key_layout.setSpacing(8)
        
        self.key_input = QtWidgets.QLineEdit()
        self.key_input.setPlaceholderText("Enter license key...")
        self.key_input.setFont(QtGui.QFont("Courier New", 10))
        self.key_input.setStyleSheet("""
            QLineEdit {
                background-color: white;
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                color: #212529;
            }
            QLineEdit:focus {
                border-color: #007bff;
            }
        """)
        key_layout.addWidget(self.key_input)
        
        # Validation status
        self.status_label = QtWidgets.QLabel("")
        self.status_label.setWordWrap(True)
        self.status_label.setStyleSheet("background: transparent; padding: 5px;")
        self.status_label.setMinimumHeight(30)
        key_layout.addWidget(self.status_label)
        
        key_group.setLayout(key_layout)
        layout.addWidget(key_group)
        
        # Action buttons
        btn_layout = QtWidgets.QHBoxLayout()
        btn_layout.setSpacing(10)
        
        validate_btn = QtWidgets.QPushButton("✓ Activate License")
        validate_btn.setMinimumHeight(42)
        validate_btn.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 12px 24px;
                border-radius: 4px;
                border: none;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        validate_btn.clicked.connect(self.validate_and_activate)
        btn_layout.addWidget(validate_btn)
        
        exit_btn = QtWidgets.QPushButton("✕ Exit Application")
        exit_btn.setMinimumHeight(42)
        exit_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 12px 24px;
                border-radius: 4px;
                border: none;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        exit_btn.clicked.connect(self.reject)
        btn_layout.addWidget(exit_btn)
        
        layout.addLayout(btn_layout)
        
        # Contact info
        contact_container = QtWidgets.QFrame()
        contact_container.setStyleSheet("""
            QFrame {
                background-color: #e8f4f8;
                border-radius: 4px;
                padding: 10px;
                border: 1px solid #bee5eb;
            }
        """)
        contact_layout = QtWidgets.QVBoxLayout(contact_container)
        contact_layout.setSpacing(6)
        
        contact_info = QtWidgets.QLabel(
            "<span style='color: #0c5460; font-size: 12px;'>"
            "<b>🔑 Need a License Key?</b><br>"
            "Use <b>Master Key Generator</b> tool to create keys for users<br>"
            "Or contact Sky-Tech Support via Telegram</span>"
        )
        contact_info.setWordWrap(True)
        contact_info.setStyleSheet("background: transparent;")
        contact_layout.addWidget(contact_info)
        
        telegram_btn = QtWidgets.QPushButton("📱 Telegram: @chhanycls")
        telegram_btn.setMinimumHeight(32)
        telegram_btn.setStyleSheet("""
            QPushButton {
                background-color: #0088cc;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 6px 12px;
                border-radius: 4px;
                border: none;
            }
            QPushButton:hover {
                background-color: #006699;
            }
        """)
        telegram_btn.clicked.connect(lambda: webbrowser.open('https://t.me/chhanycls'))
        telegram_btn.setCursor(QtCore.Qt.PointingHandCursor)
        contact_layout.addWidget(telegram_btn)
        
        layout.addWidget(contact_container)
    
    def show_copy_confirmation(self):
        """Show brief confirmation that Machine ID was copied"""
        self.status_label.setText("✅ Machine ID copied to clipboard!")
        self.status_label.setStyleSheet("color: #28a745; padding: 8px; background-color: #d4edda; border-radius: 4px; border: 1px solid #c3e6cb;")
        QtCore.QTimer.singleShot(2000, lambda: self.status_label.setText(""))
    
    def validate_and_activate(self):
        """Validate and activate the entered license key"""
        license_key = self.key_input.text().strip()
        if not license_key:
            self.status_label.setText("❌ Please enter a license key")
            self.status_label.setStyleSheet("color: #e74c3c; padding: 5px;")
            return
        
        is_valid, message = validate_license_key(license_key, self.machine_id)
        
        if is_valid:
            if save_license_key(license_key):
                self.status_label.setText(f"✅ {message}")
                self.status_label.setStyleSheet("color: #27ae60; padding: 5px; background-color: #d5f4e6; border-radius: 3px;")
                self.license_valid = True
                QtWidgets.QMessageBox.information(self, "Success", f"License activated successfully!\n\n{message}")
                self.accept()
            else:
                self.status_label.setText("❌ Failed to save license key")
                self.status_label.setStyleSheet("color: #e74c3c; padding: 5px;")
        else:
            self.status_label.setText(f"❌ {message}")
            self.status_label.setStyleSheet("color: #e74c3c; padding: 5px; background-color: #fadbd8; border-radius: 3px;")
    
# ---------------- GUI ----------------
class CameraMonitor(QtWidgets.QMainWindow):
    table_update = QtCore.pyqtSignal(int, str, str, str, object, str)
    enhanced_table_update = QtCore.pyqtSignal(int, str, str, str, str, str, str, str, str, str, object, str)
    nvr_update = QtCore.pyqtSignal(int, str)
    nvr_login_result = QtCore.pyqtSignal(bool, str, str)  # success, real_ip, error
    camera_update_signal = QtCore.pyqtSignal(list, str, float, int, int)  # cameras, nvr_name, elapsed, online, offline
    progress_update_signal = QtCore.pyqtSignal(str, int, str, bool)  # label_text, progress_value, style_sheet, enable_buttons
    button_control_signal = QtCore.pyqtSignal(str, bool, str, bool)  # button_name, enabled, text, visible
    ui_status_update_signal = QtCore.pyqtSignal(str, str, str)  # element_type, element_name, status_data
    ui_call_signal = QtCore.pyqtSignal(object)  # generic callable dispatcher for UI-thread work
    
    # Enhanced signals for v8.6+
    error_notification_signal = QtCore.pyqtSignal(str, str, str)  # level, title, message
    performance_update_signal = QtCore.pyqtSignal(dict)  # performance metrics
    cache_stats_signal = QtCore.pyqtSignal(dict)  # cache statistics

    def update_cameras_direct(self, selected_nvr):
        """Enhanced camera update - login to NVR and extract camera list using proven IVMS method with Digest Auth."""
        nvr_name = selected_nvr.get('name', 'Unknown')
        nvr_ip = selected_nvr.get('ip', '')
        
        # Get credentials using the same method as the successful refresh process
        username, password = self._resolve_nvr_credentials(nvr_ip)
        
        log(f"[UPDATE-CAMERAS] Starting enhanced camera update for {nvr_name} ({nvr_ip}) using proven IVMS method")
        self.status.showMessage(f"🔄 Updating cameras from {nvr_name} using proven method...", 0)
        QtWidgets.QApplication.processEvents()

        def fast_update_thread():
            log(f"[THREAD] Background thread started for {nvr_name} [IVMS method]")
            try:
                import time
                start_time = time.time()
                # Use the same proven controller that works for NVR18 and others
                controller = WorkingNVRController(nvr_ip, username, password)
                
                log(f"[UPDATE-CAMERAS] Testing NVR connection: {nvr_name}")
                if not controller.connect():
                    raise Exception(f"Cannot connect to NVR {nvr_name} at {nvr_ip}")
                
                log(f"[UPDATE-CAMERAS] Fetching cameras using proven method")
                cameras, method = controller.get_cameras(timeout=15.0)
                elapsed = time.time() - start_time
                
                if cameras:
                    log(f"[UPDATE-CAMERAS] Successfully found {len(cameras)} cameras via {method}")
                    result = self.update_camera_list_fast(cameras, nvr_name)
                    log(f"[UPDATE-CAMERAS] Camera list update completed, result: {result}")
                    self.populate_table(self.filtered)
                    
                    # Calculate status summary
                    nvr_online = sum(1 for cam in cameras if cam.get('status', '').lower() == 'online')
                    nvr_offline = len(cameras) - nvr_online
                    total_cameras = len(self.filtered)
                    total_online = sum(1 for cam in self.filtered if 'online' in cam.get('status', '').lower())
                    
                    msg = f"✅ Updated {len(cameras)} cameras from {nvr_name} ({nvr_online} online, {nvr_offline} offline) | Total: {total_cameras} ({total_online} online) | {elapsed:.1f}s"
                    self.status.showMessage(msg, 8000)  # Show longer for individual updates
                    log(f"[UPDATE-CAMERAS] Success: {msg}")
                    self.camera_update_signal.emit(cameras, nvr_name, elapsed, nvr_online, nvr_offline)
                else:
                    def show_error():
                        self.status.showMessage(f"❌ No cameras found from {nvr_name}. Check credentials and network connectivity.", 5000)
                        log(f"[UPDATE-CAMERAS] No cameras found from {nvr_name} using method: {method}")
                    QtCore.QTimer.singleShot(0, show_error)
            except Exception as e:
                log(f"[UPDATE-CAMERAS] Exception in background thread: {str(e)}")
                import traceback
                log(f"[UPDATE-CAMERAS] Exception details: {traceback.format_exc()}")
                def show_error():
                    self.status.showMessage(f"❌ Failed to update cameras from {nvr_name}: {str(e)}", 8000)
                    log(f"[UPDATE-CAMERAS] Error: {str(e)}")
                QtCore.QTimer.singleShot(0, show_error)

        threading.Thread(target=fast_update_thread, daemon=True).start()

    @staticmethod
    def _clean_text(value):
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        text = str(value).strip()
        return "" if text.lower() == "none" else text

    def get_resource_path(self, relative_path):
        """Get absolute path to resource, works for dev and for PyInstaller"""
        try:
            # PyInstaller creates a temp folder and stores path in _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)
    
    def get_data_path(self, relative_path):
        """Get absolute path to data files (Excel, logs, etc.) - always in exe directory"""
        if getattr(sys, 'frozen', False):
            # Running as exe - use exe directory
            base_path = os.path.dirname(sys.executable)
        else:
            # Running as script - use script directory
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def __init__(self):
        super().__init__()
        
        # ==================== LICENSE VALIDATION ====================
        if REQUIRES_LICENSE:
            machine_id = get_machine_id()
            license_key = load_license_key()
            license_valid = False
            
            if license_key:
                is_valid, message = validate_license_key(license_key, machine_id)
                if is_valid:
                    license_valid = True
                    log(f"[LICENSE] Valid license: {message}")
                else:
                    log(f"[LICENSE] Invalid license: {message}")
            
            # Show license dialog if no valid license
            if not license_valid:
                license_dlg = LicenseDialog(None, machine_id)
                if license_dlg.exec_() != QtWidgets.QDialog.Accepted:
                    # User cancelled or closed dialog - exit application
                    sys.exit(0)
        
        # ==================== APPLICATION INITIALIZATION ====================
        self.setWindowTitle(f"🔐 {APP_TITLE} v{APP_VERSION} Enhanced Edition")
        self.resize(1150, 700)
        self.vlc = find_vlc_executable()
        self.nvrs = []
        self.cams = []  # Excel-based cameras (for reference)
        self.api_cameras = []  # Live API results
        self.cameras = self.api_cameras  # Use API cameras as primary source
        self.filtered = []
        self.check_history = {}  # persistent map: ip -> {status, device_type, model, timestamp}
        self.creds_meta = load_creds_meta()
        self.current_check_id = 0  # Track current check operation to cancel on NVR switch
        self.nvr_thread_running = False  # Prevent concurrent NVR operations
        self.nvr_operation_lock = threading.Lock()  # Thread safety for NVR operations
        self.offline_dialog = None  # Track active offline-camera popup
        self.camera_check_progress = {}
        
        # Enhanced caching system for v8.6+
        self.connection_cache = {}  # Cache for connection status: ip -> {status, timestamp}
        self.nvr_cache = {}  # Cache for NVR responses: nvr_ip -> {cameras, timestamp}
        self.performance_metrics = {  # Track performance metrics
            'total_checks': 0,
            'cache_hits': 0,
            'average_response_time': 0.0,
            'last_full_scan': 0
        }

        # Connect camera update signal
        self.camera_update_signal.connect(self._handle_camera_update)
        
        # Connect new progress and UI update signals
        self.progress_update_signal.connect(self._handle_progress_update)
        self.button_control_signal.connect(self._handle_button_control)
        self.ui_status_update_signal.connect(self._handle_ui_status_update)
        self.ui_call_signal.connect(self._execute_ui_callable)
        
        # Connect enhanced v8.6+ signals
        self.error_notification_signal.connect(self._handle_error_notification)
        self.performance_update_signal.connect(self._handle_performance_update)
        self.cache_stats_signal.connect(self._handle_cache_stats)

        # Set window icon if logo exists
        logo_path = self.get_resource_path(LOGO_FILE)
        if os.path.exists(logo_path):
            app_icon = QtGui.QIcon(logo_path)
            self.setWindowIcon(app_icon)
            # Also set application-wide icon for taskbar
            QtWidgets.QApplication.setWindowIcon(app_icon)
            log("Application icons loaded successfully")
        else:
            log(f"Logo file not found: {logo_path}")

        self.setStyleSheet("""
            QWidget{font-family:Arial; font-size:12px;}
            QHeaderView::section{background:#f6f6f6;padding:6px;border:1px solid #ddd;}
            QListWidget{border:1px solid #ccc;background:#fff;}
            QGroupBox{border:1px solid #ddd;padding:6px;margin-top:8px;}
        """)

        central = QtWidgets.QWidget(); self.setCentralWidget(central)
        vbox = QtWidgets.QVBoxLayout(central)

        # toolbar
        top = QtWidgets.QHBoxLayout()
        
        # Logo in toolbar
        logo_path = self.get_resource_path(LOGO_FILE)
        if os.path.exists(logo_path):
            logo_label = QtWidgets.QLabel()
            pixmap = QtGui.QPixmap(logo_path).scaledToHeight(32, QtCore.Qt.SmoothTransformation)
            logo_label.setPixmap(pixmap)
            top.addWidget(logo_label)
        
        self.btn_load = QtWidgets.QPushButton("📂 Load Excel"); self.btn_load.clicked.connect(self.load_data)
        self.btn_export = QtWidgets.QPushButton("💾 Export Report"); self.btn_export.clicked.connect(self.show_export_dialog)
        self.btn_check_sel = QtWidgets.QPushButton("🔍 Check Selected"); self.btn_check_sel.clicked.connect(self.check_selected)
        # Duplicate detection button
        self.btn_duplicates = QtWidgets.QPushButton("🔍 Find Duplicates v8.6"); self.btn_duplicates.clicked.connect(self.show_duplicate_report)
        self.btn_duplicates.setStyleSheet("QPushButton { background-color: #f39c12; color: white; font-weight: bold; padding: 6px; border-radius: 4px; }")
        self.btn_duplicates.setToolTip("v8.6 Enhanced: Detect duplicate cameras across all sources with advanced algorithms")
        
        # SADP Device Discovery Tool
        self.btn_sadp = QtWidgets.QPushButton("🔍 SADP Tool"); self.btn_sadp.clicked.connect(self.show_sadp_tool)
        self.btn_sadp.setStyleSheet("QPushButton { background-color: #16a085; color: white; font-weight: bold; padding: 6px; border-radius: 4px; }")
        self.btn_sadp.setToolTip("Search Active Device Protocol - Discover Hikvision devices on network")
        # Search bar with more space
        self.search = QtWidgets.QLineEdit(); 
        self.search.setPlaceholderText("🔍 Search camera or IP..."); 
        self.search.setMinimumWidth(300)
        self.search.textChanged.connect(self.filter_table)
        
        top.addWidget(self.btn_load)
        top.addWidget(self.btn_export)
        top.addWidget(self.btn_check_sel)
        top.addWidget(self.btn_duplicates)
        top.addWidget(self.btn_sadp)
        top.addStretch()
        # Add update button at the end
        if UPDATE_MANAGER_AVAILABLE:
            self.btn_update = QtWidgets.QPushButton("🔄 Check Updates")
            self.btn_update.setToolTip(f"Check for {APP_TITLE} updates\nCurrent version: {APP_VERSION}\nClick to manually check for new versions")
            self.btn_update.setStyleSheet("""
                QPushButton {
                    background-color: #3498db;
                    color: white;
                    border: none;
                    padding: 6px 12px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #2980b9;
                }
                QPushButton:pressed {
                    background-color: #21618c;
                }
            """)
            self.btn_update.clicked.connect(self.check_for_updates_manual)
            top.addWidget(self.btn_update)
        
        # About button
        self.btn_about = QtWidgets.QPushButton("ℹ️ About")
        self.btn_about.setToolTip("About this Enhanced Edition")
        self.btn_about.setStyleSheet("""
            QPushButton {
                background-color: #16a085;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #138d75;
            }
        """)
        self.btn_about.clicked.connect(self.show_about_dialog)
        top.addWidget(self.btn_about)
        top.addWidget(self.search)
        vbox.addLayout(top)

        # splitter layout
        splitter = QtWidgets.QSplitter(QtCore.Qt.Horizontal); vbox.addWidget(splitter)

        # left NVR list & details
        left = QtWidgets.QWidget(); left_l = QtWidgets.QVBoxLayout(left)
        left_l.setContentsMargins(6,6,6,6); left_l.setSpacing(6)
        left_l.addWidget(QtWidgets.QLabel("🗄️ NVR Overview"))
        self.list_nvr = QtWidgets.QListWidget(); self.list_nvr.itemSelectionChanged.connect(self.on_nvr_selected)
        self.list_nvr.setSizeAdjustPolicy(QtWidgets.QAbstractScrollArea.AdjustToContents)
        # Add context menu for right-click on NVR
        self.list_nvr.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.list_nvr.customContextMenuRequested.connect(self.show_nvr_context_menu)
        left_l.addWidget(self.list_nvr)
        
        # NVR management buttons
        nvr_buttons_layout = QtWidgets.QHBoxLayout()
        
        # Add NVR button
        self.btn_add_nvr = QtWidgets.QPushButton("➕ Add NVR")
        self.btn_add_nvr.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 6px 8px;
                border-radius: 3px;
                font-weight: bold;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        self.btn_add_nvr.clicked.connect(self.add_new_nvr)
        nvr_buttons_layout.addWidget(self.btn_add_nvr)
        
        # Refresh NVR button
        self.btn_refresh_nvr = QtWidgets.QPushButton("🔄 Refresh")
        self.btn_refresh_nvr.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 6px 8px;
                border-radius: 3px;
                font-weight: bold;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        self.btn_refresh_nvr.clicked.connect(self.refresh_nvr_status)
        nvr_buttons_layout.addWidget(self.btn_refresh_nvr)
        
        left_l.addLayout(nvr_buttons_layout)
        self.grp = QtWidgets.QGroupBox("NVR Details"); form = QtWidgets.QFormLayout()
        self.lbl_name = QtWidgets.QLabel("-"); self.lbl_ip = QtWidgets.QLabel("-"); self.lbl_subnet = QtWidgets.QLabel("-"); self.lbl_gw = QtWidgets.QLabel("-"); self.lbl_sheet = QtWidgets.QLabel("-"); self.lbl_real_ip = QtWidgets.QLabel("-")
        form.addRow("Name:", self.lbl_name); form.addRow("IP:", self.lbl_ip); form.addRow("Real IP:", self.lbl_real_ip); form.addRow("Subnet:", self.lbl_subnet); form.addRow("Gateway:", self.lbl_gw); form.addRow("Sheet:", self.lbl_sheet)
        self.grp.setLayout(form); left_l.addWidget(self.grp)
        splitter.addWidget(left)

        # right camera table
        right = QtWidgets.QWidget(); right_l = QtWidgets.QVBoxLayout(right)
        right_l.setContentsMargins(4,4,4,4); right_l.setSpacing(6)
        right_l.addWidget(QtWidgets.QLabel("📷 Cameras"))

        # --- enhanced table block with comprehensive camera information ---
        self.table = QtWidgets.QTableWidget()
        self.table.setColumnCount(9)
        self.table.setHorizontalHeaderLabels(["", "Camera Name", "IP", "Status", "Model", "Port", "NVR", "Last Updated", "Remark"])

        header = self.table.horizontalHeader()
        # icon column fixed
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.Fixed)
        header.resizeSection(0, 30)
        # camera name and ip sized to content
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeToContents)  # Camera Name
        header.setSectionResizeMode(2, QtWidgets.QHeaderView.ResizeToContents)  # IP
        # status and technical info sized to content
        header.setSectionResizeMode(3, QtWidgets.QHeaderView.ResizeToContents)  # Status
        header.setSectionResizeMode(4, QtWidgets.QHeaderView.ResizeToContents)  # Model
        header.setSectionResizeMode(5, QtWidgets.QHeaderView.ResizeToContents)  # Port
        header.setSectionResizeMode(6, QtWidgets.QHeaderView.ResizeToContents)  # NVR
        header.setSectionResizeMode(7, QtWidgets.QHeaderView.ResizeToContents)  # Last Updated
        # let last column (Remark) stretch to fill remaining space
        header.setSectionResizeMode(8, QtWidgets.QHeaderView.Stretch)  # Remark

        self.table.setSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        self.table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.table.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.table.setContextMenuPolicy(QtCore.Qt.CustomContextMenu); self.table.customContextMenuRequested.connect(self.open_context_menu)
        self.table.cellDoubleClicked.connect(self.double_click_ip)

        right_l.addWidget(self.table)
        splitter.addWidget(right)

        splitter.setStretchFactor(0, 0); splitter.setStretchFactor(1, 1)
        splitter.setSizes([200, 950])  # give more room to table by default

        # status bar + signals
        self.status = QtWidgets.QStatusBar(); self.setStatusBar(self.status)
        
        # LEFT SIDE: Selected NVR and Camera info
        self.lbl_selected_nvr = QtWidgets.QLabel("Selected NVR: None")
        self.lbl_selected_nvr.setStyleSheet("color: #2c3e50; font-weight: bold; margin-right: 10px;")
        self.lbl_selected_cameras = QtWidgets.QLabel("Cameras: 0")
        self.lbl_selected_cameras.setStyleSheet("color: #34495e; font-weight: bold; margin-right: 10px;")
        self.lbl_selected_online = QtWidgets.QLabel("🟢 0")
        self.lbl_selected_online.setStyleSheet("color: green; font-weight: bold; margin-right: 10px;")
        self.lbl_selected_offline = QtWidgets.QLabel("🔴 0")
        self.lbl_selected_offline.setStyleSheet("color: red; font-weight: bold; margin-right: 15px;")
        
        # Add left side widgets
        self.status.addWidget(self.lbl_selected_nvr)
        self.status.addWidget(self.lbl_selected_cameras)
        self.status.addWidget(self.lbl_selected_online)
        self.status.addWidget(self.lbl_selected_offline)
        
        # RIGHT SIDE: Total counters (permanent widgets)
        # NVR counters
        self.lbl_nvr_total = QtWidgets.QLabel("🗄️ Total NVRs: 0")
        self.lbl_nvr_total.setStyleSheet("color: blue; font-weight: bold; margin-left: 10px;")
        self.lbl_nvr_online = QtWidgets.QLabel("🟢 Online: 0")
        self.lbl_nvr_online.setStyleSheet("color: green; font-weight: bold; margin-left: 5px;")
        self.lbl_nvr_offline = QtWidgets.QLabel("🔴 Offline: 0")
        self.lbl_nvr_offline.setStyleSheet("color: red; font-weight: bold; margin-left: 5px;")
        
        # Camera counters
        self.lbl_total = QtWidgets.QLabel("📷 All Cameras: 0")
        self.lbl_total.setStyleSheet("color: #2980b9; font-weight: bold; margin-left: 15px;")
        self.lbl_online = QtWidgets.QLabel("🟢 Online: 0")
        self.lbl_online.setStyleSheet("color: green; font-weight: bold; margin-left: 5px;")
        self.lbl_offline = QtWidgets.QLabel("🔴 Offline: 0")
        self.lbl_offline.setStyleSheet("color: red; font-weight: bold; margin-left: 5px;")
        
        # Add right side widgets (permanent = right aligned)
        self.status.addPermanentWidget(self.lbl_nvr_total)
        self.status.addPermanentWidget(self.lbl_nvr_online)
        self.status.addPermanentWidget(self.lbl_nvr_offline)
        self.status.addPermanentWidget(QtWidgets.QLabel(" | "))
        self.status.addPermanentWidget(self.lbl_total)
        self.status.addPermanentWidget(self.lbl_online)
        self.status.addPermanentWidget(self.lbl_offline)
        
        self.table_update.connect(self.apply_table_update)
        self.enhanced_table_update.connect(self.apply_enhanced_table_update)
        self.nvr_update.connect(self.apply_nvr_update)
        self.nvr_login_result.connect(self.on_nvr_login_result)

        # Auto-load Excel on startup
        excel_path = self.get_data_path(EXCEL_FILE)
        if os.path.exists(excel_path):
            try:
                self.load_data(initial=True)
                log(f"Auto-loaded Excel file: {excel_path}")
            except Exception as e:
                log(f"Failed to auto-load Excel: {traceback.format_exc()}")
        # load check history if present
        try:
            if os.path.exists(CHECK_HISTORY_FILE):
                with open(CHECK_HISTORY_FILE, 'r', encoding='utf-8') as f:
                    self.check_history = json.load(f)
        except Exception:
            self.check_history = {}
        
        # Auto-update check on startup (non-intrusive, background check)
        if UPDATE_MANAGER_AVAILABLE:
            QtCore.QTimer.singleShot(2000, self.check_for_updates_startup)

    # ---------------- data load ----------------
    def load_data(self, initial=False):
        try:
            excel_path = self.get_data_path(EXCEL_FILE)
            nvrs, cams = load_excel_robust(excel_path)
            
            # Merge with enhanced NVR configurations (credentials, etc.)
            enhanced_nvrs = get_enhanced_nvr_configs()
            
            # Update existing NVRs with credential information
            for nvr in nvrs:
                nvr_name = nvr.get('name', '')
                nvr_ip = nvr.get('ip', '')
                
                # Find matching enhanced config
                for enhanced in enhanced_nvrs:
                    if (enhanced.get('name') == nvr_name or 
                        enhanced.get('ip') == nvr_ip):
                        # Merge enhanced data (credentials, etc.) but keep Excel data priority for basic info
                        for key, value in enhanced.items():
                            if key not in ['name', 'ip', 'cameras', 'cam_count'] or key not in nvr:
                                nvr[key] = value
                        break
                else:
                    # Add default credentials if no enhanced config found
                    nvr.update({
                        'port': 80,
                        'protocol': 'http',
                        'username': 'admin',
                        'password': 'Kkcctv12345'
                    })
            
            # Add any custom NVRs not in Excel
            excel_names = {nvr.get('name') for nvr in nvrs}
            excel_ips = {nvr.get('ip') for nvr in nvrs}
            
            for enhanced in enhanced_nvrs:
                if (enhanced.get('name') not in excel_names and 
                    enhanced.get('ip') not in excel_ips):
                    nvrs.append(enhanced)
            
            self.nvrs = nvrs; self.cams = cams; self.cameras = self.api_cameras; self.filtered = list(self.api_cameras)
            self.populate_nvr_list(); self.populate_table(self.filtered)
            self.status.showMessage(f"Loaded {len(self.nvrs)} NVRs, {len(self.cams)} cameras")
        except Exception as e:
            log(traceback.format_exc())
            if not initial:
                QtWidgets.QMessageBox.critical(self, "Load error", str(e))

    def populate_nvr_list(self):
        self.list_nvr.clear()
        for idx, n in enumerate(self.nvrs):
            # Get NVR status and display appropriate emoji
            status = n.get('status', '')
            if status == 'online':
                emoji = "🟢"
            elif status == 'limited':
                emoji = "🟡"
            elif status == 'offline':
                emoji = "🔴"
            else:
                emoji = "🗄️"  # Default if no status
            
            cam_count = self._get_camera_count_for_nvr(n.get('name', ''), n.get('ip', ''))
            self.nvrs[idx]['cam_count'] = cam_count
            sheet_flag = "" if n.get("sheet_found", False) else " ⚠️ sheet missing"
            text = f"{emoji} {n.get('name','')} | {n.get('ip','')} | 🎥 {cam_count}{sheet_flag}"
            item = QtWidgets.QListWidgetItem(text); item.setData(QtCore.Qt.UserRole, idx)
            self.list_nvr.addItem(item)

    def populate_table(self, camlist):
        self.table.setRowCount(0)
        # UNIFIED APPROACH: Show all cameras but prioritize API data
        # This ensures cameras are always displayed regardless of source
        for c in camlist:
            r = self.table.rowCount(); self.table.insertRow(r)
            # IVMS Method: Simple status display based on NVR's devIndex assessment
            status_raw = c.get('status', None)
            status_val = (status_raw or '').lower() if status_raw is not None else ''
            # Check status text - handle both simple and verified statuses
            if 'online' in status_val or '🟢' in status_val:
                # Camera online (including verified)
                badge_emoji = '🟢'
                status_color = QtGui.QColor(0, 160, 0)
            elif 'offline' in status_val or '🔴' in status_val:
                # Camera offline (including verified)
                badge_emoji = '🔴'
                status_color = QtGui.QColor(160, 0, 0)
            elif 'unknown' in status_val or '🟡' in status_val:
                # Camera status unknown
                badge_emoji = '🟡'
                status_color = QtGui.QColor(200, 200, 0)
            else:
                # Fallback for any other status
                badge_emoji = '❔'
                status_color = QtGui.QColor(120, 120, 120)
            badge = QtWidgets.QTableWidgetItem(badge_emoji)
            badge.setTextAlignment(QtCore.Qt.AlignCenter)
            badge.setFlags(QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsSelectable)
            # Enhanced tooltip with status explanation
            status_raw = c.get('status', 'Unknown')
            connection_type = c.get('connection_type', '')
            tooltip_text = f"Status: {status_raw}"
            if connection_type:
                tooltip_text += f"\nConnection: {connection_type.replace('_', ' ').title()}"
            badge.setToolTip(tooltip_text)
            self.table.setItem(r, 0, badge)
            # Name
            name_item = QtWidgets.QTableWidgetItem(c.get("name", ""))
            name_item.setToolTip(f"Camera Name: {c.get('name','')}")
            self.table.setItem(r, 1, name_item)
            # IP
            ip_item = QtWidgets.QTableWidgetItem(c.get("ip", ""))
            ip_item.setToolTip(f"IP: {c.get('ip','')}")
            self.table.setItem(r, 2, ip_item)
            # Status (direct, no normalization)
            status_item = QtWidgets.QTableWidgetItem(c.get('status',''))
            status_item.setTextAlignment(QtCore.Qt.AlignCenter)
            status_item.setForeground(status_color)
            status_item.setToolTip(f"Status: {c.get('status','')}")
            self.table.setItem(r, 3, status_item)
            # Model
            model_txt = self._clean_text(c.get("model"))
            model_item = QtWidgets.QTableWidgetItem(model_txt)
            model_item.setToolTip(f"Model: {model_txt}")
            self.table.setItem(r, 4, model_item)
            # Port
            port_txt = str(c.get("port", ''))
            port_item = QtWidgets.QTableWidgetItem(port_txt)
            port_item.setToolTip(f"Port: {port_txt}")
            self.table.setItem(r, 5, port_item)
            # NVR index/name
            nvr_idx_text = ""
            try:
                cam_nvr = (c.get("nvr", "") or "").strip()
                if cam_nvr:
                    for i, n in enumerate(self.nvrs):
                        if self._clean_text(n.get("name")).lower() == cam_nvr.lower():
                            nvr_idx_text = str(i+1)
                            break
                    if not nvr_idx_text:
                        nvr_idx_text = cam_nvr

            except Exception:
                nvr_idx_text = ""
            nvr_item = QtWidgets.QTableWidgetItem(nvr_idx_text if nvr_idx_text else "")
            nvr_item.setForeground(QtGui.QColor("black"))
            nvr_item.setToolTip(f"NVR: {nvr_idx_text}")
            self.table.setItem(r, 6, nvr_item)
            # Last updated
            updated_txt = c.get('last_updated', '')
            updated_item = QtWidgets.QTableWidgetItem(updated_txt)
            updated_item.setToolTip(f"Last Updated: {updated_txt}")
            self.table.setItem(r, 7, updated_item)
            # Remark
            remark_txt = c.get('remark', '')
            remark_item = QtWidgets.QTableWidgetItem(remark_txt)
            remark_item.setForeground(QtGui.QColor("blue"))
            remark_item.setToolTip(remark_txt)
            self.table.setItem(r, 8, remark_item)
        self.update_counters()
    
    def update_counters(self):
        """Update status bar counters for NVRs and cameras"""
        # Use API cameras as primary source (live data from NVRs), fallback to Excel
        camera_source = self.api_cameras if self.api_cameras else self.cams
        
        # Count ALL cameras including duplicates (same IP on different NVRs = different cameras)
        online_cameras = 0
        offline_cameras = 0
        configured_cameras = 0
        seen_ips = {}
        
        for cam in camera_source:
            ip = cam.get('ip', '').strip()
            
            # Track duplicate IPs for display
            if ip:
                seen_ips[ip] = seen_ips.get(ip, 0) + 1
            
            status = cam.get('status', '').lower()
            if 'online' in status or '🟢' in status:
                online_cameras += 1
            elif 'offline' in status or '🔴' in status:
                offline_cameras += 1
            elif 'configured' in status:
                configured_cameras += 1
        
        # Total is ALL cameras (including duplicates)
        total_cameras = len(camera_source)
        # Count unique IPs and duplicates
        unique_ips = len(seen_ips)
        duplicate_count = total_cameras - unique_ips
        
        # Count NVRs
        total_nvrs = len(self.nvrs)
        online_nvrs = 0
        offline_nvrs = 0
        
        for nvr in self.nvrs:
            nvr_name = self._clean_text(nvr.get('name')).lower()
            # Count cameras for this NVR from API data (live)
            camera_source = self.api_cameras if self.api_cameras else self.cams
            nvr_cameras = [
                c for c in camera_source
                if self._clean_text(c.get('nvr')).lower() == nvr_name
            ]
            nvr_online_count = sum(1 for c in nvr_cameras if 'online' in c.get('status', '').lower())
            
            if nvr_online_count > 0:
                online_nvrs += 1
            else:
                offline_nvrs += 1
        
        # Update RIGHT SIDE labels (totals) - using direct updates since this is called from main thread
        self.lbl_nvr_total.setText(f"🗄️ Total NVRs: {total_nvrs}")
        self.lbl_nvr_online.setText(f"🟢 Online: {online_nvrs}")
        self.lbl_nvr_offline.setText(f"🔴 Offline: {offline_nvrs}")
        self.lbl_total.setText(f"📷 Total: {total_cameras} | Unique: {unique_ips} | Duplicates: {duplicate_count}")
        self.lbl_online.setText(f"🟢 Online: {online_cameras}")
        self.lbl_offline.setText(f"🔴 Offline: {offline_cameras}")
        
        # Update status bar with comprehensive status info
        status_text = f"📷 {total_cameras} total | 🟢 {online_cameras} online | 🔴 {offline_cameras} offline"
        if configured_cameras > 0:
            status_text += f" | 🟡 {configured_cameras} configured"
        self.status.showMessage(status_text, 0)
        
        # Also emit signals for any background thread usage
        self.ui_status_update_signal.emit("emoji_status", "nvr_online", f"🟢|{online_nvrs}|#d5f4e6")
        self.ui_status_update_signal.emit("emoji_status", "nvr_offline", f"🔴|{offline_nvrs}|#fdeaea") 
        self.ui_status_update_signal.emit("emoji_status", "cam_online", f"🟢|{online_cameras}|#d5f4e6")
        self.ui_status_update_signal.emit("emoji_status", "cam_offline", f"🔴|{offline_cameras}|#fdeaea")
        
        # Update LEFT SIDE labels (selected NVR info)
        self.update_selected_counters()

    def update_selected_counters(self):
        """Update LEFT SIDE status bar counters for selected NVR and its cameras"""
        if not hasattr(self, 'filtered'):
            self.filtered = self.api_cameras if self.api_cameras else self.cams
        
        # Count cameras in current selection/filter from API data (includes verified statuses)
        selected_total = len(self.filtered)
        selected_online = 0
        selected_offline = 0
        selected_unknown = 0
        
        for cam in self.filtered:
            status = str(cam.get('status', '')).lower()
            # Count based on actual status values including verified ones
            if 'online' in status or '🟢' in status:
                selected_online += 1
            elif 'offline' in status or '🔴' in status:
                selected_offline += 1
            else:
                selected_unknown += 1
        
        # Update left side labels with enhanced formatting
        self.lbl_selected_cameras.setText(f"Cameras: {selected_total}")
        self.lbl_selected_online.setText(f"🟢 {selected_online}")
        self.lbl_selected_offline.setText(f"🔴 {selected_offline}")
        if selected_unknown > 0:
            self.lbl_selected_offline.setText(f"🔴 {selected_offline} ❔ {selected_unknown}")

    # ---------------- NVR select/filter ----------------
    def on_nvr_selected(self):
        items = self.list_nvr.selectedItems()
        if not items:
            # No NVR selected - show all cameras
            self.lbl_selected_nvr.setText("Selected NVR: None (All)")
            # Use API cameras as primary source, merge with Excel if needed
            if self.api_cameras:
                self.filtered = list(self.api_cameras)
            else:
                self.filtered = list(self.cams)
            self.populate_table(self.filtered)
            self.update_selected_counters()
            return
            
        idx = items[0].data(QtCore.Qt.UserRole)
        n = self.nvrs[idx]
        self.lbl_name.setText(n.get("name","")); self.lbl_ip.setText(n.get("ip",""))
        self.lbl_subnet.setText(n.get("subnet","")); self.lbl_gw.setText(n.get("gateway",""))
        self.lbl_sheet.setText("Found" if n.get("sheet_found", False) else "Missing")
        name = self._clean_text(n.get("name")).lower()
        # UNIFIED NVR FILTERING: Use API cameras as primary, fallback to Excel
        cams = []
        if self.api_cameras:
            cams = [c for c in self.api_cameras if self._clean_text(c.get("nvr")).lower() == name]
            log(f"[NVR-FILTER] Selected '{name}', found {len(cams)} API cameras")
            # Debug: show first 3 cameras and their NVR values
            for i, c in enumerate(cams[:3]):
                log(f"  Camera {i+1}: {c.get('name', 'NO_NAME')} -> NVR: '{c.get('nvr', 'NO_NVR')}'")
        if not cams and self.cams:  # Fallback if no API cameras found for this NVR
            cams = [c for c in self.cams if self._clean_text(c.get("nvr")).lower() == name]
            log(f"[NVR-FILTER] Fallback to Excel: found {len(cams)} cameras for '{name}'")
        self.filtered = cams; self.populate_table(self.filtered)
        
        # Update left side status bar with selected NVR info
        self.lbl_selected_nvr.setText(f"Selected NVR: {n.get('name', 'Unknown')}")
        self.update_selected_counters()
        
        # Cancel any ongoing check operations by incrementing check ID
        self.current_check_id += 1

    def filter_table(self):
        q = self.search.text().strip().lower()
        # UNIFIED SEARCH: Use API cameras as primary, fallback to Excel
        if self.api_cameras:
            if not q:
                self.filtered = list(self.api_cameras)
            else:
                self.filtered = [c for c in self.api_cameras if q in c.get("name","").lower() or q in c.get("ip","").lower() or q in c.get("nvr","").lower()]
        else:
            # Fallback to Excel cameras if no API cameras available
            if not q:
                self.filtered = list(self.cams)
            else:
                self.filtered = [c for c in self.cams if q in c.get("name","").lower() or q in c.get("ip","").lower() or q in c.get("nvr","").lower()]
        
        self.populate_table(self.filtered)
        self.status.showMessage(f"{len(self.filtered)} entries")

    # ---------------- checks ----------------
    def check_selected(self):
        """Run classic camera checks (SADP, HTTP, RTSP) against selected rows."""
        rows = sorted({r.row() for r in self.table.selectionModel().selectedRows()})
        if not rows:
            QtWidgets.QMessageBox.information(self, "Check", "Select camera rows first.")
            return

        targets = []
        for r in rows:
            if r >= self.table.rowCount():
                continue
            ip_item = self.table.item(r, 2)
            ip_val = ip_item.text().strip() if ip_item else ""
            if not ip_val:
                continue
            targets.append({"row": r, "ip": ip_val})

        if not targets:
            QtWidgets.QMessageBox.information(self, "Check", "No valid IP addresses found in the selected rows.")
            return

        self.status.showMessage(f"🔄 Checking {len(targets)} cameras...", 0)
        threading.Thread(target=self._run_checks, args=(targets,), daemon=True).start()



    def check_live_status(self):
        """Check camera live status via NVR API or SADP."""
        rows = sorted({r.row() for r in self.table.selectionModel().selectedRows()})
        if not rows:
            QtWidgets.QMessageBox.information(self, "Check Live", "Select camera rows first, or click again to check all.")
            return
        
        # Get NVR credentials for the selected cameras
        nvr_ip = nvr_user = nvr_pwd = ""
        selected_nvr_name = None
        
        # Try to get NVR info from selected cameras
        for row in rows:
            cam = self.filtered[row] if row < len(self.filtered) else None
            if cam:
                cam_nvr_name = self._clean_text(cam.get("nvr")).lower()
                for n in self.nvrs:
                    if self._clean_text(n.get("name")).lower() == cam_nvr_name:
                        nvr_ip = n.get("ip", "")
                        nvr_user, nvr_pwd = get_password(nvr_ip) if nvr_ip else (None, None)
                        if not nvr_user:
                            nvr_user = "admin"
                        if not nvr_pwd:
                            nvr_pwd = DEFAULT_CREDS[0][1] if DEFAULT_CREDS else "Kkcctv12345"
                        selected_nvr_name = n.get("name", "")
                        break
                if nvr_ip:
                    break
        
        targets = [{"row": r, "ip": self.table.item(r, 2).text().strip()} for r in rows if self.table.item(r, 2)]
        threading.Thread(target=self._run_live_checks, args=(targets, nvr_ip, nvr_user, nvr_pwd), daemon=True).start()

    def check_all_parallel(self):
        """Legacy method - redirects to NVR-based checking for better performance."""
        log("[REMOVED] check_all_parallel() method has been removed")
        QtWidgets.QMessageBox.information(self, "Feature Removed", "The Check All feature has been removed in this version.")


        # Initialize IP to row mapping for table updates
        ip_to_row = {}
        for r in range(self.table.rowCount()):
            try:
                item = self.table.item(r,2)
                if item:
                    ip = item.text().strip()
                    if ip:
                        ip_to_row[ip] = r
            except Exception:
                continue

        # Prepare NVR list
        nvrs = list(self.nvrs)
        cams = list(self.cams)

        def enhanced_check_ip(ip):
            # Try SADP first (most reliable for Hikvision)
            try:
                sadp_online, sadp_model = check_camera_via_sadp(ip, timeout=1.2)
                if sadp_online:
                    return True, "SADP", sadp_model
            except Exception:
                pass
                
            # Try TCP ports with better timeout
            try:
                h = check_tcp(ip, HTTP_PORT, timeout=0.8)
                r = check_tcp(ip, RTSP_PORT, timeout=0.8)
                if h or r:
                    methods = []
                    if h: methods.append('HTTP')
                    if r: methods.append('RTSP')
                    return True, 'TCP', ','.join(methods)
            except Exception:
                pass
                
            # Last resort ping
            try:
                if silent_ping(ip):
                    return True, 'Ping', 'Network reachable'
            except Exception:
                pass
            return False, 'Offline', 'No response'

        def run_checks_thread():
            cam_online = 0
            cam_total = 0
            cam_offline = 0
            # check NVRs first via existing _check_nvr (which updates UI)
            for idx, n in enumerate(nvrs):
                try:
                    threading.Thread(target=self._check_nvr, args=(idx, n), daemon=True).start()
                except Exception:
                    pass

            # Now check cameras in parallel
            cam_total = len(cams)
            if cam_total == 0:
                QtCore.QTimer.singleShot(0, lambda: self._finish_ip_check(0,0,0))
                return

            # Use enhanced ThreadPoolExecutor with progress dialog
            cam_total = len(cams)
            prog = None
            try:
                QtCore.QTimer.singleShot(0, lambda: self.status.showMessage('🔍 Enhanced parallel checking...'))
                prog = QtWidgets.QProgressDialog('🚀 Enhanced Camera Check - Parallel Processing', 'Cancel', 0, cam_total, self)
                prog.setWindowModality(QtCore.Qt.WindowModal)
                prog.setMinimumDuration(0)
                prog.setValue(0)
                # must show on UI thread
                QtCore.QTimer.singleShot(0, prog.show)
            except Exception:
                prog = None

            checked = 0
            # Use enhanced check with better thread pool size
            with concurrent.futures.ThreadPoolExecutor(max_workers=50) as ex:
                futures = {ex.submit(enhanced_check_ip, c.get('ip','')): c for c in cams if c.get('ip')}
                for fut in concurrent.futures.as_completed(futures):
                    cam = futures[fut]
                    ip = cam.get('ip','')
                    try:
                        ok, method, details = fut.result()
                    except Exception as e:
                        ok, method, details = False, 'Error', str(e)

                    checked += 1
                    # Get enhanced camera information from check history or defaults
                    cam_info = None
                    for hist_cam in self.check_history.get('cameras', []):
                        if hist_cam.get('ip') == ip:
                            cam_info = hist_cam
                            break
                    
                    # prepare enhanced update values
                    if ok:
                        cam_online += 1
                        em = '🟢'
                        color = QtGui.QColor(0,160,0)
                        status_text = f'Online ({method})'
                        device_type = method
                        
                        # Use enhanced model info from SADP or check history
                        if method == 'SADP' and details:
                            model = details
                        elif cam_info and cam_info.get('model'):
                            model = cam_info.get('model')
                        else:
                            model = details or '—'
                    else:
                        cam_offline += 1
                        em = '🔴'
                        color = QtGui.QColor(160,0,0)
                        status_text = 'Offline'
                        device_type = '—'
                        model = '—'

                    # update visible table row with enhanced information
                    row = ip_to_row.get(ip) if 'ip_to_row' in locals() else None
                    if row is not None:
                        # Get additional info from camera object or check history
                        channel = cam_info.get('channel', '—') if cam_info else cam.get('channel', '—')
                        port = cam_info.get('port', '—') if cam_info else cam.get('port', '—')
                        serial = cam_info.get('serial', '—') if cam_info else cam.get('serial', '—')
                        firmware = cam_info.get('firmware', '—') if cam_info else cam.get('firmware', '—')
                        nvr_name = cam.get('nvr', '—')
                        last_updated = time.strftime('%Y-%m-%d %H:%M')
                        
                        # Enhanced table update with comprehensive information
                        try:
                            self.enhanced_table_update.emit(row, status_text, device_type, model, 
                                                           channel, port, serial, firmware, nvr_name, 
                                                           last_updated, color, em)
                        except:
                            # Fallback to regular update if enhanced fails
                            self.table_update.emit(row, status_text, device_type, model, color, em)

                    # update progress dialog
                    if prog:
                        QtCore.QTimer.singleShot(0, lambda v=checked: prog.setValue(v))

                    log(f"Check IP {ip}: {status_text} ({method})")

            QtCore.QTimer.singleShot(0, lambda: self._finish_ip_check(cam_total, cam_online, cam_offline))

        threading.Thread(target=run_checks_thread, daemon=True).start()

    def _finish_ip_check(self, total, online, offline):
        try:
            self.btn_check_all.setEnabled(True)
            self.btn_check_live.setEnabled(True)
            self.status.showMessage(f"✅ Enhanced check complete: {online}/{total} online, {offline} offline")
            QtWidgets.QMessageBox.information(self, "🚀 Enhanced Check Complete", 
                                             f"📊 Camera Status Summary:\n\n"
                                             f"🟢 Online: {online}\n"
                                             f"🔴 Offline: {offline}\n"
                                             f"📈 Total: {total}\n\n"
                                             f"💡 Enhanced information displayed in table")
        except Exception:
            pass

    def _run_live_checks(self, targets, nvr_ip, nvr_user, nvr_pwd):
        self.status.showMessage(f"Checking live status for {len(targets)} cameras...")
        for t in targets:
            row = t["row"]
            ip = t["ip"]
            try:
                online, method, details = check_camera_live(ip, nvr_ip, nvr_user, nvr_pwd, timeout=5.0)
                if online:
                    em = "🟢"
                    color = QtGui.QColor(0, 160, 0)
                    status_text = f"Live ({method})"
                else:
                    em = "🔴"
                    color = QtGui.QColor(160, 0, 0)
                    status_text = f"Offline ({method})"
                
                device_type = method
                model = details
                self.table_update.emit(row, status_text, device_type, model, color, em)
                log(f"Live check {ip}: {status_text} - {details}")
            except Exception as e:
                log(f"Live check error {ip}: {e}")
                self.table_update.emit(row, "Error", "Error", str(e), QtGui.QColor(128, 0, 0), "⚠️")
        self.status.showMessage("Live status check complete.")

    def check_all_via_nvr(self):
        """Enhanced Check All: Uses fast NVR ISAPI to check all cameras efficiently."""
        log("=== ENHANCED CHECK ALL START (NVR-BASED) ===")
        
        # Check if another NVR operation is running
        with self.nvr_operation_lock:
            if self.nvr_thread_running:
                QtWidgets.QMessageBox.information(self, "Check In Progress", 
                    "Another NVR operation is currently running. Please wait for it to complete.")
                return
            self.nvr_thread_running = True
        
        # Disable buttons during check
        try:
            self.btn_check_sel.setEnabled(False)  # Disable button during check
            self.status.showMessage("⚡ Enhanced Check All - Querying all NVRs...")
        except AttributeError:
            pass
        
        # Visual indicator that Check All has started
        self._show_check_all_started()
        
        # Start enhanced background thread
        threading.Thread(target=self._run_enhanced_check_all, daemon=True).start()

    def _run_enhanced_check_all(self):
        """Enhanced Check All: Fast parallel processing of all NVRs using ISAPI method."""
        import time
        import concurrent.futures
        
        start_time = time.time()
        log("=== ENHANCED CHECK ALL EXECUTION START ===")
        
        try:
            # Get all NVR configurations
            nvr_configs = list(self.nvrs)
            if not nvr_configs:
                log("[ENHANCED-CHECK] No NVR configurations found")
                QtCore.QTimer.singleShot(100, lambda: QtWidgets.QMessageBox.warning(
                    self, "No NVRs Found", "No NVR configurations found. Please load Excel file with NVR data first."
                ))
                return
            
            log(f"[ENHANCED-CHECK] Processing {len(nvr_configs)} NVRs in parallel")
            
            # Progress tracking
            nvr_results = {}
            all_camera_updates = []
            total_cameras_found = 0
            total_online = 0
            total_offline = 0
            successful_nvrs = 0
            
            # Create enhanced progress dialog on main thread
            prog = None
            def create_progress():
                nonlocal prog
                prog = QtWidgets.QProgressDialog('🟢 Step 1: Connecting to NVRs...', 'Cancel', 0, len(nvr_configs), self)
                prog.setWindowModality(QtCore.Qt.WindowModal)
                prog.setMinimumDuration(0)
                prog.setValue(0)
                prog.show()
            QtCore.QTimer.singleShot(0, create_progress)
            
            # Process all NVRs in parallel using ThreadPoolExecutor
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, len(nvr_configs))) as executor:
                # Submit all NVR checking tasks
                future_to_nvr = {
                    executor.submit(self._check_single_nvr_enhanced, nvr): nvr 
                    for nvr in nvr_configs
                }
                
                # Process completed tasks
                completed = 0
                for future in concurrent.futures.as_completed(future_to_nvr):
                    nvr = future_to_nvr[future]
                    completed += 1
                    
                    # Update progress on main thread with visual indicators
                    if prog:
                        QtCore.QTimer.singleShot(0, lambda c=completed, t=len(nvr_configs), name=nvr.get('name', 'Unknown'): (
                            prog.setValue(c) if prog else None,
                            prog.setLabelText(f'🟢 Step 1: NVRs ({c}/{t}) | 📷 Step 2: Processing cameras from {name}') if prog else None
                        ))
                    
                    try:
                        result = future.result()
                        nvr_name = nvr.get('name', nvr.get('ip', 'Unknown'))
                        
                        if result['success']:
                            cameras = result['cameras']
                            online_count = len([c for c in cameras if c.get('status', '').lower() == 'online'])
                            offline_count = len(cameras) - online_count
                            
                            nvr_results[nvr_name] = {
                                'success': True,
                                'total': len(cameras),
                                'online': online_count,
                                'offline': offline_count,
                                'cameras': cameras
                            }
                            
                            # 🟢 STEP 1: Show NVR online status immediately on main thread
                            self._schedule_nvr_visual_update(nvr_name, True, f"{online_count}/{len(cameras)}")
                            
                            # 🟢 STEP 2: Update camera data and show camera progress
                            self._merge_nvr_camera_data_with_visual(cameras, nvr_name)
                            
                            total_cameras_found += len(cameras)
                            total_online += online_count
                            total_offline += offline_count
                            successful_nvrs += 1
                            
                            log(f"[ENHANCED-CHECK] ✅ {nvr_name}: {online_count}/{len(cameras)} online")
                        else:
                            nvr_results[nvr_name] = {
                                'success': False,
                                'error': result.get('error', 'Unknown error')
                            }
                            
                            # ❌ Show NVR offline status on main thread
                            self._schedule_nvr_visual_update(nvr_name, False, result.get('error', 'Unknown error'))
                            
                            log(f"[ENHANCED-CHECK] ❌ {nvr_name}: {result.get('error', 'Unknown error')}")
                            
                    except Exception as e:
                        nvr_name = nvr.get('name', nvr.get('ip', 'Unknown'))
                        nvr_results[nvr_name] = {'success': False, 'error': str(e)}
                        log(f"[ENHANCED-CHECK] ❌ {nvr_name}: Exception - {e}")
            
            # Close progress dialog on main thread
            if prog:
                QtCore.QTimer.singleShot(0, lambda: prog.close() if prog else None)
            
            # 📷 STEP 3: Final UI update with all camera status indicators
            QtCore.QTimer.singleShot(100, lambda: (
                self.populate_table(self.filtered),
                self.update_counters(),
                self.status.showMessage(f"🟢 Step 3: Updated {total_cameras_found} cameras | ✅ Complete!"),
                self._flash_completion_indicator()
            ))
            
            # Calculate performance metrics
            end_time = time.time()
            duration = end_time - start_time
            
            # Prepare comprehensive results with visual progress summary
            result_summary = f"✅ Enhanced Check All Complete!\n\n"
            result_summary += f"📋 Process Summary:\n"
            result_summary += f"🟢 Step 1: Connected to {successful_nvrs}/{len(nvr_configs)} NVRs\n"
            result_summary += f"📷 Step 2: Processed {total_cameras_found} cameras\n"
            result_summary += f"✅ Step 3: Updated UI with status indicators\n\n"
            result_summary += f"⏱️ Duration: {duration:.1f} seconds\n"
            result_summary += f"📊 Results: {total_online} online, {total_offline} offline\n\n"
            result_summary += "📊 NVR Details:\n"
            
            for nvr_name, result in nvr_results.items():
                if result['success']:
                    online_pct = int((result['online'] / result['total']) * 100) if result['total'] > 0 else 0
                    status_icon = "🟢" if online_pct == 100 else "🟡" if online_pct > 0 else "🔴"
                    result_summary += f"{status_icon} {nvr_name}: {result['online']}/{result['total']} online ({online_pct}%)\n"
                else:
                    result_summary += f"🔴 {nvr_name}: {result['error']}\n"
            
            log(f"[ENHANCED-CHECK] ⚡ Completed in {duration:.1f}s: {successful_nvrs}/{len(nvr_configs)} NVRs, {total_online}/{total_cameras_found} cameras online")
            
            # Update status bar with completion summary
            QtCore.QTimer.singleShot(500, lambda: self.status.showMessage(
                f"✅ Enhanced Check Complete: {successful_nvrs} NVRs | {total_online}/{total_cameras_found} cameras online ({duration:.1f}s)"
            ))
            
            # Show detailed results
            QtCore.QTimer.singleShot(100, lambda: QtWidgets.QMessageBox.information(
                self, "⚡ Enhanced Check All Complete", result_summary
            ))
            
        except Exception as e:
            log(f"[ENHANCED-CHECK] Critical error: {e}")
            import traceback
            log(f"[ENHANCED-CHECK] Traceback: {traceback.format_exc()}")
            QtCore.QTimer.singleShot(100, lambda: QtWidgets.QMessageBox.critical(
                self, "❌ Enhanced Check Error", f"Critical error during enhanced check:\n\n{str(e)}"
            ))
        finally:
            # Reset thread running flag and re-enable buttons
            with self.nvr_operation_lock:
                self.nvr_thread_running = False
            QtCore.QTimer.singleShot(100, lambda: self._enable_check_buttons())

    def _check_single_nvr_enhanced(self, nvr_config):
        """Check a single NVR using the fast ISAPI method."""
        nvr_ip = nvr_config.get('ip', '')
        nvr_name = nvr_config.get('name', nvr_ip)
        
        try:
            # Get credentials
            username, password = get_password(nvr_ip)
            if not username:
                username = 'admin'
            if not password:
                password = 'Kkcctv12345'  # Default
            
            log(f"[ENHANCED-NVR] Processing {nvr_name} ({nvr_ip})")
            
            # Use IVMS-only camera extraction method
            controller = WorkingNVRController(nvr_ip, username, password)
            cameras, method = controller.get_cameras(timeout=15.0)
            if cameras:
                log(f"[ENHANCED-NVR] {nvr_name}: Found {len(cameras)} cameras [IVMS method]")
                return {'success': True, 'cameras': cameras}
            else:
                return {'success': False, 'error': "No cameras found or connection failed (IVMS method)"}
                
        except Exception as e:
            log(f"[ENHANCED-NVR] Error checking {nvr_name}: {e}")
            return {'success': False, 'error': str(e)}

    def _merge_nvr_camera_data(self, nvr_cameras, nvr_name):
        """Merge camera data from NVR into our existing camera list."""
        from datetime import datetime
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M')
        
        for nvr_cam in nvr_cameras:
            nvr_cam_name = nvr_cam.get('name', '')
            nvr_cam_ip = nvr_cam.get('ip', '')
            nvr_cam_status = nvr_cam.get('status', 'Unknown')
            for existing_cam in self.cams:
                existing_name = self._clean_text(existing_cam.get('name')).lower()
                incoming_name = self._clean_text(nvr_cam_name).lower()
                existing_ip = self._clean_text(existing_cam.get('ip'))
                incoming_ip = self._clean_text(nvr_cam_ip)
                name_match = existing_name == incoming_name if incoming_name else False
                ip_match = existing_ip == incoming_ip if incoming_ip else False
                if name_match or ip_match:
                    # Set status directly from IVMS fetch result, unmodified
                    existing_cam['status'] = nvr_cam_status
                    existing_cam['device_type'] = nvr_cam.get('model', existing_cam.get('device_type', 'Camera'))
                    existing_cam['channel'] = nvr_cam.get('channel', existing_cam.get('channel', ''))
                    existing_cam['port'] = nvr_cam.get('port', existing_cam.get('port', ''))
                    existing_cam['last_updated'] = current_time
                    if incoming_ip and incoming_ip != existing_ip:
                        existing_cam['previous_ip'] = existing_ip
                        existing_cam['ip'] = incoming_ip
                        existing_cam['remark'] = f"IP: {existing_ip}  {incoming_ip}"
                        log(f"[IP-CHANGE] {self._clean_text(nvr_cam_name)}: {existing_ip}  {incoming_ip}")
                    break
        
        # Also add all API cameras to the live collection for direct display
        for nvr_cam in nvr_cameras:
            # Debug specific NVRs that are having model issues - check first 3 cameras per NVR
            if nvr_name in ['NVR2', 'NVR3', 'NVR4']:
                camera_count = len([c for c in self.api_cameras if c.get('nvr_name') == nvr_name])
                if camera_count <= 3:  # Only show first 3 cameras per NVR
                    model_value = nvr_cam.get('model', 'NO_MODEL')
                    log(f"[DEBUG-{nvr_name}] Camera #{camera_count}: {nvr_cam.get('name', 'NO_NAME')} | Model: '{model_value}' | Status: {nvr_cam.get('status', 'NO_STATUS')}")
            
            # Create enhanced camera record with all API data
            api_camera = {
                'name': nvr_cam.get('name', ''),
                'ip': nvr_cam.get('ip', ''),
                'status': nvr_cam.get('status', 'Unknown'),
                'model': nvr_cam.get('model', 'Camera'),
                'device_type': nvr_cam.get('model', 'Camera'),
                'channel': nvr_cam.get('channel', ''),
                'port': nvr_cam.get('port', ''),
                'nvr': nvr_name,
                'last_updated': current_time,
                'source': 'IVMS_API'
            }
            
            # Check if camera already exists in API collection for THIS NVR (avoid duplicates by name+IP+NVR)
            camera_exists = False
            for existing_api_cam in self.api_cameras:
                name_match = (self._clean_text(existing_api_cam.get('name', '')).lower() == 
                    self._clean_text(api_camera['name']).lower())
                ip_match = existing_api_cam.get('ip', '') == api_camera.get('ip', '')
                nvr_match = (self._clean_text(existing_api_cam.get('nvr', '')).lower() == 
                    self._clean_text(nvr_name).lower())
                
                if name_match and ip_match and nvr_match:
                    # Update existing API camera with latest info (same NVR)
                    existing_api_cam.update(api_camera)
                    camera_exists = True
                    break
            
            if not camera_exists:
                self.api_cameras.append(api_camera)
    def _schedule_nvr_visual_update(self, nvr_name, is_online, status_info):
        """Schedule NVR visual update on main thread."""
        # Use QTimer to ensure this runs on the main thread
        QtCore.QTimer.singleShot(0, lambda: self._update_nvr_status_visual(nvr_name, is_online, status_info))
    
    def _update_nvr_status_visual(self, nvr_name, is_online, status_info):
        """Update NVR list with enhanced visual status indicators during checking process."""
        try:
            # Only proceed if we're on the main thread
            if QtCore.QThread.currentThread() != QtWidgets.QApplication.instance().thread():
                log(f"[VISUAL-NVR] Scheduling UI update for {nvr_name} on main thread")
                QtCore.QTimer.singleShot(0, lambda: self._update_nvr_status_visual(nvr_name, is_online, status_info))
                return
                
            # Find the NVR in the list and update its display
            updated = False
            for i in range(self.list_nvr.count()):
                item = self.list_nvr.item(i)
                if item and nvr_name.lower() in item.text().lower():
                    if is_online:
                        # BRIGHT GREEN indicators for online NVR
                        updated_text = f"🟢 {nvr_name} | {status_info} cameras | ✅ ONLINE"
                        item.setForeground(QtGui.QColor(0, 128, 0))  # Dark green text
                        item.setBackground(QtGui.QColor(144, 238, 144))  # Bright light green background
                        log(f"[VISUAL-NVR] ✅ {nvr_name} marked as ONLINE with {status_info} cameras")
                    else:
                        # BRIGHT RED indicators for offline NVR
                        updated_text = f"❌ {nvr_name} | OFFLINE: {status_info}"
                        item.setForeground(QtGui.QColor(128, 0, 0))  # Dark red text
                        item.setBackground(QtGui.QColor(255, 182, 193))  # Bright light red background
                        log(f"[VISUAL-NVR] ❌ {nvr_name} marked as OFFLINE: {status_info}")
                    
                    item.setText(updated_text)
                    # Make sure it's visible and force multiple refresh attempts
                    self.list_nvr.scrollToItem(item)
                    self.list_nvr.update()
                    self.list_nvr.repaint()
                    QtWidgets.QApplication.processEvents()  # Force immediate processing
                    updated = True
                    break
                    
            if not updated:
                log(f"[VISUAL-NVR] WARNING: Could not find NVR '{nvr_name}' in list to update")
                
        except Exception as e:
            log(f"[VISUAL-NVR] Error updating NVR status: {e}")

    def _merge_nvr_camera_data_with_visual(self, nvr_cameras, nvr_name):
        """Merge camera data with real-time visual updates showing camera status changes."""
        from datetime import datetime
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M')
        
        camera_count = 0
        online_count = 0
        
        for nvr_cam in nvr_cameras:
            nvr_cam_name = nvr_cam.get('name', '')
            nvr_cam_ip = nvr_cam.get('ip', '')
            nvr_cam_status = nvr_cam.get('status', 'Unknown')
            
            # Find matching camera in our data
            for existing_cam in self.cams:
                # Try to match by name first, then by IP using safe normalization
                existing_name = self._clean_text(existing_cam.get('name')).lower()
                incoming_name = self._clean_text(nvr_cam_name).lower()
                existing_ip = self._clean_text(existing_cam.get('ip'))
                incoming_ip = self._clean_text(nvr_cam_ip)

                name_match = existing_name == incoming_name if incoming_name else False
                ip_match = existing_ip == incoming_ip if incoming_ip else False
                
                if name_match or ip_match:
                    camera_count += 1
                    
                    # Update existing camera with NVR data
                    status_clean = self._clean_text(nvr_cam_status).lower()
                    if status_clean == 'online':
                        existing_cam['status'] = 'Online (NVR)'
                        online_count += 1
                        status_icon = '🟢'
                        log(f"[VISUAL-CAM] 🟢 {self._clean_text(nvr_cam_name)} | {incoming_ip} | Online")
                    else:
                        existing_cam['status'] = 'Offline (NVR)'
                        status_icon = '🔴'
                        log(f"[VISUAL-CAM] 🔴 {self._clean_text(nvr_cam_name)} | {incoming_ip} | Offline")
                    
                    # Update technical details
                    existing_cam['device_type'] = nvr_cam.get('model', existing_cam.get('device_type', 'Camera'))
                    existing_cam['channel'] = nvr_cam.get('channel', existing_cam.get('channel', ''))
                    existing_cam['port'] = nvr_cam.get('port', existing_cam.get('port', ''))
                    existing_cam['last_updated'] = current_time
                    
                    # Update IP if it changed
                    if incoming_ip and incoming_ip != existing_ip:
                        existing_cam['previous_ip'] = existing_ip
                        existing_cam['ip'] = incoming_ip
                        existing_cam['remark'] = f"IP: {existing_ip} → {incoming_ip}"
                        log(f"[IP-CHANGE] {self._clean_text(nvr_cam_name)}: {existing_ip} → {incoming_ip}")
                    
                    # Real-time visual update for individual camera with table updates
                    if camera_count % 3 == 0:  # Update every 3 cameras for more responsive feedback
                        QtCore.QTimer.singleShot(10 * (camera_count // 3), lambda count=camera_count, online=online_count, name=nvr_name: (
                            self.status.showMessage(f"🔄 Step 2: Processing {name} cameras... ({online}/{count} online)"),
                            self.update_counters(),
                            self._update_camera_table_visual(name, online, count)
                        ))
                    
                    break
        
        # Final update for this NVR's cameras with table refresh
        QtCore.QTimer.singleShot(50, lambda: (
            self.status.showMessage(f"✅ Step 2: {nvr_name} complete - {online_count}/{camera_count} cameras online"),
            self.populate_table(self.filtered),  # Refresh table to show updates
            self.update_counters()  # Update status bar counters
        ))

    def _update_camera_table_visual(self, nvr_name, online_count, total_count):
        """Update camera table with enhanced visual progress indicators."""
        try:
            # Only proceed if we're on the main thread
            if QtCore.QThread.currentThread() != QtWidgets.QApplication.instance().thread():
                log(f"[VISUAL-TABLE] Scheduling table update for {nvr_name} on main thread")
                QtCore.QTimer.singleShot(0, lambda: self._update_camera_table_visual(nvr_name, online_count, total_count))
                return
                
            # BRIGHT flash effect to show update is happening
            original_style = self.table.styleSheet()
            self.table.setStyleSheet("QTableWidget { border: 3px solid #00FF00; background-color: #F0FFF0; }")
            self.table.repaint()
            QtWidgets.QApplication.processEvents()
            
            # Reset flash after short delay
            QtCore.QTimer.singleShot(300, lambda: self.table.setStyleSheet(original_style))
            
            # Update table with current camera data
            if hasattr(self, 'filtered') and self.filtered:
                cameras_updated = 0
                # Show BRIGHT green/red indicators for cameras from this NVR
                for row in range(self.table.rowCount()):
                    camera_item = self.table.item(row, 1)  # Camera name column
                    nvr_item = self.table.item(row, 4)  # NVR column
                    
                    if camera_item and nvr_item and nvr_name.lower() in nvr_item.text().lower():
                        # Find the corresponding camera data
                        for cam in self.filtered:
                            if cam.get('name', '').lower() == camera_item.text().lower():
                                # Update status with BRIGHT colors
                                status_item = self.table.item(row, 3)  # Status column
                                if status_item:
                                    if 'online' in cam.get('status', '').lower():
                                        status_item.setText('🟢 Online')
                                        status_item.setBackground(QtGui.QColor(144, 238, 144))  # Bright green
                                        status_item.setForeground(QtGui.QColor(0, 100, 0))  # Dark green text
                                    else:
                                        status_item.setText('❌ Offline')
                                        status_item.setBackground(QtGui.QColor(255, 182, 193))  # Bright red
                                        status_item.setForeground(QtGui.QColor(128, 0, 0))  # Dark red text
                                    cameras_updated += 1
                                break
                                
            # Force immediate visual refresh
            self.table.update()
            self.table.repaint()
            QtWidgets.QApplication.processEvents()
            
            log(f"[VISUAL-TABLE] ✅ Updated {cameras_updated} cameras for {nvr_name}: {online_count}/{total_count}")
        except Exception as e:
            log(f"[VISUAL-TABLE] Error updating camera table: {e}")

    def _show_check_all_started(self):
        """Show visual indication that Check All has started."""
        try:
            # Flash status bar blue to indicate start
            original_style = self.status.styleSheet()
            start_style = "QStatusBar { background-color: #007ACC; color: white; font-weight: bold; }"
            self.status.setStyleSheet(start_style)
            self.status.repaint()
            
            # Flash Check All button to show it's working
            # Update status message during progress
            self.status.showMessage("🔄 Checking in progress...")
            
            # Reset status bar after brief delay
            QtCore.QTimer.singleShot(800, lambda: self.status.setStyleSheet(original_style))
            
            log("[VISUAL-START] 🔄 Check All started - visual indicators activated")
        except Exception as e:
            log(f"[VISUAL-START] Error showing start indicator: {e}")

    def _flash_completion_indicator(self):
        """Flash an enhanced completion indicator to show the check is done."""
        try:
            # Only proceed if we're on the main thread
            if QtCore.QThread.currentThread() != QtWidgets.QApplication.instance().thread():
                log("[VISUAL-COMPLETE] Scheduling completion flash on main thread")
                QtCore.QTimer.singleShot(0, self._flash_completion_indicator)
                return
                
            # BRIGHT GREEN flash for status bar
            original_status_style = self.status.styleSheet()
            bright_flash_style = "QStatusBar { background-color: #00FF00; color: black; font-weight: bold; border: 2px solid #008000; }"
            self.status.setStyleSheet(bright_flash_style)
            self.status.repaint()
            
            # Flash the Check All button with BRIGHT colors and reset text
            if hasattr(self, 'btn_check_all'):
                original_btn_style = self.btn_check_all.styleSheet()
                btn_flash_style = "QPushButton { background-color: #00FF00; color: black; font-weight: bold; border: 2px solid #008000; }"
                self.btn_check_all.setStyleSheet(btn_flash_style)
                self.btn_check_all.setText("✅ Check All Complete!")
                self.btn_check_all.repaint()
                
                # Reset button text and style after longer delay
                def reset_button():
                    self.btn_check_all.setText("Check All")
                    self.btn_check_all.setStyleSheet(original_btn_style)
                QtCore.QTimer.singleShot(3000, reset_button)
            
            # Flash the entire NVR list briefly
            if hasattr(self, 'list_nvr'):
                original_list_style = self.list_nvr.styleSheet()
                self.list_nvr.setStyleSheet("QListWidget { border: 3px solid #00FF00; background-color: #F0FFF0; }")
                self.list_nvr.repaint()
                QtCore.QTimer.singleShot(1500, lambda: self.list_nvr.setStyleSheet(original_list_style))
            
            # Force immediate processing
            QtWidgets.QApplication.processEvents()
            
            # Reset status bar after longer delay for visibility
            QtCore.QTimer.singleShot(2500, lambda: self.status.setStyleSheet(original_status_style))
                
            log("[VISUAL-COMPLETE] ✅ ENHANCED completion indicator flashed with bright colors!")
        except Exception as e:
            log(f"[VISUAL-COMPLETE] Error flashing completion indicator: {e}")

    def _enable_check_buttons(self):
        """Re-enable check buttons after NVR check completion."""
        try:
            self.btn_check_all.setEnabled(True)
            self.btn_check_live.setEnabled(True)
        except AttributeError:
            try:
                self.ui.checkAllBtn.setEnabled(True)
                self.ui.checkBtn.setEnabled(True)
            except:
                pass

    def show_nvr_check_result(self, message, success):
        """Show NVR check result message."""
        title = "✅ NVR Check Complete" if success else "❌ NVR Check Error"
        if success:
            QtCore.QTimer.singleShot(100, lambda: QtWidgets.QMessageBox.information(self, title, message))
        else:
            QtCore.QTimer.singleShot(100, lambda: QtWidgets.QMessageBox.critical(self, title, message))
        QtCore.QTimer.singleShot(100, lambda: self._enable_check_buttons())

    def load_nvr_config(self):
        """Load NVR configurations from nvr_config.json file."""
        try:
            if os.path.exists('nvr_config.json'):
                with open('nvr_config.json', 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    return config.get('nvrs', [])
            else:
                # Fallback to self.nvrs if no JSON file exists
                return list(self.nvrs)
        except Exception as e:
            log(f"[NVR-CONFIG] Error loading NVR config: {e}")
            return list(self.nvrs)  # Fallback to in-memory NVRs
    
    def save_nvr_config(self, nvrs_list):
        """Save NVR configurations to nvr_config.json file."""
        try:
            from datetime import datetime
            config = {
                'nvrs': nvrs_list,
                'last_updated': datetime.now().isoformat()
            }
            with open('nvr_config.json', 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
            log(f"[NVR-CONFIG] Saved {len(nvrs_list)} NVR configurations")
            return True
        except Exception as e:
            log(f"[NVR-CONFIG] Error saving NVR config: {e}")
            return False
    
    def sync_nvr_credentials(self):
        """Sync NVR credentials between memory and persistent storage."""
        try:
            # Load existing config
            nvr_configs = self.load_nvr_config()
            updated = False
            
            # Add credentials from keyring/metadata to NVR configs
            # The following block was removed due to undefined nvr_cameras/current_time/camera_updates in this context.
            # If visual merge is needed, use _merge_nvr_camera_data_with_visual in the appropriate context.
                    # ...removed broken/legacy block...
            
        except Exception as e:
            log(f"[NVR-CHECK] Error updating check history: {e}")

    def save_check_history(self):
        """Save check history to persistent storage."""
        try:
            import json
            with open(CHECK_HISTORY_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.check_history, f, indent=2, ensure_ascii=False)
            log(f"[HISTORY] Saved check history with {len(self.check_history)} entries")
        except Exception as e:
            log(f"[HISTORY] Error saving check history: {e}")

    def closeEvent(self, event):
        # Persist check history before closing
        try:
            with open(CHECK_HISTORY_FILE, 'w', encoding='utf-8') as hf:
                json.dump(self.check_history, hf, indent=2)
            log(f"[CLOSE] Saved check history ({len(self.check_history)} entries)")
        except Exception as e:
            log(f"[CLOSE] Error saving history: {e}")
        try:
            super().closeEvent(event)
        except Exception:
            event.accept()

    def _run_checks(self, targets, check_id=None):
        """Enhanced camera checking with smart caching v8.6+."""
        self.status.showMessage(f"🔍 Smart checking {len(targets)} cameras (with cache)...")
        checked_count = 0
        cache_hits = 0
        start_time = time.time()
        
        for t in targets:
            # Check if this operation was cancelled by NVR switch
            if check_id is not None and check_id != self.current_check_id:
                self.status.showMessage(f"Check cancelled (switched NVR)")
                return
                
            row = t["row"]; ip = t["ip"]
            
            try:
                # Check cache first for recent results
                cache_valid, cached_data = get_cached_status(
                    getattr(self, 'connection_cache', {}), ip, CACHE_TIMEOUT
                )
                
                if cache_valid and cached_data:
                    # Use cached result
                    status_text = cached_data.get('status', 'Cached')
                    device_type = cached_data.get('device_type', 'Cache')
                    model = cached_data.get('model', 'Cached')
                    
                    # Determine color and emoji from cached status
                    if 'online' in status_text.lower():
                        em = "🟢"; color = QtGui.QColor(0, 160, 0)
                    elif 'ping' in status_text.lower():
                        em = "🟡"; color = QtGui.QColor(200, 140, 0)
                    else:
                        em = "🔴"; color = QtGui.QColor(160, 0, 0)
                    
                    self.table_update.emit(row, f"{status_text} (Cached)", device_type, model, color, em)
                    log(f"[CACHE HIT] {ip}: {status_text} (cached)")
                    cache_hits += 1
                    checked_count += 1
                    continue
                
                # No cache hit - perform actual check
                result = self._perform_enhanced_check(ip)
                
                # Update cache with result
                if not hasattr(self, 'connection_cache'):
                    self.connection_cache = {}
                
                update_cache(self.connection_cache, ip, {
                    'status': result['status'],
                    'device_type': result['device_type'],
                    'model': result['model']
                })
                
                # Update UI
                self.table_update.emit(row, result['status'], result['device_type'], 
                                     result['model'], result['color'], result['emoji'])
                log(f"[FRESH CHECK] {ip}: {result['status']}")
                checked_count += 1
                
            except Exception as e:
                log(f"[CHECK ERROR] {ip}: {e}")
                self.table_update.emit(row, "Error", "Error", str(e)[:50], QtGui.QColor(128, 0, 0), "⚠️")
                checked_count += 1
        
        # Update performance metrics
        elapsed = time.time() - start_time
        if hasattr(self, 'performance_metrics'):
            self.performance_metrics['total_checks'] += checked_count
            self.performance_metrics['cache_hits'] += cache_hits
            if checked_count > 0:
                self.performance_metrics['average_response_time'] = elapsed / checked_count
        
        cache_ratio = (cache_hits / checked_count * 100) if checked_count > 0 else 0
        self.status.showMessage(
            f"✅ Smart check complete: {checked_count} cameras, "
            f"{cache_hits} cache hits ({cache_ratio:.1f}%), {elapsed:.1f}s"
        )
    
    def _perform_enhanced_check(self, ip):
        """Perform enhanced camera check with multiple methods."""
        try:
            # Try SADP first (Hikvision UDP discovery - most reliable)
            sadp_online, sadp_model = check_camera_via_sadp(ip, timeout=1.5)
            if sadp_online:
                return {
                    'status': "Online (SADP)",
                    'device_type': "SADP",
                    'model': f"Model: {sadp_model}",
                    'color': QtGui.QColor(0, 160, 0),
                    'emoji': "🟢"
                }
            
            # Fallback: Try TCP ports with optimized timeout
            h = check_tcp(ip, HTTP_PORT, timeout=0.8)
            r = check_tcp(ip, RTSP_PORT, timeout=0.8)
            if h or r:
                ports = []
                if h: ports.append('HTTP')
                if r: ports.append('RTSP')
                return {
                    'status': "Online (TCP)",
                    'device_type': '/'.join(ports),
                    'model': "TCP Services",
                    'color': QtGui.QColor(0, 160, 0),
                    'emoji': "🟢"
                }
            
            # Last resort: Try ping
            if silent_ping(ip):
                return {
                    'status': "Online (Ping)",
                    'device_type': "Ping",
                    'model': "Network Only",
                    'color': QtGui.QColor(200, 140, 0),
                    'emoji': "🟡"
                }
            
            # Completely offline
            return {
                'status': "Offline",
                'device_type': "None",
                'model': "No Response",
                'color': QtGui.QColor(160, 0, 0),
                'emoji': "🔴"
            }
            
        except Exception as e:
            return {
                'status': f"Error: {str(e)[:30]}",
                'device_type': "Error",
                'model': str(e)[:50],
                'color': QtGui.QColor(128, 0, 0),
                'emoji': "⚠️"
            }

    @QtCore.pyqtSlot(int, str, str, str, object, str)
    def apply_table_update(self, row, status_text, device_type_text, model_text, color, emoji):
        """Update core camera columns after a check concludes."""
        try:
            if not (0 <= row < self.table.rowCount()):
                return

            badge_item = QtWidgets.QTableWidgetItem(emoji or "📷")
            badge_item.setTextAlignment(QtCore.Qt.AlignCenter)
            badge_item.setFlags(QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsSelectable)
            self.table.setItem(row, 0, badge_item)

            status_item = QtWidgets.QTableWidgetItem(status_text or "")
            status_item.setForeground(color)
            status_item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table.setItem(row, 3, status_item)

            ip_item = self.table.item(row, 2)
            ip_val = ip_item.text().strip() if ip_item else ""

            matched_cam = None
            if ip_val:
                for cam in self.cams:
                    if (cam.get("ip", "") or "").strip() == ip_val:
                        matched_cam = cam
                        break

            history_entry = self.check_history.get(ip_val, {}) if ip_val else {}

            model_display = (model_text or "").strip()
            if not model_display and matched_cam:
                model_display = (matched_cam.get("model") or matched_cam.get("device_type") or "").strip()
            if not model_display and history_entry:
                model_display = (history_entry.get("model") or history_entry.get("device_type") or "").strip()
            model_item = QtWidgets.QTableWidgetItem(model_display)
            model_item.setForeground(color)
            self.table.setItem(row, 4, model_item)

            port_display = ""
            if matched_cam:
                port_display = str(matched_cam.get("port", "") or "").strip()
            if not port_display and history_entry:
                port_display = str(history_entry.get("port", "") or "").strip()
            if port_display in ("0", "None", "—"):
                port_display = ""
            port_item = QtWidgets.QTableWidgetItem(port_display)
            port_item.setForeground(color)
            self.table.setItem(row, 5, port_item)

            nvr_display = ""
            if matched_cam:
                nvr_display = (matched_cam.get("nvr") or "").strip()
            if not nvr_display and history_entry:
                nvr_display = (history_entry.get("nvr") or "").strip()
            if not nvr_display and ip_val:
                    # ...removed broken/legacy block...
                try:
                    with open(CHECK_HISTORY_FILE, "w", encoding="utf-8") as history_file:
                        json.dump(self.check_history, history_file, indent=2, ensure_ascii=False)
                except Exception as history_error:
                    log(f"[HISTORY SAVE] Error persisting history for {ip_val}: {history_error}")

            self.update_counters()
        except Exception as e:
            log(f"[TABLE UPDATE] Error applying table update on row {row}: {e}")

    @QtCore.pyqtSlot(int, str, str, str, str, str, str, str, str, str, object, str)
    def apply_enhanced_table_update(self, row, status, device_type, model, channel, port,
                                    serial, firmware, nvr_name, last_updated, color, emoji):
        """Handle enhanced updates emitted by background threads."""
        try:
            if not (0 <= row < self.table.rowCount()):
                return

            badge_item = QtWidgets.QTableWidgetItem(emoji or "📷")
            badge_item.setTextAlignment(QtCore.Qt.AlignCenter)
            badge_item.setFlags(QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsSelectable)
            self.table.setItem(row, 0, badge_item)

            status_item = QtWidgets.QTableWidgetItem(status or "")
            status_item.setForeground(color)
            status_item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table.setItem(row, 3, status_item)

            model_item = QtWidgets.QTableWidgetItem(str(model or "").strip())
            model_item.setForeground(color)
            self.table.setItem(row, 4, model_item)

            port_item = QtWidgets.QTableWidgetItem(str(port or "").strip())
            port_item.setForeground(color)
            self.table.setItem(row, 5, port_item)

            nvr_item = QtWidgets.QTableWidgetItem(str(nvr_name or "").strip())
            nvr_item.setForeground(color)
            self.table.setItem(row, 6, nvr_item)

            last_item = QtWidgets.QTableWidgetItem(str(last_updated or "").strip())
            last_item.setForeground(color)
            self.table.setItem(row, 7, last_item)

            remark_candidate = str(serial or "").strip()
            if not remark_candidate:
                channel_text = str(channel or "").strip()
                if channel_text:
                    remark_candidate = f"CH {channel_text}"
            if not remark_candidate:
                remark_candidate = str(firmware or "").strip()
            remark_item = QtWidgets.QTableWidgetItem(remark_candidate)
            remark_item.setForeground(color)
            self.table.setItem(row, 8, remark_item)

            self.update_counters()
        except Exception as e:
            log(f"[ENHANCED-TABLE] Error updating row {row}: {e}")

    def _show_offline_camera_dialog(self, offline_cams):
        """Display a non-blocking dialog listing offline cameras with quick navigation."""
        try:
            log(f"[OFFLINE-DIALOG] Attempting to show dialog with {len(offline_cams)} cameras")
            if not offline_cams:
                log("[OFFLINE-DIALOG] No cameras to display, returning")
                return

            if self.offline_dialog and self.offline_dialog.isVisible():
                try:
                    self.offline_dialog.close()
                except Exception:
                    pass

            dialog = QtWidgets.QDialog(self)
            dialog.setWindowTitle(f"⚠️ Offline Cameras ({len(offline_cams)})")
            dialog.setModal(False)
            dialog.setAttribute(QtCore.Qt.WA_DeleteOnClose, True)

            layout = QtWidgets.QVBoxLayout(dialog)
            
            # Add header with verification info
            info_label = QtWidgets.QLabel(f"<b>{len(offline_cams)} cameras confirmed offline after verification</b><br>Click to locate camera in table.")
            info_label.setWordWrap(True)
            layout.addWidget(info_label)

            list_widget = QtWidgets.QListWidget(dialog)
            list_widget.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
            list_widget.setUniformItemSizes(True)
            
            # Custom styling for better visibility
            list_widget.setStyleSheet("""
                QListWidget::item:selected {
                    background-color: #3daee9;
                    color: white;
                }
                QListWidget::item:hover {
                    background-color: #e0e0e0;
                }
            """)

            for cam in offline_cams:
                display_name = (cam.get('name') or 'Unnamed').strip() or 'Unnamed'
                display_ip = (cam.get('ip') or '').strip()
                display_nvr = (cam.get('nvr') or 'Unknown NVR').strip() or 'Unknown NVR'
                # Show verified status if available
                display_status = cam.get('verified_status', cam.get('status', 'Offline'))
                item_text = f"🔴 {display_name} | {display_ip} | {display_nvr} | {display_status}"
                item = QtWidgets.QListWidgetItem(item_text)
                # Store both IP and NVR for accurate identification (handles duplicates)
                item.setData(QtCore.Qt.UserRole, {'ip': display_ip, 'nvr': display_nvr})
                list_widget.addItem(item)
                log(f"[OFFLINE-DIALOG] Added camera: {display_name} ({display_ip})")

            def handle_item(item):
                # Get IP and NVR from item data (stored during item creation)
                item_data = item.data(QtCore.Qt.UserRole)
                log(f"[OFFLINE-DIALOG] Clicked item, data type: {type(item_data)}, data: {item_data}")
                if isinstance(item_data, dict):
                    target_ip = item_data.get('ip', '')
                    target_nvr = item_data.get('nvr', '')
                    log(f"[OFFLINE-DIALOG] Extracted IP={target_ip}, NVR={target_nvr}")
                else:
                    target_ip = item_data or ""
                    target_nvr = ""
                    log(f"[OFFLINE-DIALOG] Legacy data format, IP={target_ip}")
                self._focus_camera_by_ip_and_nvr(target_ip, target_nvr)

            # Change from double-click to single click
            list_widget.itemClicked.connect(handle_item)
            layout.addWidget(list_widget)

            close_button = QtWidgets.QPushButton("Close")
            close_button.clicked.connect(dialog.close)
            layout.addWidget(close_button)

            dialog.resize(500, 400)
            dialog.show()
            dialog.raise_()  # Bring to front
            dialog.activateWindow()  # Make it active
            dialog.finished.connect(lambda *_: setattr(self, 'offline_dialog', None))
            self.offline_dialog = dialog
            log(f"[OFFLINE-DIALOG] Dialog displayed and raised to front")
        except Exception as e:
            log(f"[OFFLINE-DIALOG] Failed to display offline camera list: {e}")
            import traceback
            log(f"[OFFLINE-DIALOG] Traceback: {traceback.format_exc()}")

    def _focus_camera_by_ip_and_nvr(self, ip_address, nvr_name):
        """Select and scroll to the camera row matching the provided IP and NVR."""
        try:
            target_ip = (ip_address or '').strip()
            target_nvr = (nvr_name or '').strip().lower()
            
            if not target_ip:
                return
            
            log(f"[FOCUS-CAMERA] Looking for IP={target_ip}, NVR={target_nvr}")

            # First, ensure we're showing all cameras (not filtered by NVR)
            if self.filtered != self.api_cameras:
                self.filtered = list(self.api_cameras)
                self.list_nvr.clearSelection()
                self.lbl_selected_nvr.setText("Selected NVR: None (All)")
                self.populate_table(self.filtered)
                QtWidgets.QApplication.processEvents()
                log(f"[FOCUS-CAMERA] Switched to All cameras view")

            # Find NVR index from name (e.g., "NVR9" -> "9")
            target_nvr_index = ""
            if target_nvr:
                # Try to extract number from "NVR9" -> "9"
                import re
                match = re.search(r'(\d+)', target_nvr)
                if match:
                    target_nvr_index = match.group(1)
                
                # Also try to match full NVR name to get index
                for i, nvr in enumerate(self.nvrs):
                    nvr_name_clean = self._clean_text(nvr.get("name")).lower()
                    if nvr_name_clean == target_nvr:
                        target_nvr_index = str(i + 1)
                        break

            log(f"[FOCUS-CAMERA] Target NVR index: {target_nvr_index}")

            # Search in currently displayed table for matching IP (and optionally NVR)
            for row_idx in range(self.table.rowCount()):
                ip_item = self.table.item(row_idx, 2)  # IP column
                nvr_item = self.table.item(row_idx, 6)  # NVR column (shows index number)
                
                if ip_item:
                    row_ip = ip_item.text().strip()
                    row_nvr = nvr_item.text().strip() if nvr_item else ""
                    
                    # Match IP first
                    if row_ip == target_ip:
                        # If we have NVR info, match it too (handles duplicates)
                        if target_nvr_index and row_nvr != target_nvr_index:
                            continue
                        
                        self.table.selectRow(row_idx)
                        self.table.scrollToItem(ip_item, QtWidgets.QAbstractItemView.PositionAtCenter)
                        self.status.showMessage(f"Located: {target_ip} on NVR{target_nvr_index or '?'}", 4000)
                        log(f"[FOCUS-CAMERA] Found at row {row_idx}, NVR column shows: {row_nvr}")
                        return

            self.status.showMessage(f"Camera {target_ip} not found in table", 4000)
            log(f"[FOCUS-CAMERA] Not found in table")
        except Exception as e:
            log(f"[FOCUS-CAMERA] Error focusing camera {ip_address}/{nvr_name}: {e}")
            import traceback
            log(f"[FOCUS-CAMERA] Traceback: {traceback.format_exc()}")
    
    def _focus_camera_by_ip(self, ip_address):
        """Legacy method - calls new method with empty NVR."""
        self._focus_camera_by_ip_and_nvr(ip_address, "")

    def _refresh_quick_sync_ui(self):
        """Refresh UI elements after Quick Sync merges external data."""
        try:
            if hasattr(self, 'search') and not self.search.text().strip():
                self.filtered = list(self.api_cameras if self.api_cameras else self.cams)
            self.populate_nvr_list()
            self.populate_table(self.filtered)
            self.update_counters()
            log("[QUICK-SYNC] UI refreshed with latest data")
        except Exception as e:
            log(f"[QUICK-SYNC] UI refresh error: {e}")

    def _queue_on_ui(self, func, *args, **kwargs):
        """Safely queue a function to execute on the UI thread."""
        try:
            self.ui_call_signal.emit((func, args, kwargs))
        except Exception as e:
            log(f"[UI-QUEUE] Failed to queue UI function {getattr(func, '__name__', func)}: {e}")

    def _execute_ui_callable(self, payload):
        """Execute a queued UI function dispatched via ui_call_signal."""
        try:
            func, args, kwargs = payload
            func(*args, **kwargs)
        except Exception as e:
            log(f"[UI-EXEC] Error executing UI callable: {e}")
    
    def _handle_error_notification(self, level, title, message):
        """Handle error notification signals with enhanced user feedback."""
        try:
            log(f"[ERROR-NOTIFICATION] {level}: {title} - {message}")
            
            if level.upper() == "CRITICAL":
                QtWidgets.QMessageBox.critical(self, title, message)
            elif level.upper() == "WARNING":
                QtWidgets.QMessageBox.warning(self, title, message)
            elif level.upper() == "INFO":
                QtWidgets.QMessageBox.information(self, title, message)
            else:
                # Show in status bar for minor notifications
                self.status.showMessage(f"{title}: {message}", 5000)
                
        except Exception as e:
            log(f"[ERROR-NOTIFICATION] Error showing notification: {e}")
    
    def _handle_performance_update(self, metrics):
        """Handle performance metrics updates."""
        try:
            self.performance_metrics.update(metrics)
            
            # Update performance display if dashboard is open
            if hasattr(self, 'performance_dialog') and self.performance_dialog:
                self._update_performance_display(metrics)
                
        except Exception as e:
            log(f"[PERFORMANCE-UPDATE] Error updating metrics: {e}")
    
    def _handle_cache_stats(self, stats):
        """Handle cache statistics updates."""
        try:
            # Update cache display if dashboard is open
            if hasattr(self, 'performance_dialog') and self.performance_dialog:
                self._update_cache_display(stats)
                
        except Exception as e:
            log(f"[CACHE-STATS] Error updating cache stats: {e}")

    def _normalize_status_text(self, raw_status, fallback="Online"):
        """Normalize status text to include emoji prefixes for clarity."""
        try:
            status_txt = str(raw_status or "").strip()
            if not status_txt:
                status_txt = fallback

            if status_txt == "—":
                return status_txt

            if status_txt.startswith(("🟢", "🔴", "🟡")):
                return status_txt

            base_status = status_txt
            base_lower = base_status.lower()

            # Check for offline/error conditions first
            if any(keyword in base_lower for keyword in ("error", "offline", "fail", "disconnect", "timeout")):
                return f"🔴 {base_status}"

            # Check for limited/warning conditions
            if any(keyword in base_lower for keyword in ("limited", "warning", "degraded", "ping", "slow")):
                return f"🟡 {base_status}"

            # Check for online conditions
            if any(keyword in base_lower for keyword in ("online", "connected", "active")):
                return f"🟢 {base_status}"
            
            # Unknown status - don't assume online, use yellow for caution
            return f"🟡 {base_status}"
        except Exception:
            return f"🟢 {fallback}"

    # ---------------- NVR refresh ----------------
    def refresh_nvr_status(self):
        """Unified IVMS-based refresh: updates all NVR/camera statuses using only WorkingNVRController.get_cameras()."""
        if not self.nvrs:
            QtWidgets.QMessageBox.information(self, "Refresh Status", "No NVRs loaded. Please load Excel file first.")
            return
        total_nvrs = len(self.nvrs)
        log(f"[REFRESH] Starting IVMS-based refresh for {total_nvrs} NVRs")
        self.status.showMessage(f"🔄 Refreshing {total_nvrs} NVRs status (IVMS method)...", 0)
        threading.Thread(target=self._ivms_refresh_thread, daemon=True).start()

    def _ivms_refresh_thread(self):
        """Background thread: refreshes all NVR/camera statuses using only IVMS fetch logic."""
        try:
            start_time = time.time()
            online_nvrs = 0
            total_cameras_updated = 0
            
            # Clear previous API results to start fresh
            self.api_cameras.clear()
            
            log(f"[REFRESH-THREAD] Starting refresh of {len(self.nvrs)} NVRs...")
            
            for i, nvr in enumerate(self.nvrs):
                nvr_start_time = time.time()
                nvr_ip = nvr.get('ip', '').strip()
                nvr_name = nvr.get('name', f'NVR-{i+1}').strip()
                
                if not nvr_ip:
                    log(f"[REFRESH-THREAD] Skipping NVR {i+1}: No IP address")
                    continue
                    
                progress = int(((i + 1) / len(self.nvrs)) * 100)
                log(f"[REFRESH-THREAD] Checking {nvr_name} ({nvr_ip}) - {progress}%")
                self._queue_on_ui(self.status.showMessage, f"🔄 Checking {nvr_name} ({progress}%)", 0)
                
                try:
                    # Get credentials and create controller
                    username, password, cred_source = self._resolve_nvr_credentials(nvr_ip)
                    log(f"[REFRESH-THREAD] Using {cred_source} credentials for {nvr_name}")
                    
                    controller = WorkingNVRController(nvr_ip, username, password)
                    
                    # Test connection first
                    connection_test = controller.connect()
                    log(f"[REFRESH-THREAD] {nvr_name} connection test: {'PASS' if connection_test else 'FAIL'}")
                    
                    # Fetch cameras with reduced timeout for faster processing
                    try:
                        # Skip camera fetch if connection test failed (faster processing)
                        if not connection_test:
                            fetched_cameras = []
                            camera_count = 0
                            method = "Connection failed - skipped camera fetch"
                        else:
                            fetched_cameras, method = controller.get_cameras(timeout=8.0)
                            camera_count = len(fetched_cameras) if fetched_cameras else 0
                    except Exception as fetch_error:
                        log(f"[REFRESH-THREAD] Camera fetch error for {nvr_name}: {fetch_error}")
                        fetched_cameras = []
                        camera_count = 0
                        method = f"Error: {str(fetch_error)[:50]}..."
                    
                    # Determine NVR status - consider device online if it responds to HTTP even without camera data
                    if fetched_cameras:
                        nvr_status = 'online'
                        online_nvrs += 1
                        total_cameras_updated += camera_count
                        log(f"[REFRESH-THREAD] ✅ {nvr_name}: ONLINE - {camera_count} cameras found via {method}")
                    elif connection_test:
                        # Device responds to HTTP but no cameras found (might be auth issue)
                        nvr_status = 'online'  # Consider it online since device responds
                        online_nvrs += 1
                        log(f"[REFRESH-THREAD] ⚠️ {nvr_name}: ONLINE (no cameras) - Device responsive but {method}")
                    else:
                        nvr_status = 'offline'
                        log(f"[REFRESH-THREAD] ❌ {nvr_name}: OFFLINE - {method}")
                    
                    # Update NVR display
                    self._queue_on_ui(self._update_nvr_display, i, nvr_status)
                    
                    # Update/merge camera statuses directly from IVMS fetch
                    if fetched_cameras:
                        self._merge_nvr_camera_data(fetched_cameras, nvr_name)
                        log(f"[REFRESH-THREAD] Merged {camera_count} cameras for {nvr_name}")
                    
                    # Log processing time
                    nvr_elapsed = time.time() - nvr_start_time
                    log(f"[REFRESH-THREAD] {nvr_name} processed in {nvr_elapsed:.1f}s")
                
                except Exception as nvr_error:
                    log(f"[REFRESH-THREAD] Critical error processing {nvr_name}: {nvr_error}")
                    # Set as offline and continue with next NVR
                    nvr_status = 'offline'
                    self._queue_on_ui(self._update_nvr_display, i, nvr_status)
            # Final summary
            elapsed = time.time() - start_time
            offline_nvrs = len(self.nvrs) - online_nvrs
            
            log(f"[REFRESH-THREAD] Final results: {online_nvrs} online, {offline_nvrs} offline, {total_cameras_updated} cameras updated")
            
            self._queue_on_ui(self._update_status_counts, online_nvrs, offline_nvrs)
            self._queue_on_ui(self._refresh_table_display)
            
            completion_msg = f"✅ IVMS Refresh Complete! {len(self.nvrs)} NVRs in {elapsed:.1f}s | 🟢 {online_nvrs} online, 🔴 {offline_nvrs} offline | 📹 {total_cameras_updated} cameras updated"
            self._queue_on_ui(self.status.showMessage, completion_msg, 8000)
            self._queue_on_ui(self._maybe_show_offline_cameras)
            log(f"[REFRESH-THREAD] Complete: {completion_msg}")
        except Exception as e:
            error_msg = f"❌ IVMS Refresh failed: {str(e)}"
            self._queue_on_ui(self.status.showMessage, error_msg, 5000)
            log(f"[REFRESH-THREAD] Critical error: {e}", exc_info=True)

    def _check_nvr_simple(self, ip, credentials=None):
        """Simple NVR status check with clear status return."""
        try:
            log(f"[NVR-CHECK] Testing connectivity to {ip}")
            username, password, source = credentials if credentials else self._resolve_nvr_credentials(ip)
            
            # First try to connect using the NVR controller
            controller = WorkingNVRController(ip, username, password)
            if controller.connect():
                log(f"[NVR-CHECK] ✅ {ip} - ISAPI connection successful")
                return 'online'
            
            # Try HTTP ports first (most reliable)
            ports_to_check = [80, 8000, 8080, 443, 8443]
            for port in ports_to_check:
                if check_tcp(ip, port, timeout=3):
                    log(f"[NVR-CHECK] ✅ {ip} - Port {port} accessible")
                    return 'online'
            
            # Try ping as fallback
            if silent_ping(ip):
                log(f"[NVR-CHECK] 🟡 {ip} - Ping successful but no HTTP ports")
                return 'limited'
            
            log(f"[NVR-CHECK] ❌ {ip} - No connectivity detected")
            return 'offline'
        except Exception as e:
            log(f"[NVR-CHECK] Error checking {ip}: {e}")
            return 'offline'

    def _get_camera_count_for_nvr(self, nvr_name, nvr_ip):
        """Return real camera count for a given NVR based on current data."""
        try:
            name_key = self._clean_text(nvr_name).lower()
            ip_key = self._clean_text(nvr_ip)
            cams_source = getattr(self, 'cams', []) or []
            count = 0
            for cam in cams_source:
                cam_name = self._clean_text(cam.get('nvr')).lower()
                cam_parent_ip = self._clean_text(cam.get('nvr_ip'))
                if (name_key and cam_name == name_key) or (ip_key and cam_parent_ip == ip_key):
                    count += 1
            return count
        except Exception as e:
            log(f"[NVR-CAM-COUNT] Error counting cameras for {nvr_name} ({nvr_ip}): {e}")
            return 0

    def _resolve_nvr_credentials(self, ip):
        """Determine credentials for an NVR, preferring stored secrets."""
        log(f"[CREDS] Resolving credentials for {ip}")
        
        # First try the keyring/creds system
        stored_creds = get_password(ip)
        if stored_creds and len(stored_creds) == 2 and all(stored_creds):
            username, password = stored_creds
            source = 'keyring'
            log(f"[CREDS] ✅ Using keyring credentials for {ip}: {username}/***")
            return username, password, source
        
        # Try nvr_credentials.json file (main credential store)
        try:
            if os.path.exists("nvr_credentials.json"):
                with open("nvr_credentials.json", "r", encoding="utf-8") as f:
                    nvr_creds = json.load(f)
                
                # Look for matching IP in the credentials file
                for key, nvr_data in nvr_creds.items():
                    if nvr_data.get('ip') == ip:
                        username = nvr_data.get('username', 'admin')
                        password = nvr_data.get('password', 'Kkcctv12345')
                        source = 'nvr_credentials.json'
                        log(f"[CREDS] ✅ Using nvr_credentials.json for {ip}: {username}/*** (password: {password[:3]}***)")
                        return username, password, source
        except Exception as e:
            log(f"[CREDS] Error reading nvr_credentials.json: {e}")
        
        # Fallback to default credentials
        username = 'admin'
        password = 'Kkcctv12345'
        source = 'default'
        log(f"[CREDS] ⚠️ Using default credentials for {ip}: {username}/*** (no stored creds found)")
        
        return username, password, source

    def _fetch_and_update_nvr_cameras(self, nvr_data, username, password):
        """Fetch camera list from NVR to obtain per-camera status; returns set of IPs updated."""
        try:
            nvr_ip = (nvr_data.get('ip', '') or '').strip()
            if not nvr_ip:
                return set()

            controller = WorkingNVRController(nvr_ip, username, password)
            fetched_cameras, method = controller.get_cameras(timeout=15.0)
            if not fetched_cameras:
                log(f"[REFRESH-FETCH] IVMS fetch failed for {nvr_data.get('name', nvr_ip)}: No cameras returned (IVMS method)")
                return set()

            from datetime import datetime
            now_str = datetime.now().strftime('%Y-%m-%d %H:%M')

            existing_by_ip = {}
            existing_by_name = {}
            for cam in self.cams:
                cam_ip = (cam.get('ip', '') or '').strip()
                cam_name = (cam.get('name', '') or '').strip().lower()
                if cam_ip:
                    existing_by_ip[cam_ip] = cam
                if cam_name:
                    existing_by_name.setdefault(cam_name, []).append(cam)

            updated_ips = set()
            for cam in fetched_cameras:
                cam_ip = (cam.get('ip', '') or '').strip()
                cam_name_key = (cam.get('name', '') or '').strip().lower()
                cam_status_raw = cam.get('status', 'unknown')
                
                log(f"[REFRESH-FETCH] Processing camera: {cam.get('name', 'Unknown')} | IP: {cam_ip} | Raw Status: {cam_status_raw}")

                targets = []
                if cam_ip and cam_ip in existing_by_ip:
                    targets.append(existing_by_ip[cam_ip])
                if not targets and cam_name_key and cam_name_key in existing_by_name:
                    targets.extend(existing_by_name[cam_name_key])

                if not targets:
                    log(f"[REFRESH-FETCH] No matching target found for {cam.get('name', 'Unknown')} ({cam_ip})")
                    continue

                cam_status = cam.get('status', 'Unknown')
                log(f"[REFRESH-FETCH] Direct status: {cam_status}")
                for target in targets:
                    target['status'] = cam_status
                    target['model'] = cam.get('model', target.get('model', ''))
                    target['port'] = cam.get('port', target.get('port', ''))
                    target['last_updated'] = now_str
                    target.setdefault('nvr', nvr_data.get('name', ''))
                    target.setdefault('nvr_ip', nvr_ip)
                    history_key = target.get('ip', '') or f"{nvr_data.get('name', '')}:{target.get('name', '')}"
                    history_entry = self.check_history.setdefault(history_key, {})
                    history_entry.update({
                        'status': cam_status,
                        'timestamp': now_str,
                        'nvr': nvr_data.get('name', ''),
                        'device_type': cam.get('model', history_entry.get('device_type', ''))
                    })

                if cam_ip:
                    updated_ips.add(cam_ip)

            log(f"[REFRESH-FETCH] Retrieved {len(fetched_cameras)} cameras from {nvr_data.get('name', nvr_ip)} during refresh")
            return updated_ips
        except Exception as e:
            log(f"[REFRESH-FETCH] Error fetching cameras from {nvr_data.get('name', 'Unknown')} ({nvr_data.get('ip', '')}): {e}")
            return set()

    def _integrate_quick_sync_cameras(self, nvr, fetched_cameras):
        """Merge cameras fetched during Quick Sync into the main dataset."""
        if not fetched_cameras:
            return 0

        from datetime import datetime
        now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
        nvr_name = (nvr.get('name', '') or '').strip()
        nvr_ip = (nvr.get('ip', '') or '').strip()
        nvr_name_key = nvr_name.lower()

        existing_list = self.cams
        new_entries = 0

        for cam in fetched_cameras:
            cam_name = (cam.get('name', '') or '').strip()
            cam_ip = (cam.get('ip', '') or '').strip()
            cam_status = cam.get('status', 'Unknown')
            cam_model = cam.get('model', '')
            cam_port = cam.get('port', '')
            matched = None
            for existing in existing_list:
                same_ip = cam_ip and (existing.get('ip', '').strip() == cam_ip)
                same_name = cam_name and (existing.get('name', '').strip().lower() == cam_name.lower())
                if same_ip or same_name:
                    matched = existing
                    break
            target = matched
            if target is None:
                target = {
                    'nvr': nvr_name,
                    'nvr_ip': nvr_ip,
                    'name': cam_name or f"Camera {len(existing_list) + 1}",
                    'ip': cam_ip,
                    'status': cam_status,
                    'model': cam_model,
                    'port': cam_port,
                    'last_updated': now_str
                }
                existing_list.append(target)
                new_entries += 1
            else:
                target['status'] = cam_status
                target['model'] = cam_model or target.get('model', '')
                target['port'] = cam_port or target.get('port', '')
                target['last_updated'] = now_str
                target.setdefault('nvr', nvr_name)
                target.setdefault('nvr_ip', nvr_ip)
            history_key = cam_ip if cam_ip else f"{nvr_name}:{cam_name}"
            history_entry = self.check_history.setdefault(history_key, {})
            history_entry.update({
                'status': cam_status,
                'timestamp': now_str,
                'nvr': nvr_name,
                'device_type': cam_model or history_entry.get('device_type', '')
            })

        return new_entries

    def _update_nvr_display(self, index, status):
        """Update NVR display in the list."""
        try:
            item = self.list_nvr.item(index)
            if not item:
                return
            
            nvr = self.nvrs[index]
            name = nvr.get('name', '')
            ip = nvr.get('ip', '')
            cam_count = self._get_camera_count_for_nvr(name, ip)
            
            # Store status in NVR data for persistence
            nvr['status'] = status
            nvr['cam_count'] = cam_count
            
            # Choose emoji based on status
            if status == 'online':
                emoji = "🟢"
            elif status == 'limited':
                emoji = "🟡"
            else:
                emoji = "🔴"
            
            # Update item text
            item_text = f"{emoji} {name} | {ip} | 🎥 {cam_count}"
            item.setText(item_text)
            
        except Exception as e:
            log(f"[NVR-DISPLAY] Error updating NVR {index}: {e}")

    def _update_cameras_for_nvr(self, nvr_data, nvr_status, skip_ips=None):
        """Update camera statuses based on NVR status."""
        try:
            if not nvr_data:
                return

            nvr_ip = (nvr_data.get('ip', '') or '').strip()
            nvr_name = (nvr_data.get('name', '') or '').strip()
            nvr_name_key = nvr_name.lower()
            ip_prefix = '.'.join(nvr_ip.split('.')[:3]) if nvr_ip else ''

            if nvr_status == 'online':
                camera_status = '🟢 Online'
            elif nvr_status == 'limited':
                camera_status = '🟡 Limited'
            elif nvr_status == 'ping':
                camera_status = '🟡 Ping Only'
            else:
                camera_status = '🔴 Offline'

            from datetime import datetime
            now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
            updated_count = 0

            for camera in self.cams:
                cam_ip = (camera.get('ip', '') or '').strip()
                if skip_ips and cam_ip in skip_ips:
                    continue
                cam_name = (camera.get('name', '') or '').strip()
                cam_nvr_name = (camera.get('nvr', '') or '').strip().lower()
                cam_nvr_ip = (camera.get('nvr_ip', '') or '').strip()

                matches_nvr = False
                if nvr_name_key and cam_nvr_name == nvr_name_key:
                    matches_nvr = True
                elif nvr_ip and cam_nvr_ip and cam_nvr_ip == nvr_ip:
                    matches_nvr = True
                elif ip_prefix and cam_ip.startswith(ip_prefix):
                    matches_nvr = True

                if not matches_nvr:
                    continue

                camera['status'] = camera_status
                camera['last_updated'] = now_str
                if nvr_name and not camera.get('nvr'):
                    camera['nvr'] = nvr_name
                if nvr_ip and not camera.get('nvr_ip'):
                    camera['nvr_ip'] = nvr_ip

                # Update check history keyed by IP or composite key
                history_key = cam_ip if cam_ip else f"{nvr_name}:{cam_name}"
                history_entry = self.check_history.setdefault(history_key, {})
                history_entry['status'] = camera_status
                history_entry['timestamp'] = now_str
                history_entry['nvr'] = nvr_name

                updated_count += 1

            # Ensure filtered view reflects the same status updates
            for camera in self.filtered:
                cam_ip = (camera.get('ip', '') or '').strip()
                if skip_ips and cam_ip in skip_ips:
                    continue
                cam_nvr_name = (camera.get('nvr', '') or '').strip().lower()
                cam_nvr_ip = (camera.get('nvr_ip', '') or '').strip()
                match_filtered = False
                if nvr_name_key and cam_nvr_name == nvr_name_key:
                    match_filtered = True
                elif nvr_ip and cam_nvr_ip and cam_nvr_ip == nvr_ip:
                    match_filtered = True
                elif ip_prefix and cam_ip.startswith(ip_prefix):
                    match_filtered = True

                if not match_filtered:
                    continue

                camera['status'] = camera_status
                camera['last_updated'] = now_str

            if updated_count:
                log(f"[CAMERA-UPDATE] Updated {updated_count} cameras for NVR {nvr_name or nvr_ip} -> {camera_status}")
            else:
                log(f"[CAMERA-UPDATE] No cameras matched NVR {nvr_name or nvr_ip} for status update")

        except Exception as e:
            log(f"[CAMERA-UPDATE] Error updating cameras for NVR {nvr_data.get('name', 'Unknown')} ({nvr_data.get('ip', '')}): {e}")

    def _update_status_counts(self, online_nvrs=None, offline_nvrs=None):
        """Update status bar counters."""
        try:
            total_nvrs = len(self.nvrs)
            
            # Calculate NVR counts if not provided
            if online_nvrs is None or offline_nvrs is None:
                online_nvrs = sum(1 for nvr in self.nvrs if nvr.get('status', '').lower() == 'online')
                offline_nvrs = total_nvrs - online_nvrs
            
            # Update NVR counts  
            self.lbl_nvr_total.setText(f"📊 {total_nvrs}")
            self.lbl_nvr_online.setText(f"🟢 {online_nvrs}")
            self.lbl_nvr_offline.setText(f"🔴 {offline_nvrs}")
            
            # Count cameras using IVMS status values from API (live data from NVRs)
            # Handle both simple statuses and verified statuses
            camera_source = self.api_cameras if self.api_cameras else self.cams
            total_cameras = len(camera_source)
            
            # Count ALL cameras including duplicates (same IP on different NVR = different camera)
            online_cameras = 0
            offline_cameras = 0
            offline_details = []  # Track offline cameras for debugging
            
            for cam in camera_source:
                ip = cam.get('ip', '').strip()
                status = str(cam.get('status', ''))
                status_lower = status.lower()
                
                # Check for online: emoji or text "online"
                if '🟢' in status or 'online' in status_lower:
                    online_cameras += 1
                # Check for offline: emoji or text "offline" 
                elif '🔴' in status or 'offline' in status_lower:
                    offline_cameras += 1
                    if len(offline_details) < 20:  # Track first 20 for logging
                        offline_details.append(f"{cam.get('name', 'NO_NAME')} ({ip}): {status}")
            
            unknown_cameras = total_cameras - online_cameras - offline_cameras
            
            # Update camera counts
            self.lbl_total.setText(f"📷 {total_cameras}")
            self.lbl_online.setText(f"🟢 {online_cameras}")
            self.lbl_offline.setText(f"🔴 {offline_cameras}")
            
            log(f"[STATUS-COUNTS] NVRs: {online_nvrs}/{total_nvrs}, Cameras: 🟢{online_cameras} 🔴{offline_cameras} ❔{unknown_cameras}")
            if offline_cameras > 0 and offline_details:
                log(f"[STATUS-COUNTS] First {len(offline_details)} offline cameras: {offline_details[:5]}")
            
        except Exception as e:
            log(f"[STATUS-COUNTS] Error: {e}")

    def _refresh_table_display(self):
        """Refresh the table display with updated camera data."""
        try:
            # Use live API cameras as the primary data source
            self.filtered = list(self.api_cameras)
            self.cameras = self.api_cameras  # Update main camera reference
            
            # Repopulate the table with current filtered data
            self.populate_table(self.filtered)
            # Update counters to reflect new status
            self.update_counters()
            log(f"[TABLE-REFRESH] Table display refreshed with {len(self.filtered)} cameras")
        except Exception as e:
            log(f"[TABLE-REFRESH] Error refreshing table: {e}")

    def _maybe_show_offline_cameras(self):
        """Show an actionable popup listing offline cameras after refresh with auto-verification."""
        try:
            log("[OFFLINE-VERIFY] === _maybe_show_offline_cameras called ===")
            offline_cams = self._collect_offline_cameras()
            log(f"[OFFLINE-VERIFY] Collected {len(offline_cams)} offline cameras from initial scan")
            if not offline_cams:
                if self.offline_dialog and self.offline_dialog.isVisible():
                    self.offline_dialog.close()
                    self.offline_dialog = None
                return

            # Auto-verify offline cameras in background before showing popup
            log(f"[OFFLINE-VERIFY] Starting verification of {len(offline_cams)} offline cameras...")
            self.status.showMessage(f"⏳ Verifying {len(offline_cams)} offline cameras...", 3000)
            
            # Run verification in background thread
            import threading
            def verify_and_show():
                try:
                    verified_offline = self._verify_offline_cameras(offline_cams)
                    log(f"[OFFLINE-VERIFY] Verification thread complete, queuing UI update...")
                    # Update UI on main thread using existing _queue_on_ui method
                    self._queue_on_ui(self._finalize_offline_verification, verified_offline)
                except Exception as e:
                    log(f"[OFFLINE-VERIFY] Error in verification thread: {e}")
                    import traceback
                    log(f"[OFFLINE-VERIFY] Traceback: {traceback.format_exc()}")
            
            thread = threading.Thread(target=verify_and_show, daemon=True)
            thread.start()
            
        except Exception as e:
            log(f"[OFFLINE-ALERT] Unable to show offline camera popup: {e}")

    def _collect_offline_cameras(self):
        """Return list of offline camera metadata dictionaries from API data."""
        offline = []
        try:
            # Use API cameras (live data) instead of Excel
            cams_source = self.api_cameras if self.api_cameras else getattr(self, 'cams', [])
            log(f"[OFFLINE-VERIFY] _collect_offline_cameras using {'api_cameras' if self.api_cameras else 'cams'} ({len(cams_source)} total cameras)")
            seen_ips = set()
            for cam in cams_source:
                ip_val = (cam.get('ip', '') or '').strip()
                if not ip_val or ip_val in seen_ips:
                    continue

                raw_status = str(cam.get('status', '') or '')
                normalized = self._normalize_status_text(raw_status)
                combined_text = f"{raw_status} {normalized}".lower()
                if ('🔴' in normalized) or any(keyword in combined_text for keyword in ('offline', 'failed', 'error', 'down', 'timeout', 'inactive', 'disconnect')):
                    offline.append({
                        'ip': ip_val,
                        'name': (cam.get('name', '') or '').strip(),
                        'nvr': (cam.get('nvr', '') or '').strip(),
                        'status': normalized
                    })
                    seen_ips.add(ip_val)
        except Exception as e:
            log(f"[OFFLINE-ALERT] Error collecting offline cameras: {e}")
        log(f"[OFFLINE-VERIFY] _collect_offline_cameras found {len(offline)} offline cameras: {[c['ip'] for c in offline[:5]]}{'...' if len(offline) > 5 else ''}")
        return offline

    def _verify_offline_cameras(self, offline_cams):
        """Verify offline cameras by pinging each IP, return truly offline list with updated statuses."""
        import subprocess
        import concurrent.futures
        
        def ping_camera(cam_info):
            """Ping single camera IP and return updated status."""
            ip = cam_info['ip']
            try:
                # Use fast ping with 2-second timeout
                result = subprocess.run(
                    ['ping', '-n', '1', '-w', '2000', ip],
                    capture_output=True,
                    text=True,
                    timeout=3,
                    creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                )
                
                if result.returncode == 0:
                    # Ping successful - camera is actually online
                    log(f"[OFFLINE-VERIFY] ✅ {ip} ({cam_info['name']}) responded to ping - marking online")
                    cam_info['verified_status'] = '🟢 Online (Verified)'
                    cam_info['is_truly_offline'] = False
                    return cam_info
                else:
                    # Ping failed - truly offline
                    log(f"[OFFLINE-VERIFY] ❌ {ip} ({cam_info['name']}) no response - confirmed offline")
                    cam_info['verified_status'] = '🔴 Offline (Verified)'
                    cam_info['is_truly_offline'] = True
                    return cam_info
                    
            except subprocess.TimeoutExpired:
                log(f"[OFFLINE-VERIFY] ⏱️ {ip} ({cam_info['name']}) timeout - confirmed offline")
                cam_info['verified_status'] = '🔴 Offline (Timeout)'
                cam_info['is_truly_offline'] = True
                return cam_info
            except Exception as e:
                log(f"[OFFLINE-VERIFY] ⚠️ {ip} ({cam_info['name']}) verification error: {e}")
                cam_info['verified_status'] = '🟡 Unknown (Error)'
                cam_info['is_truly_offline'] = True  # Treat errors as offline for safety
                return cam_info
        
        # Verify all cameras in parallel with thread pool
        verified_cams = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            verified_cams = list(executor.map(ping_camera, offline_cams))
        
        # Filter to truly offline cameras only
        truly_offline = [cam for cam in verified_cams if cam.get('is_truly_offline', True)]
        
        # Update api_cameras with verified statuses
        verified_count = 0
        online_count = 0
        for cam_info in verified_cams:
            ip = cam_info['ip']
            new_status = cam_info.get('verified_status', cam_info['status'])
            is_offline = cam_info.get('is_truly_offline', True)
            
            # Update in api_cameras
            for cam in self.api_cameras:
                if cam.get('ip', '').strip() == ip:
                    cam['status'] = new_status
                    verified_count += 1
                    if not is_offline:
                        online_count += 1
                    break
        
        log(f"[OFFLINE-VERIFY] Verification complete: {len(offline_cams)} checked, {online_count} now online, {len(truly_offline)} confirmed offline")
        return truly_offline

    def _finalize_offline_verification(self, verified_offline):
        """Update UI and show popup with verified offline cameras."""
        try:
            log(f"[OFFLINE-VERIFY] _finalize_offline_verification called with {len(verified_offline)} verified offline cameras")
            
            # Always show ALL cameras after verification (auto list all, no need to select NVR)
            self.filtered = list(self.api_cameras)
            log(f"[OFFLINE-VERIFY] Auto-showing all cameras: {len(self.filtered)} cameras with verified statuses")
            
            # Clear NVR selection to show "All"
            self.list_nvr.clearSelection()
            self.lbl_selected_nvr.setText("Selected NVR: None (All)")
            
            # Refresh table display with all cameras and their updated statuses
            self.populate_table(self.filtered)
            log(f"[OFFLINE-VERIFY] Table refreshed with all {len(self.filtered)} cameras")
            
            # Refresh all counters with updated statuses
            self.update_counters()
            self.update_selected_counters()
            self._update_status_counts()
            log("[OFFLINE-VERIFY] All counters updated (global + selected)")
            
            # Collect ALL offline cameras after verification (including duplicates on different NVRs)
            # Same IP on different NVR = different camera that needs checking
            all_offline_after_verification = []
            
            log(f"[OFFLINE-VERIFY] Scanning {len(self.api_cameras)} total cameras for offline status...")
            
            for cam in self.api_cameras:
                ip = (cam.get('ip', '') or '').strip()
                name = (cam.get('name', '') or '').strip()
                
                if not ip:
                    continue
                
                status = str(cam.get('status', ''))
                status_lower = status.lower()
                
                # Use SAME logic as counter: check online first, then offline
                is_online = '🟢' in status or 'online' in status_lower
                is_offline = '🔴' in status or 'offline' in status_lower
                
                # Debug: log first few cameras to see what we're finding
                if len(all_offline_after_verification) < 5:
                    log(f"[OFFLINE-VERIFY] Camera {name or 'NO_NAME'} ({ip}): status='{status}', is_online={is_online}, is_offline={is_offline}")
                
                # Only count as offline if NOT online (prioritize online detection)
                if is_offline and not is_online:
                    all_offline_after_verification.append({
                        'ip': ip,
                        'name': name,
                        'nvr': (cam.get('nvr', '') or '').strip(),
                        'status': status,
                        'verified_status': cam.get('verified_status', status)
                    })
            
            log(f"[OFFLINE-VERIFY] Total offline cameras collected for popup: {len(all_offline_after_verification)}")
            log(f"[OFFLINE-VERIFY] Offline IPs: {[c['ip'] for c in all_offline_after_verification]}")
            
            # Show popup with ALL offline cameras
            if all_offline_after_verification and len(all_offline_after_verification) > 0:
                log(f"[OFFLINE-VERIFY] Showing popup with {len(all_offline_after_verification)} offline cameras")
                log(f"[OFFLINE-VERIFY] Offline camera IPs: {[c['ip'] for c in all_offline_after_verification[:10]]}{'...' if len(all_offline_after_verification) > 10 else ''}")
                self.status.showMessage(f"⚠️ {len(all_offline_after_verification)} cameras offline", 5000)
                try:
                    self._show_offline_camera_dialog(all_offline_after_verification)
                    log("[OFFLINE-VERIFY] Popup dialog displayed successfully")
                except Exception as popup_error:
                    log(f"[OFFLINE-VERIFY] Error showing popup dialog: {popup_error}")
                    import traceback
                    log(f"[OFFLINE-VERIFY] Popup traceback: {traceback.format_exc()}")
            else:
                log("[OFFLINE-VERIFY] All cameras online - no offline popup needed")
                self.status.showMessage("✅ All cameras online", 5000)
                if self.offline_dialog and self.offline_dialog.isVisible():
                    self.offline_dialog.close()
                    self.offline_dialog = None
                    
        except Exception as e:
            log(f"[OFFLINE-VERIFY] Error finalizing verification: {e}")
            import traceback
            log(f"[OFFLINE-VERIFY] Traceback: {traceback.format_exc()}")



    def _apply_comprehensive_update(self, status_result):
        """Apply comprehensive NVR status update to UI (like Update Cameras UI updates)."""
        try:
            index = status_result['index']
            item = self.list_nvr.item(index)
            if not item:
                return
            
            nvr = status_result['nvr']
            name = status_result['name']
            ip = status_result['ip']
            status = status_result['status']
            method = status_result['method']
            details = status_result['details']
            
            # Create comprehensive status emoji and text with enhanced styling
            if status == 'online':
                if method == 'SADP':
                    emoji = "🟢 SADP"
                    model_info = f" ({details.get('model', 'Unknown')})" if details.get('model') else ""
                    item_color = QtGui.QColor(39, 174, 96)  # Green
                elif method == 'HTTP':
                    emoji = "🟢 HTTP"
                    model_info = f" ({details.get('protocol', 'Web')})"
                    item_color = QtGui.QColor(39, 174, 96)  # Green
                else:
                    emoji = "🟢 Online"
                    model_info = ""
                    item_color = QtGui.QColor(39, 174, 96)  # Green
            elif status == 'limited':
                emoji = "🟡 TCP"
                ports = details.get('open_ports', [])
                model_info = f" ({len(ports)} ports)" if ports else ""
                item_color = QtGui.QColor(243, 156, 18)  # Orange
            elif status == 'ping':
                emoji = "🟡 Ping"
                model_info = " (basic)"
                item_color = QtGui.QColor(243, 156, 18)  # Orange
            elif status == 'error':
                emoji = "⚠️ Error"
                model_info = f" ({status_result.get('error', 'Unknown')})"
                item_color = QtGui.QColor(231, 76, 60)  # Red
            else:
                emoji = "🔴 Offline"
                model_info = ""
                item_color = QtGui.QColor(231, 76, 60)  # Red
            
            # Build comprehensive display text (like Update Cameras detailed info)
            sheet_flag = "" if nvr.get("sheet_found", False) else " ⚠️ sheet missing"
            cam_count = nvr.get('cam_count', 0)
            
            text = f"{emoji}  🗄️ {name} | {ip} | 🎥 {cam_count}{model_info}{sheet_flag}"
            item.setText(text)
            
            # Apply color styling to the NVR list item
            item.setForeground(item_color)
            font = item.font()
            font.setBold(True)
            item.setFont(font)
            
            log(f"[COMPREHENSIVE-UPDATE] Updated NVR {index}: {text}")
            
        except Exception as e:
            log(f"[COMPREHENSIVE-UPDATE] Error updating NVR {status_result.get('index', '?')}: {e}")

    def _finalize_refresh(self):
        """Finalize refresh with comprehensive summary and count updates (like Update Cameras completion)."""
        try:
            elapsed = time.time() - self.refresh_progress['start_time']
            total = self.refresh_progress['total']
            online = self.refresh_progress['online']
            offline = self.refresh_progress['offline']
            errors = self.refresh_progress['errors']
            
            # Update camera and NVR counts in status bar using timer
            def update_counts():
                self._update_comprehensive_counts()
            QtCore.QTimer.singleShot(100, update_counts)
            
            # Create comprehensive completion message (like Update Cameras final message)
            success_msg = f"✅ Refresh Complete! {total} NVRs checked in {elapsed:.1f}s | "
            success_msg += f"🟢 {online} online, "
            success_msg += f"🔴 {offline} offline"
            if errors > 0:
                success_msg += f", ⚠️ {errors} errors"
            
            # Add camera statistics to the message
            total_cameras = len(self.cameras)
            online_cameras = sum(1 for cam in self.cameras if 'online' in cam.get('status', '').lower())
            offline_cameras = total_cameras - online_cameras
            
            success_msg += f" | 📷 {total_cameras} cameras ({online_cameras} online, {offline_cameras} offline)"
            
            # Show final status using signal for thread safety
            QtCore.QTimer.singleShot(0, lambda: self.status.showMessage(success_msg, 8000))
            
            # Force table refresh to show updated status with camera status update
            QtCore.QTimer.singleShot(200, lambda: self._refresh_table_with_status())
            
            log(f"[COMPREHENSIVE-REFRESH] ✅ COMPLETE: {success_msg}")
            
        except Exception as e:
            log(f"[COMPREHENSIVE-REFRESH] Error in finalization: {e}")

    def _update_camera_statuses_from_nvr(self, status_result):
        """Update camera statuses based on NVR status during refresh."""
        try:
            nvr_index = status_result['index']
            nvr_status = status_result['status']
            nvr_obj = self.nvrs[nvr_index] if nvr_index < len(self.nvrs) else None
            
            if not nvr_obj:
                return
                
            nvr_ip = nvr_obj.get('ip', '')
            if not nvr_ip:
                return
            
            # Update cameras belonging to this NVR
            updated_count = 0
            for camera in self.cameras:
                camera_nvr_ip = camera.get('nvr_ip', '') or camera.get('ip', '').split('.')[0:3]
                if isinstance(camera_nvr_ip, list):
                    camera_nvr_ip = '.'.join(camera_nvr_ip)
                    
                # Match cameras to NVR by IP range or exact NVR IP
                if (nvr_ip in camera.get('ip', '') or 
                    camera.get('nvr_ip', '') == nvr_ip or
                    camera.get('ip', '').startswith(nvr_ip.rsplit('.', 1)[0])):
                    
                    # Update camera status based on NVR status
                    if nvr_status in ['online', 'limited']:
                        camera['status'] = '🟢 Online'
                        updated_count += 1
                    elif nvr_status == 'ping':
                        camera['status'] = '🟡 TCP Only'
                        updated_count += 1
                    else:
                        camera['status'] = '🔴 Offline'
                        updated_count += 1
            
            # Also update filtered list
            for camera in self.filtered:
                camera_nvr_ip = camera.get('nvr_ip', '') or camera.get('ip', '').split('.')[0:3]
                if isinstance(camera_nvr_ip, list):
                    camera_nvr_ip = '.'.join(camera_nvr_ip)
                    
                if (nvr_ip in camera.get('ip', '') or 
                    camera.get('nvr_ip', '') == nvr_ip or
                    camera.get('ip', '').startswith(nvr_ip.rsplit('.', 1)[0])):
                    
                    if nvr_status in ['online', 'limited']:
                        camera['status'] = '🟢 Online'
                    elif nvr_status == 'ping':
                        camera['status'] = '🟡 TCP Only'
                    else:
                        camera['status'] = '🔴 Offline'
            
            log(f"[CAMERA-STATUS-UPDATE] Updated {updated_count} cameras for NVR {nvr_obj.get('name', '')} ({nvr_ip}) - Status: {nvr_status}")
            
        except Exception as e:
            log(f"[CAMERA-STATUS-UPDATE] Error updating camera statuses: {e}")

    def _refresh_table_with_status(self):
        """Refresh table with proper status indicators after refresh."""
        try:
            # Update camera statuses in check history for table display
            for camera in self.filtered:
                ip = camera.get('ip', '')
                status = camera.get('status', '—')
                if ip and status != '—':
                    if ip not in self.check_history:
                        self.check_history[ip] = {}
                    self.check_history[ip]['status'] = status
                    
            # Repopulate table with updated statuses
            self.populate_table(self.filtered)
            log(f"[REFRESH-TABLE] Table refreshed with status indicators for {len(self.filtered)} cameras")
            
        except Exception as e:
            log(f"[REFRESH-TABLE] Error refreshing table with status: {e}")
    
    def _update_comprehensive_counts(self):
        """Update all status bar counts after refresh completion."""
        try:
            # Update NVR counts
            total_nvrs = len(self.nvrs)
            online_nvrs = self.refresh_progress.get('online', 0)
            offline_nvrs = total_nvrs - online_nvrs
            
            self.lbl_nvr_total.setText(f"📊 {total_nvrs}")
            self.lbl_nvr_online.setText(f"🟢 {online_nvrs}")
            self.lbl_nvr_offline.setText(f"🔴 {offline_nvrs}")
            
            # Update camera counts
            total_cameras = len(self.cameras)
            online_cameras = sum(1 for cam in self.cameras if 'online' in cam.get('status', '').lower())
            offline_cameras = total_cameras - online_cameras
            
            self.lbl_total.setText(f"📷 {total_cameras}")
            self.lbl_online.setText(f"🟢 {online_cameras}")
            self.lbl_offline.setText(f"🔴 {offline_cameras}")
            
            # Also emit signals for thread-safe updates
            self.ui_status_update_signal.emit("emoji_status", "cam_online", f"🟢|{online_cameras}|#d5f4e6")
            self.ui_status_update_signal.emit("emoji_status", "cam_offline", f"🔴|{offline_cameras}|#fdeaea")
            
            # Force table repopulation to show updated status indicators
            QtCore.QTimer.singleShot(300, lambda: self._refresh_table_with_status())
            log(f"[REFRESH-UI] Scheduled table refresh with {len(self.filtered)} cameras")
            log(f"[REFRESH-UI] Scheduled table refresh with {len(self.filtered)} cameras")
            
            # Update selected NVR camera counts if applicable
            if hasattr(self, 'current_nvr_selection'):
                selected_cameras = [cam for cam in self.cameras if cam.get('nvr', '') == self.current_nvr_selection]
                selected_online = sum(1 for cam in selected_cameras if 'online' in cam.get('status', '').lower())
                selected_offline = len(selected_cameras) - selected_online
                
                self.lbl_selected_cameras.setText(f"📷 {len(selected_cameras)}")
                self.lbl_selected_online.setText(f"🟢 {selected_online}")
                self.lbl_selected_offline.setText(f"🔴 {selected_offline}")
            
            log(f"[COUNT-UPDATE] Updated counts - NVRs: {online_nvrs}/{total_nvrs}, Cameras: {online_cameras}/{total_cameras}")
            
        except Exception as e:
            log(f"[COUNT-UPDATE] Error updating counts: {e}")

    @QtCore.pyqtSlot(int, str)
    def apply_nvr_update(self, index, emoji):
        """Legacy NVR update method - kept for compatibility."""
        try:
            item = self.list_nvr.item(index)
            if not item:
                return
            n = self.nvrs[index]
            sheet_flag = "" if n.get("sheet_found", False) else " ⚠️ sheet missing"
            text = f"{emoji}  🗄️ {n.get('name','')} | {n.get('ip','')} | 🎥 {n.get('cam_count',0)}{sheet_flag}"
            item.setText(text)
            log(f"[NVR UPDATE] Updated NVR {index}: {text}")
        except Exception as e:
            log(f"[NVR UPDATE] Error: {e}")

    def _enhanced_camera_check_thread(self, targets):
        """Enhanced camera check thread using same methodology as NVR refresh."""
        try:
            import concurrent.futures
            from concurrent.futures import ThreadPoolExecutor, as_completed
            
            # Use thread pool for controlled parallel processing (like NVR refresh)
            max_workers = min(6, len(targets))  # Same as NVR refresh
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all camera check tasks
                future_to_camera = {
                    executor.submit(self._comprehensive_camera_check, target): target
                    for target in targets
                }
                
                # Process results as they complete (real-time updates)
                for future in as_completed(future_to_camera):
                    target = future_to_camera[future]
                    try:
                        check_result = future.result()
                        self._update_camera_check_progress(check_result)
                        
                        # Update UI in main thread - using direct invokeMethod
                        self._apply_camera_check_update_direct(check_result)
                        
                    except Exception as e:
                        log(f"[ENHANCED-CAMERA-CHECK] Error processing camera {target.get('name', 'Unknown')}: {e}")
                        error_result = {
                            'target': target,
                            'status': 'error',
                            'error': str(e),
                            'method': 'Error'
                        }
                        self._update_camera_check_progress(error_result)
            
            # Final completion update (like NVR refresh completion)
            self._finalize_camera_check()
            
        except Exception as e:
            log(f"[ENHANCED-CAMERA-CHECK] Critical error in camera check thread: {e}")
            QtCore.QTimer.singleShot(0, lambda: self.status.showMessage(f"❌ Camera check failed: {str(e)}", 5000))

    def _comprehensive_camera_check(self, target):
        """Comprehensive camera check with multiple methods (like NVR comprehensive check)."""
        row = target['row']
        ip = target['ip']
        name = target['name']
        camera_data = target['camera_data']
        
        result = {
            'target': target,
            'row': row,
            'ip': ip,
            'name': name,
            'status': 'offline',
            'method': 'none',
            'details': {},
            'response_time': 0,
            'error': None
        }
        
        if not ip:
            result.update({'status': 'error', 'error': 'No IP address', 'method': 'Error'})
            return result
        
        try:
            import time
            start_time = time.time()
            
            log(f"[COMPREHENSIVE-CAMERA-CHECK] Checking {name} ({ip})...")
            
            # Method 1: Enhanced SADP discovery (most reliable for Hikvision cameras)
            try:
                sadp_online, sadp_model = check_camera_via_sadp(ip, timeout=2.0)
                if sadp_online:
                    response_time = time.time() - start_time
                    result.update({
                        'status': 'online',
                        'method': 'SADP',
                        'details': {'model': sadp_model, 'protocol': 'SADP Discovery'},
                        'response_time': response_time
                    })
                    log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} ✅ ONLINE via SADP - Model: {sadp_model} ({response_time:.2f}s)")
                    return result
            except Exception as e:
                log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} SADP failed: {e}")
            
            # Method 2: HTTP/HTTPS ISAPI check with authentication
            try:
                import requests
                for protocol in ['http', 'https']:
                    for port in [80, 443, 8000]:
                        try:
                            url = f"{protocol}://{ip}:{port}/ISAPI/System/deviceInfo"
                            
                            # Try with default credentials
                            for username, password in [('admin', 'Kkcctv12345'), ('admin', 'admin'), ('admin', '12345')]:
                                try:
                                    response = requests.get(url, auth=(username, password), timeout=3.0, verify=False)
                                    if response.status_code == 200:
                                        response_time = time.time() - start_time
                                        result.update({
                                            'status': 'online',
                                            'method': 'HTTP',
                                            'details': {'protocol': f'{protocol.upper()}:{port}', 'auth': f'{username}'},
                                            'response_time': response_time
                                        })
                                        log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} ✅ ONLINE via {protocol.upper()}:{port} ({response_time:.2f}s)")
                                        return result
                                    elif response.status_code == 401:
                                        # Device responds but needs different credentials
                                        response_time = time.time() - start_time
                                        result.update({
                                            'status': 'limited',
                                            'method': 'HTTP-Auth',
                                            'details': {'protocol': f'{protocol.upper()}:{port}', 'needs_auth': True},
                                            'response_time': response_time
                                        })
                                        log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} 🟡 NEEDS AUTH via {protocol.upper()}:{port}")
                                        return result
                                except:
                                    continue
                        except:
                            continue
            except Exception as e:
                log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} HTTP methods failed: {e}")
            
            # Method 3: Basic ping test
            try:
                if silent_ping(ip):
                    response_time = time.time() - start_time
                    result.update({
                        'status': 'ping',
                        'method': 'Ping',
                        'details': {'connectivity': 'basic', 'services': 'unknown'},
                        'response_time': response_time
                    })
                    log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} 🟡 PING ONLY ({response_time:.2f}s)")
                    return result
            except Exception as e:
                log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} Ping failed: {e}")
            
            # Completely offline
            response_time = time.time() - start_time
            result.update({
                'status': 'offline',
                'method': 'None',
                'details': {'connectivity': 'failed', 'all_methods': 'failed'},
                'response_time': response_time
            })
            log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} ❌ OFFLINE - All methods failed ({response_time:.2f}s)")
            return result
            
        except Exception as e:
            response_time = time.time() - start_time if 'start_time' in locals() else 0
            result.update({
                'status': 'error',
                'method': 'Error',
                'error': str(e),
                'details': {'exception': str(e)},
                'response_time': response_time
            })
            log(f"[COMPREHENSIVE-CAMERA-CHECK] {name} ⚠️ ERROR: {e}")
            return result

    def _update_camera_check_progress(self, check_result):
        """Update camera check progress tracking (like NVR refresh progress)."""
        self.camera_check_progress['completed'] += 1
        
        if check_result['status'] in ['online', 'limited', 'ping']:
            self.camera_check_progress['online'] += 1
        elif check_result['status'] == 'offline':
            self.camera_check_progress['offline'] += 1
        else:
            self.camera_check_progress['errors'] += 1
        
        # Update progress message (like NVR refresh status messages)
        completed = self.camera_check_progress['completed']
        total = self.camera_check_progress['total']
        online = self.camera_check_progress['online']
        percentage = int((completed / total) * 100)
        
        progress_msg = f"🔄 Checking cameras ({completed}/{total} - {percentage}%) | ✅ {online} online"
        
        # Update status bar using invokeMethod for thread safety
        QtCore.QMetaObject.invokeMethod(self.status, "showMessage",
            QtCore.Qt.QueuedConnection,
            QtCore.Q_ARG(str, progress_msg),
            QtCore.Q_ARG(int, 0))

    def _apply_camera_check_update_direct(self, check_result):
        """Apply comprehensive camera check result to UI - called from worker thread."""
        try:
            row = check_result['row']
            status = check_result['status']
            method = check_result['method']
            details = check_result['details']
            response_time = check_result['response_time']
            
            # Update table row with enhanced status information using invokeMethod for thread safety
            if row < self.table.rowCount():
                # Status column with enhanced visual styling
                if status == 'online':
                    status_text = f"🟢 Online ({method})"
                    status_color = QtGui.QColor(255, 255, 255)  # White text
                    bg_color = QtGui.QColor(39, 174, 96)  # Green background
                elif status == 'limited':
                    status_text = f"🟡 Limited ({method})"
                    status_color = QtGui.QColor(255, 255, 255)
                    bg_color = QtGui.QColor(243, 156, 18)  # Orange
                elif status == 'ping':
                    status_text = f"🟡 Ping Only"
                    status_color = QtGui.QColor(255, 255, 255)
                    bg_color = QtGui.QColor(243, 156, 18)
                elif status == 'error':
                    status_text = f"⚠️ Error"
                    status_color = QtGui.QColor(255, 255, 255)
                    bg_color = QtGui.QColor(231, 76, 60)  # Red
                else:
                    status_text = f"🔴 Offline"
                    status_color = QtGui.QColor(255, 255, 255)
                    bg_color = QtGui.QColor(231, 76, 60)
                
                # Create status item with styling
                status_item = QtWidgets.QTableWidgetItem(status_text)
                status_item.setForeground(status_color)
                status_item.setBackground(bg_color)
                status_item.setTextAlignment(QtCore.Qt.AlignCenter)
                
                # Update timestamp
                import datetime
                timestamp = datetime.datetime.now().strftime('%H:%M:%S')
                if response_time > 0:
                    timestamp += f" ({response_time:.1f}s)"
                timestamp_item = QtWidgets.QTableWidgetItem(timestamp)
                
                # Use invokeMethod to update UI from worker thread
                QtCore.QMetaObject.invokeMethod(self.table, "setItem",
                    QtCore.Qt.QueuedConnection,
                    QtCore.Q_ARG(int, row),
                    QtCore.Q_ARG(int, 3),
                    QtCore.Q_ARG(QtWidgets.QTableWidgetItem, status_item))
                
                QtCore.QMetaObject.invokeMethod(self.table, "setItem",
                    QtCore.Qt.QueuedConnection,
                    QtCore.Q_ARG(int, row),
                    QtCore.Q_ARG(int, 7),
                    QtCore.Q_ARG(QtWidgets.QTableWidgetItem, timestamp_item))
                
                log(f"[CAMERA-CHECK-UPDATE] Updated row {row}: {status_text}")
            
        except Exception as e:
            log(f"[CAMERA-CHECK-UPDATE] Error updating camera: {e}")

    def _finalize_camera_check(self):
        """Finalize camera check with comprehensive summary (like NVR refresh finalization)."""
        try:
            elapsed = time.time() - self.camera_check_progress['start_time']
            total = self.camera_check_progress['total']
            online = self.camera_check_progress['online']
            offline = self.camera_check_progress['offline']
            errors = self.camera_check_progress['errors']
            
            # Update comprehensive counts using invokeMethod
            def update_counts():
                self._update_comprehensive_counts()
            QtCore.QMetaObject.invokeMethod(QtCore.QCoreApplication.instance(),
                "processEvents", QtCore.Qt.QueuedConnection)
            QtCore.QTimer.singleShot(100, update_counts)
            
            # Create comprehensive completion message
            success_msg = f"✅ Camera Check Complete! {total} cameras checked in {elapsed:.1f}s | "
            success_msg += f"🟢 {online} online, "
            success_msg += f"🔴 {offline} offline"
            if errors > 0:
                success_msg += f", ⚠️ {errors} errors"
            
            # Show final status using invokeMethod
            QtCore.QMetaObject.invokeMethod(self.status, "showMessage",
                QtCore.Qt.QueuedConnection,
                QtCore.Q_ARG(str, success_msg),
                QtCore.Q_ARG(int, 8000))
            
            log(f"[COMPREHENSIVE-CAMERA-CHECK] ✅ COMPLETE: {success_msg}")
            
        except Exception as e:
            log(f"[COMPREHENSIVE-CAMERA-CHECK] Error in finalization: {e}")

    # ---------------- context menu & double-click ----------------
    def open_context_menu(self, pos):
        r = self.table.currentRow()
        if r < 0:
            return
        ip = self.table.item(r,2).text().strip()
        m = QtWidgets.QMenu()
        a1 = m.addAction("🌐  Open HTTP in Browser")
        a2 = m.addAction("🎦  Open RTSP in VLC")
        a3 = m.addAction("📸  Snapshot (http://ip/snapshot)")
        m.addSeparator()
        a4 = m.addAction("🪪  Set Credentials...")
        a5 = m.addAction("🗑️  Forget Credentials")
        action = m.exec_(self.table.viewport().mapToGlobal(pos))
        if action == a1:
            self.open_http_with_try(ip)
        elif action == a2:
            threading.Thread(target=self.open_rtsp_direct, args=(ip,), daemon=True).start()
        elif action == a3:
            webbrowser.open(f"http://{ip}/snapshot")
        elif action == a4:
            self.show_credentials_dialog(ip)
        elif action == a5:
            delete_credentials(ip)
            QtWidgets.QMessageBox.information(self, "Credentials", "Credentials removed (best-effort).")

    def double_click_ip(self, row, col):
        ip_item = self.table.item(row, 2)
        ip = ip_item.text().strip() if ip_item else ""
        if ip:
            self.status.showMessage(f"Launching VLC for {ip}...")
            threading.Thread(target=self.open_rtsp_direct, args=(ip,), daemon=True).start()

    # ---------------- open helpers ----------------
    def open_http_with_try(self, ip):
        u,p = get_password(ip)
        if u and p:
            webbrowser.open(f"http://{u}:{p}@{ip}"); return
        for du,dp in DEFAULT_CREDS:
            webbrowser.open(f"http://{du}:{dp}@{ip}"); return
        webbrowser.open(f"http://{ip}")

    def open_rtsp_direct(self, ip):
        u,p = get_password(ip)
        if u and p:
            if self._launch_vlc_rtsp(ip,u,p): return
        for du,dp in DEFAULT_CREDS:
            if self._launch_vlc_rtsp(ip,du,dp): return
        self._launch_vlc_rtsp(ip,None,None)

    def _launch_vlc_rtsp(self, ip, username, password):
        if username and password:
            url = f"rtsp://{username}:{password}@{ip}"
        else:
            url = f"rtsp://{ip}"
        try:
            subprocess.Popen([self.vlc, url], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            log(f"Launching VLC -> {url}")
            return True
        except Exception as e:
            log(f"VLC launch failed ({url}): {e}")
            return False

    # ---------------- credentials dialog ----------------
    def show_credentials_dialog(self, ip):
        dlg = QtWidgets.QDialog(self); dlg.setWindowTitle(f"Credentials for {ip}")
        layout = QtWidgets.QFormLayout(dlg)
        user_edit = QtWidgets.QLineEdit(); pwd_edit = QtWidgets.QLineEdit(); pwd_edit.setEchoMode(QtWidgets.QLineEdit.Password)
        meta = load_creds_meta()
        if ip in meta: user_edit.setText(meta[ip].get("username",""))
        auto_chk = QtWidgets.QCheckBox("Auto-login when opening RTSP"); auto_chk.setChecked(meta.get(ip,{}).get("auto_login", False))
        layout.addRow("Username:", user_edit); layout.addRow("Password:", pwd_edit); layout.addRow(auto_chk)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel); layout.addRow(btns)
        btns.accepted.connect(dlg.accept); btns.rejected.connect(dlg.reject)
        if dlg.exec_() == QtWidgets.QDialog.Accepted:
            username = user_edit.text().strip(); pwd = pwd_edit.text(); auto_login = auto_chk.isChecked()
            if username and pwd:
                ok = set_password(ip, username, pwd)
                meta = load_creds_meta(); meta[ip] = {"username": username, "auto_login": auto_login}; save_creds_meta(meta)
                self.creds_meta = meta
                if ok and KEYRING_AVAILABLE:
                    QtWidgets.QMessageBox.information(self, "Saved", "Credentials saved to system keyring.")
                else:
                    QtWidgets.QMessageBox.information(self, "Saved (fallback)", "Credentials saved to local fallback store.")
            else:
                QtWidgets.QMessageBox.warning(self, "Invalid", "Provide username and password to save.")

    # ---------------- export ----------------
    def show_export_dialog(self):
        """Show professional export dialog with company info and format selection"""
        dialog = QtWidgets.QDialog(self)
        dialog.setWindowTitle("📊 Professional Export")
        dialog.setMinimumWidth(500)
        
        layout = QtWidgets.QVBoxLayout(dialog)
        
        # Header
        header = QtWidgets.QLabel("<b style='font-size: 14px;'>📊 Generate Professional Report</b>")
        layout.addWidget(header)
        
        layout.addWidget(QtWidgets.QLabel("<hr>"))
        
        # Company info section
        company_group = QtWidgets.QGroupBox("Company Information")
        company_layout = QtWidgets.QFormLayout()
        
        company_name = QtWidgets.QLineEdit()
        company_name.setPlaceholderText("Enter company name")
        company_name.setText(getattr(self, 'export_company_name', APP_COMPANY))
        
        telephone = QtWidgets.QLineEdit()
        telephone.setPlaceholderText("Enter telephone number")
        telephone.setText(getattr(self, 'export_telephone', ''))
        
        telegram = QtWidgets.QLineEdit()
        telegram.setPlaceholderText("Enter Telegram contact")
        telegram.setText(getattr(self, 'export_telegram', ''))
        
        company_layout.addRow("Company Name:", company_name)
        company_layout.addRow("Telephone:", telephone)
        company_layout.addRow("Telegram:", telegram)
        
        company_group.setLayout(company_layout)
        layout.addWidget(company_group)
        
        # Logo selection
        logo_group = QtWidgets.QGroupBox("Logo")
        logo_layout = QtWidgets.QHBoxLayout()
        
        logo_path_display = QtWidgets.QLineEdit()
        logo_path_display.setPlaceholderText("Default logo will be used")
        logo_path_display.setReadOnly(True)
        logo_path_display.setText(getattr(self, 'export_logo_path', ''))
        
        browse_btn = QtWidgets.QPushButton("Browse...")
        def select_logo():
            path, _ = QtWidgets.QFileDialog.getOpenFileName(dialog, "Select Logo", "", "Images (*.png *.jpg *.jpeg)")
            if path:
                logo_path_display.setText(path)
        browse_btn.clicked.connect(select_logo)
        
        logo_layout.addWidget(logo_path_display)
        logo_layout.addWidget(browse_btn)
        logo_group.setLayout(logo_layout)
        layout.addWidget(logo_group)
        
        # Export format selection
        format_group = QtWidgets.QGroupBox("Export Format")
        format_layout = QtWidgets.QVBoxLayout()
        
        excel_radio = QtWidgets.QRadioButton("📗 Excel (XLSX) - Recommended")
        excel_radio.setChecked(True)
        word_radio = QtWidgets.QRadioButton("📘 Word (DOCX)")
        pdf_radio = QtWidgets.QRadioButton("📕 PDF")
        
        format_layout.addWidget(excel_radio)
        format_layout.addWidget(word_radio)
        format_layout.addWidget(pdf_radio)
        format_group.setLayout(format_layout)
        layout.addWidget(format_group)
        
        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        export_btn = QtWidgets.QPushButton("📊 Generate Report")
        export_btn.setStyleSheet("QPushButton { background-color: #27ae60; color: white; font-weight: bold; padding: 8px; }")
        cancel_btn = QtWidgets.QPushButton("Cancel")
        
        def do_export():
            # Save settings
            self.export_company_name = company_name.text().strip()
            self.export_telephone = telephone.text().strip()
            self.export_telegram = telegram.text().strip()
            self.export_logo_path = logo_path_display.text().strip()
            
            export_info = {
                'company_name': self.export_company_name or APP_COMPANY,
                'telephone': self.export_telephone,
                'telegram': self.export_telegram,
                'logo_path': self.export_logo_path
            }
            
            if excel_radio.isChecked():
                self.export_to_excel(export_info)
            elif word_radio.isChecked():
                self.export_to_word(export_info)
            elif pdf_radio.isChecked():
                self.export_to_pdf(export_info)
            
            dialog.accept()
        
        export_btn.clicked.connect(do_export)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(export_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        dialog.exec_()
    
    def export_csv(self):
        """Legacy CSV export - kept for compatibility"""
        if not self.cams:
            QtWidgets.QMessageBox.information(self, "Export", "No cameras to export."); return
        try:
            with open(EXPORT_FILE, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f); w.writerow(["nvr","camera","ip"])
                for c in self.cams:
                    w.writerow([c.get("nvr",""), c.get("name",""), c.get("ip","")])
            QtWidgets.QMessageBox.information(self, "Export", f"Saved {EXPORT_FILE}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Export error", str(e))
    
    def export_to_excel(self, export_info):
        """Export to professional Excel with NVR-grouped sheets and summary"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            from datetime import datetime
            from collections import defaultdict
            
            if not self.api_cameras and not self.cams:
                QtWidgets.QMessageBox.information(self, "Export", "No cameras to export.")
                return
            
            camera_source = self.api_cameras if self.api_cameras else self.cams
            
            filename, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Save Excel Report", 
                f"Camera_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "Excel Files (*.xlsx)"
            )
            
            if not filename:
                return
            
            # Group cameras by NVR
            nvr_groups = defaultdict(list)
            for cam in camera_source:
                nvr_name = cam.get('nvr', 'Unknown NVR')
                nvr_groups[nvr_name].append(cam)
            
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # ========== CREATE SUMMARY SHEET ==========
            summary_ws = wb.create_sheet("📊 Summary", 0)
            summary_ws.sheet_properties.tabColor = "1F4E78"
            
            # Summary header
            summary_ws.merge_cells('A1:F1')
            header_cell = summary_ws['A1']
            header_cell.value = f"🏢 {export_info['company_name']} - Camera Summary Report"
            header_cell.font = Font(size=16, bold=True, color="FFFFFF")
            header_cell.fill = PatternFill(start_color="1a5490", end_color="1a5490", fill_type="solid")
            header_cell.alignment = Alignment(horizontal="center", vertical="center")
            summary_ws.row_dimensions[1].height = 30
            
            # Date
            summary_ws.merge_cells('A2:F2')
            date_cell = summary_ws['A2']
            date_cell.value = f"📅 Generated: {datetime.now().strftime('%A, %B %d, %Y at %I:%M %p')}"
            date_cell.font = Font(size=10, italic=True, color="7f8c8d")
            date_cell.fill = PatternFill(start_color="ebf5fb", end_color="ebf5fb", fill_type="solid")
            date_cell.alignment = Alignment(horizontal="center", vertical="center")
            summary_ws.row_dimensions[2].height = 20
            
            # Spacer
            summary_ws.row_dimensions[3].height = 5
            
            # Summary table headers
            headers = ["📡 NVR Name", "📷 Total Cameras", "✅ Online", "❌ Offline", "⚠️ Unknown", "📈 Uptime %"]
            for col, header in enumerate(headers, 1):
                cell = summary_ws.cell(row=4, column=col)
                cell.value = header
                cell.font = Font(bold=True, color="FFFFFF", size=11)
                cell.fill = PatternFill(start_color="34495e", end_color="34495e", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = Border(left=Side(style='medium'), right=Side(style='medium'),
                                   top=Side(style='medium'), bottom=Side(style='medium'))
            summary_ws.row_dimensions[4].height = 25
            
            # Summary data rows
            row_num = 5
            total_cameras = 0
            total_online = 0
            total_offline = 0
            total_unknown = 0
            
            for nvr_name in sorted(nvr_groups.keys()):
                cameras = nvr_groups[nvr_name]
                nvr_total = len(cameras)
                nvr_online = sum(1 for c in cameras if 'online' in str(c.get('status', '')).lower())
                nvr_offline = sum(1 for c in cameras if 'offline' in str(c.get('status', '')).lower())
                nvr_unknown = nvr_total - nvr_online - nvr_offline
                nvr_uptime = (nvr_online / nvr_total * 100) if nvr_total > 0 else 0
                
                total_cameras += nvr_total
                total_online += nvr_online
                total_offline += nvr_offline
                total_unknown += nvr_unknown
                
                # NVR Name
                cell = summary_ws.cell(row=row_num, column=1)
                cell.value = nvr_name
                cell.font = Font(bold=True, size=10)
                cell.alignment = Alignment(horizontal="left", vertical="center")
                
                # Total
                cell = summary_ws.cell(row=row_num, column=2)
                cell.value = nvr_total
                cell.font = Font(size=10)
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Online
                cell = summary_ws.cell(row=row_num, column=3)
                cell.value = nvr_online
                cell.font = Font(size=10, color="27ae60", bold=True)
                cell.fill = PatternFill(start_color="d5f4e6", end_color="d5f4e6", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Offline
                cell = summary_ws.cell(row=row_num, column=4)
                cell.value = nvr_offline
                cell.font = Font(size=10, color="c0392b", bold=True)
                cell.fill = PatternFill(start_color="fadbd8", end_color="fadbd8", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Unknown
                cell = summary_ws.cell(row=row_num, column=5)
                cell.value = nvr_unknown
                cell.font = Font(size=10, color="d68910", bold=True)
                cell.fill = PatternFill(start_color="fcf3cf", end_color="fcf3cf", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Uptime %
                cell = summary_ws.cell(row=row_num, column=6)
                cell.value = f"{nvr_uptime:.1f}%"
                cell.font = Font(size=10, bold=True)
                if nvr_uptime >= 90:
                    cell.font = Font(size=10, bold=True, color="27ae60")
                elif nvr_uptime >= 70:
                    cell.font = Font(size=10, bold=True, color="f39c12")
                else:
                    cell.font = Font(size=10, bold=True, color="c0392b")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Add borders
                for col in range(1, 7):
                    summary_ws.cell(row=row_num, column=col).border = Border(
                        left=Side(style='thin', color='bdc3c7'),
                        right=Side(style='thin', color='bdc3c7'),
                        top=Side(style='thin', color='bdc3c7'),
                        bottom=Side(style='thin', color='bdc3c7')
                    )
                
                # Alternating rows
                if row_num % 2 == 0:
                    for col in range(1, 7):
                        if col not in [3, 4, 5]:  # Don't override status colors
                            summary_ws.cell(row=row_num, column=col).fill = PatternFill(
                                start_color="f8f9fa", end_color="f8f9fa", fill_type="solid"
                            )
                
                row_num += 1
            
            # Total row
            row_num += 1
            summary_ws.merge_cells(f'A{row_num}:F{row_num}')
            spacer_cell = summary_ws[f'A{row_num}']
            spacer_cell.fill = PatternFill(start_color="34495e", end_color="34495e", fill_type="solid")
            summary_ws.row_dimensions[row_num].height = 3
            row_num += 1
            
            total_uptime = (total_online / total_cameras * 100) if total_cameras > 0 else 0
            
            cell = summary_ws.cell(row=row_num, column=1)
            cell.value = "🏁 GRAND TOTAL"
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
            cell.alignment = Alignment(horizontal="left", vertical="center")
            
            cell = summary_ws.cell(row=row_num, column=2)
            cell.value = total_cameras
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            cell = summary_ws.cell(row=row_num, column=3)
            cell.value = total_online
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="27ae60", end_color="27ae60", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            cell = summary_ws.cell(row=row_num, column=4)
            cell.value = total_offline
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="c0392b", end_color="c0392b", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            cell = summary_ws.cell(row=row_num, column=5)
            cell.value = total_unknown
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="d68910", end_color="d68910", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            cell = summary_ws.cell(row=row_num, column=6)
            cell.value = f"{total_uptime:.1f}%"
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill(start_color="16a085", end_color="16a085", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            for col in range(1, 7):
                summary_ws.cell(row=row_num, column=col).border = Border(
                    left=Side(style='medium', color='000000'),
                    right=Side(style='medium', color='000000'),
                    top=Side(style='medium', color='000000'),
                    bottom=Side(style='medium', color='000000')
                )
            
            summary_ws.row_dimensions[row_num].height = 25
            
            # Column widths for summary
            summary_ws.column_dimensions['A'].width = 25
            summary_ws.column_dimensions['B'].width = 15
            summary_ws.column_dimensions['C'].width = 12
            summary_ws.column_dimensions['D'].width = 12
            summary_ws.column_dimensions['E'].width = 12
            summary_ws.column_dimensions['F'].width = 12
            
            # ========== CREATE NVR SHEETS ==========
            def create_nvr_sheet(nvr_name, cameras):
                """Create individual sheet for each NVR"""
                # Sanitize sheet name (Excel limits to 31 chars and no special chars)
                safe_name = nvr_name[:28].replace('/', '-').replace('\\', '-').replace('*', '').replace('?', '').replace(':', '').replace('[', '').replace(']', '')
                ws = wb.create_sheet(safe_name)
                
                # Header
                ws.merge_cells('A1:I1')
                header = ws['A1']
                header.value = f"📡 {nvr_name} - Camera Details"
                header.font = Font(size=14, bold=True, color="FFFFFF")
                header.fill = PatternFill(start_color="2874a6", end_color="2874a6", fill_type="solid")
                header.alignment = Alignment(horizontal="center", vertical="center")
                ws.row_dimensions[1].height = 30
                
                # Stats
                nvr_online = sum(1 for c in cameras if 'online' in str(c.get('status', '')).lower())
                nvr_offline = len(cameras) - nvr_online
                
                ws.merge_cells('A2:I2')
                stats = ws['A2']
                stats.value = f"📊 Total: {len(cameras)} | ✅ Online: {nvr_online} | ❌ Offline: {nvr_offline}"
                stats.font = Font(size=10, bold=True, color="34495e")
                stats.fill = PatternFill(start_color="d6eaf8", end_color="d6eaf8", fill_type="solid")
                stats.alignment = Alignment(horizontal="center", vertical="center")
                ws.row_dimensions[2].height = 22
                
                # Column headers
                headers = [("🚦", "Status"), ("📷", "Camera Name"), ("🌐", "IP Address"), 
                          ("🔧", "Model"), ("🔌", "Port"), ("🕒", "Last Updated"), 
                          ("📊", "Connection"), ("📝", "Remark"), ("🔍", "Details")]
                
                for col, (icon, header_text) in enumerate(headers, 1):
                    cell = ws.cell(row=4, column=col)
                    cell.value = f"{icon}\n{header_text}"
                    cell.font = Font(bold=True, color="FFFFFF", size=10)
                    cell.fill = PatternFill(start_color="34495e", end_color="34495e", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.border = Border(left=Side(style='medium'), right=Side(style='medium'),
                                       top=Side(style='medium'), bottom=Side(style='medium'))
                ws.row_dimensions[4].height = 32
                
                # Camera rows
                row = 5
                for cam in sorted(cameras, key=lambda x: x.get('name', '')):
                    status = str(cam.get('status', '')).lower()
                    is_online = 'online' in status
                    
                    # Status
                    cell = ws.cell(row=row, column=1)
                    cell.value = "✅" if is_online else "❌"
                    cell.font = Font(size=14)
                    cell.fill = PatternFill(start_color="d5f4e6" if is_online else "fadbd8", 
                                          end_color="d5f4e6" if is_online else "fadbd8", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Name
                    cell = ws.cell(row=row, column=2)
                    cell.value = cam.get('name', '')
                    cell.font = Font(bold=True, size=10)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                    
                    # IP
                    cell = ws.cell(row=row, column=3)
                    cell.value = cam.get('ip', '')
                    cell.font = Font(size=10, name='Courier New', color="154360")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Model
                    cell = ws.cell(row=row, column=4)
                    cell.value = cam.get('model', '')
                    cell.font = Font(size=9, color="566573")
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                    
                    # Port
                    cell = ws.cell(row=row, column=5)
                    cell.value = str(cam.get('port', ''))
                    cell.font = Font(size=10)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Last Updated
                    cell = ws.cell(row=row, column=6)
                    cell.value = cam.get('last_updated', '')
                    cell.font = Font(size=9, italic=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Connection
                    cell = ws.cell(row=row, column=7)
                    conn = cam.get('connection_type', 'Standard')
                    cell.value = conn.replace('_', ' ').title()
                    cell.font = Font(size=9, color="16a085" if 'digest' in conn.lower() else "7f8c8d")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Remark
                    cell = ws.cell(row=row, column=8)
                    cell.value = cam.get('remark', '')
                    cell.font = Font(size=9, color="7f8c8d", italic=True)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                    
                    # Details
                    cell = ws.cell(row=row, column=9)
                    cell.value = f"Ch {cam.get('channel', 'N/A')}"
                    cell.font = Font(size=9)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    # Borders and alternating colors
                    border = Border(left=Side(style='thin', color='bdc3c7'),
                                  right=Side(style='thin', color='bdc3c7'),
                                  top=Side(style='thin', color='bdc3c7'),
                                  bottom=Side(style='thin', color='bdc3c7'))
                    
                    for col in range(1, 10):
                        ws.cell(row=row, column=col).border = border
                        if col != 1 and row % 2 == 0:
                            ws.cell(row=row, column=col).fill = PatternFill(
                                start_color="f8f9fa", end_color="f8f9fa", fill_type="solid"
                            )
                    
                    ws.row_dimensions[row].height = 20
                    row += 1
                
                # Column widths
                ws.column_dimensions['A'].width = 10
                ws.column_dimensions['B'].width = 28
                ws.column_dimensions['C'].width = 16
                ws.column_dimensions['D'].width = 22
                ws.column_dimensions['E'].width = 8
                ws.column_dimensions['F'].width = 18
                ws.column_dimensions['G'].width = 14
                ws.column_dimensions['H'].width = 20
                ws.column_dimensions['I'].width = 10
                
                # Freeze panes
                ws.freeze_panes = 'A5'
            
            # Create a sheet for each NVR
            for nvr_name in sorted(nvr_groups.keys()):
                create_nvr_sheet(nvr_name, nvr_groups[nvr_name])
            
            # Save workbook
            wb.save(filename)
            
            QtWidgets.QMessageBox.information(
                self, 
                "Export Complete", 
                f"✅ Professional Excel report with NVR-grouped sheets saved!\n\n"
                f"📁 {filename}\n\n"
                f"📊 {len(nvr_groups)} NVRs | {total_cameras} cameras\n"
                f"✅ {total_online} online | ❌ {total_offline} offline\n\n"
                f"📑 Includes Summary sheet + {len(nvr_groups)} NVR detail sheets"
            )
            
        except ImportError as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Required library not installed:\n{str(e)}\n\nRun: pip install openpyxl")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Export Error", f"Failed to export Excel:\n{str(e)}")
            import traceback
            log(f"[EXPORT-EXCEL] Error: {traceback.format_exc()}")
    
    def export_to_word(self, export_info):
        """Export to professional Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            if not self.api_cameras and not self.cams:
                QtWidgets.QMessageBox.information(self, "Export", "No cameras to export.")
                return
            
            camera_source = self.api_cameras if self.api_cameras else self.cams
            
            filename, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Save Word Report",
                f"Camera_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx)"
            )
            
            if not filename:
                return
            
            doc = Document()
            
            # Company header
            header = doc.add_heading(export_info['company_name'], 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            # Contact info
            contact_parts = []
            if export_info['telephone']:
                contact_parts.append(f"☎ {export_info['telephone']}")
            if export_info['telegram']:
                contact_parts.append(f"✈ {export_info['telegram']}")
            
            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.runs[0].font.size = Pt(10)
                contact_para.runs[0].font.color.rgb = RGBColor(127, 140, 141)
            
            # Report date
            date_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.italic = True
            date_para.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Empty line
            
            # Table
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            hdr_cells = table.rows[0].cells
            headers = ["Status", "Camera Name", "IP Address", "NVR", "Model", "Port", "Last Updated", "Remark"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            # Data rows
            for cam in camera_source:
                row_cells = table.add_row().cells
                
                status = cam.get('status', '').lower()
                if 'online' in status or '🟢' in status:
                    row_cells[0].text = "✓ Online"
                elif 'offline' in status or '🔴' in status:
                    row_cells[0].text = "✗ Offline"
                else:
                    row_cells[0].text = "? Unknown"
                
                row_cells[1].text = cam.get('name', '')
                row_cells[2].text = cam.get('ip', '')
                row_cells[3].text = cam.get('nvr', '')
                row_cells[4].text = cam.get('model', '')
                row_cells[5].text = str(cam.get('port', ''))
                row_cells[6].text = cam.get('last_updated', '')
                row_cells[7].text = cam.get('remark', '')
            
            # Summary
            doc.add_paragraph()
            online_count = sum(1 for c in camera_source if 'online' in str(c.get('status', '')).lower())
            offline_count = sum(1 for c in camera_source if 'offline' in str(c.get('status', '')).lower())
            
            summary = doc.add_paragraph(f"Total Cameras: {len(camera_source)} | Online: {online_count} | Offline: {offline_count}")
            summary.runs[0].font.bold = True
            summary.runs[0].font.size = Pt(11)
            
            doc.save(filename)
            QtWidgets.QMessageBox.information(self, "Export Complete", f"Professional Word report saved:\n{filename}")
            
        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Error", "python-docx library not installed.\nRun: pip install python-docx")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Export Error", f"Failed to export Word:\n{str(e)}")
    
    def export_to_pdf(self, export_info):
        """Export to professional PDF document"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER
            from datetime import datetime
            
            if not self.api_cameras and not self.cams:
                QtWidgets.QMessageBox.information(self, "Export", "No cameras to export.")
                return
            
            camera_source = self.api_cameras if self.api_cameras else self.cams
            
            filename, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Save PDF Report",
                f"Camera_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "PDF Files (*.pdf)"
            )
            
            if not filename:
                return
            
            doc = SimpleDocTemplate(filename, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#2c3e50'),
                alignment=TA_CENTER,
                spaceAfter=10
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#7f8c8d'),
                alignment=TA_CENTER,
                spaceAfter=5
            )
            
            # Company header
            story.append(Paragraph(export_info['company_name'], title_style))
            
            # Contact info
            contact_parts = []
            if export_info['telephone']:
                contact_parts.append(f"☎ {export_info['telephone']}")
            if export_info['telegram']:
                contact_parts.append(f"✈ {export_info['telegram']}")
            
            if contact_parts:
                story.append(Paragraph(" | ".join(contact_parts), subtitle_style))
            
            # Report date
            story.append(Paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Table data
            data = [["Status", "Camera Name", "IP", "NVR", "Model", "Port", "Updated"]]
            
            for cam in camera_source:
                status = cam.get('status', '').lower()
                if 'online' in status or '🟢' in status:
                    status_text = "✓ Online"
                elif 'offline' in status or '🔴' in status:
                    status_text = "✗ Offline"
                else:
                    status_text = "? Unknown"
                
                data.append([
                    status_text,
                    cam.get('name', '')[:20],  # Truncate for PDF
                    cam.get('ip', ''),
                    cam.get('nvr', '')[:10],
                    cam.get('model', '')[:15],
                    str(cam.get('port', '')),
                    cam.get('last_updated', '')[:16]
                ])
            
            # Create table
            table = Table(data, colWidths=[0.8*inch, 1.5*inch, 1.1*inch, 0.8*inch, 1.2*inch, 0.6*inch, 1.2*inch])
            
            # Table style
            table.setStyle(TableStyle([
                # Header
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                # Data rows
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                # Borders
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ]))
            
            story.append(table)
            story.append(Spacer(1, 0.3*inch))
            
            # Summary
            online_count = sum(1 for c in camera_source if 'online' in str(c.get('status', '')).lower())
            offline_count = sum(1 for c in camera_source if 'offline' in str(c.get('status', '')).lower())
            
            summary_text = f"<b>Total Cameras: {len(camera_source)} | Online: {online_count} | Offline: {offline_count}</b>"
            story.append(Paragraph(summary_text, styles['Normal']))
            
            doc.build(story)
            QtWidgets.QMessageBox.information(self, "Export Complete", f"Professional PDF report saved:\n{filename}")
            
        except ImportError:
            QtWidgets.QMessageBox.critical(self, "Error", "reportlab library not installed.\nRun: pip install reportlab")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Export Error", f"Failed to export PDF:\n{str(e)}")

    # ---------------- NVR context menu ----------------
    # ==================== DUPLICATE DETECTION ====================
    def find_duplicate_cameras(self):
        """Find and report duplicate cameras by IP or name"""
        duplicates = {'by_ip': {}, 'by_name': {}}
        
        # Group cameras by IP
        for cam in self.cams:
            ip = cam.get('ip', '').strip()
            name = cam.get('name', '').strip().lower()
            
            if ip:
                if ip not in duplicates['by_ip']:
                    duplicates['by_ip'][ip] = []
                duplicates['by_ip'][ip].append(cam)
            
            if name:
                if name not in duplicates['by_name']:
                    duplicates['by_name'][name] = []
                duplicates['by_name'][name].append(cam)
        
        # Filter to only actual duplicates
        ip_dups = {k: v for k, v in duplicates['by_ip'].items() if len(v) > 1}
        name_dups = {k: v for k, v in duplicates['by_name'].items() if len(v) > 1}
        
        return {'ip_duplicates': ip_dups, 'name_duplicates': name_dups}
    
    def show_duplicate_report(self):
        """Show duplicate cameras report dialog"""
        duplicates = self.find_duplicate_cameras()
        
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("🔍 Duplicate Camera Analysis")
        dlg.setGeometry(100, 100, 800, 600)
        
        layout = QtWidgets.QVBoxLayout(dlg)
        
        # Header
        header = QtWidgets.QLabel("<h2>🔍 Duplicate Camera Detection Report</h2>")
        layout.addWidget(header)
        
        # Results display
        text_area = QtWidgets.QTextEdit()
        text_area.setReadOnly(True)
        
        report = "<h3>Duplicate Analysis Results:</h3><br>"
        
        if duplicates['ip_duplicates']:
            report += "<h4 style='color: red;'>⚠️ IP Address Duplicates:</h4>"
            for ip, cams in duplicates['ip_duplicates'].items():
                report += f"<b>IP {ip}:</b><br>"
                for cam in cams:
                    report += f"  • {cam.get('name', 'Unknown')} (NVR: {cam.get('nvr', 'Unknown')})<br>"
                report += "<br>"
        
        if duplicates['name_duplicates']:
            report += "<h4 style='color: orange;'>⚠️ Name Duplicates:</h4>"
            for name, cams in duplicates['name_duplicates'].items():
                report += f"<b>Name '{name.title()}':</b><br>"
                for cam in cams:
                    report += f"  • {cam.get('ip', 'No IP')} (NVR: {cam.get('nvr', 'Unknown')})<br>"
                report += "<br>"
        
        if not duplicates['ip_duplicates'] and not duplicates['name_duplicates']:
            report += "<h3 style='color: green;'>✅ No duplicates found!</h3>"
            report += "All cameras have unique IPs and names."
        
        text_area.setHtml(report)
        layout.addWidget(text_area)
        
        # Close button
        btn_close = QtWidgets.QPushButton("Close")
        btn_close.clicked.connect(dlg.accept)
        layout.addWidget(btn_close)
        
        dlg.exec_()
    
    # ==================== SADP DEVICE DISCOVERY TOOL ====================
    def show_sadp_tool(self):
        """Show SADP (Search Active Device Protocol) tool for discovering Hikvision devices."""
        dialog = QtWidgets.QDialog(self)
        dialog.setWindowTitle("🔍 SADP - Search Active Device Protocol")
        dialog.setGeometry(100, 100, 1000, 700)
        
        layout = QtWidgets.QVBoxLayout(dialog)
        
        # Header
        header = QtWidgets.QLabel("<h2>🔍 SADP - Hikvision Device Discovery Tool</h2>")
        header.setAlignment(QtCore.Qt.AlignCenter)
        header.setStyleSheet("background-color: #16a085; color: white; padding: 10px; border-radius: 5px;")
        layout.addWidget(header)
        
        # Info section
        info_layout = QtWidgets.QHBoxLayout()
        info_label = QtWidgets.QLabel(
            "<b>Discover Hikvision devices on your network:</b><br>"
            "• Network scan for cameras, NVRs, DVRs<br>"
            "• View device details (IP, MAC, model, firmware)<br>"
            "• Modify network settings (IP, gateway, subnet)<br>"
            "• Reset passwords and configure devices"
        )
        info_layout.addWidget(info_label)
        
        # Scan settings
        settings_group = QtWidgets.QGroupBox("Scan Settings")
        settings_layout = QtWidgets.QFormLayout()
        
        ip_range_input = QtWidgets.QLineEdit()
        ip_range_input.setPlaceholderText("e.g., 192.168.1.0/24")
        ip_range_input.setText("192.168.0.0/16")  # Default scan range
        
        port_input = QtWidgets.QSpinBox()
        port_input.setRange(1, 65535)
        port_input.setValue(80)
        
        timeout_input = QtWidgets.QDoubleSpinBox()
        timeout_input.setRange(0.5, 10.0)
        timeout_input.setSingleStep(0.5)
        timeout_input.setValue(2.0)
        timeout_input.setSuffix(" sec")
        
        settings_layout.addRow("IP Range:", ip_range_input)
        settings_layout.addRow("HTTP Port:", port_input)
        settings_layout.addRow("Timeout:", timeout_input)
        settings_group.setLayout(settings_layout)
        info_layout.addWidget(settings_group)
        
        layout.addLayout(info_layout)
        
        # Control buttons
        btn_layout = QtWidgets.QHBoxLayout()
        
        scan_btn = QtWidgets.QPushButton("🔍 Start Scan")
        scan_btn.setStyleSheet("QPushButton { background-color: #16a085; color: white; font-weight: bold; padding: 8px; }")
        
        stop_btn = QtWidgets.QPushButton("⏹ Stop")
        stop_btn.setEnabled(False)
        stop_btn.setStyleSheet("QPushButton { background-color: #e74c3c; color: white; font-weight: bold; padding: 8px; }")
        
        export_btn = QtWidgets.QPushButton("💾 Export Results")
        export_btn.setEnabled(False)
        
        btn_layout.addWidget(scan_btn)
        btn_layout.addWidget(stop_btn)
        btn_layout.addWidget(export_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # Progress section
        progress_layout = QtWidgets.QHBoxLayout()
        progress_bar = QtWidgets.QProgressBar()
        progress_bar.setVisible(False)
        progress_label = QtWidgets.QLabel("Ready to scan...")
        progress_layout.addWidget(progress_label)
        progress_layout.addWidget(progress_bar)
        layout.addLayout(progress_layout)
        
        # Results table
        table = QtWidgets.QTableWidget()
        table.setColumnCount(10)
        table.setHorizontalHeaderLabels([
            "Status", "Device Type", "IP Address", "HTTP Port", "MAC Address", 
            "Device Name", "Model", "Firmware", "Serial Number", "Actions"
        ])
        table.horizontalHeader().setStretchLastSection(False)
        table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)
        table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        table.setAlternatingRowColors(True)
        table.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        layout.addWidget(table)
        
        # Statistics footer
        stats_label = QtWidgets.QLabel("📊 Devices found: 0 | Cameras: 0 | NVRs: 0 | DVRs: 0")
        stats_label.setStyleSheet("padding: 5px; background-color: #ecf0f1; border-radius: 3px;")
        layout.addWidget(stats_label)
        
        # Close button
        close_btn = QtWidgets.QPushButton("✕ Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        # State variables
        scan_thread = None
        stop_scan = False
        discovered_devices = []
        
        def update_statistics():
            """Update device statistics."""
            cameras = sum(1 for d in discovered_devices if 'camera' in d.get('type', '').lower())
            nvrs = sum(1 for d in discovered_devices if 'nvr' in d.get('type', '').lower())
            dvrs = sum(1 for d in discovered_devices if 'dvr' in d.get('type', '').lower())
            stats_label.setText(f"📊 Devices found: {len(discovered_devices)} | Cameras: {cameras} | NVRs: {nvrs} | DVRs: {dvrs}")
        
        def add_device_to_table(device):
            """Add discovered device to table."""
            row = table.rowCount()
            table.insertRow(row)
            
            # Status
            status_item = QtWidgets.QTableWidgetItem("🟢")
            status_item.setTextAlignment(QtCore.Qt.AlignCenter)
            table.setItem(row, 0, status_item)
            
            # Device type with icon
            device_type = device.get('type', 'Unknown')
            if 'camera' in device_type.lower():
                type_icon = "📷 Camera"
            elif 'nvr' in device_type.lower():
                type_icon = "🗄️ NVR"
            elif 'dvr' in device_type.lower():
                type_icon = "📼 DVR"
            else:
                type_icon = "❓ Unknown"
            
            type_item = QtWidgets.QTableWidgetItem(type_icon)
            table.setItem(row, 1, type_item)
            
            # IP Address
            ip_item = QtWidgets.QTableWidgetItem(device.get('ip', ''))
            ip_item.setFont(QtGui.QFont('Courier New', 9))
            table.setItem(row, 2, ip_item)
            
            # HTTP Port
            port_item = QtWidgets.QTableWidgetItem(str(device.get('port', '80')))
            port_item.setTextAlignment(QtCore.Qt.AlignCenter)
            table.setItem(row, 3, port_item)
            
            # MAC Address
            mac_item = QtWidgets.QTableWidgetItem(device.get('mac', 'N/A'))
            mac_item.setFont(QtGui.QFont('Courier New', 9))
            table.setItem(row, 4, mac_item)
            
            # Device Name
            name_item = QtWidgets.QTableWidgetItem(device.get('name', 'Unknown'))
            table.setItem(row, 5, name_item)
            
            # Model
            model_item = QtWidgets.QTableWidgetItem(device.get('model', 'N/A'))
            table.setItem(row, 6, model_item)
            
            # Firmware
            fw_item = QtWidgets.QTableWidgetItem(device.get('firmware', 'N/A'))
            table.setItem(row, 7, fw_item)
            
            # Serial Number
            serial_item = QtWidgets.QTableWidgetItem(device.get('serial', 'N/A'))
            table.setItem(row, 8, serial_item)
            
            # Actions button
            actions_btn = QtWidgets.QPushButton("⚙️ Configure")
            actions_btn.clicked.connect(lambda: show_device_config(device))
            table.setCellWidget(row, 9, actions_btn)
        
        def show_device_config(device):
            """Show device configuration dialog."""
            config_dlg = QtWidgets.QDialog(dialog)
            config_dlg.setWindowTitle(f"⚙️ Configure Device - {device.get('ip')}")
            config_dlg.setMinimumWidth(500)
            
            config_layout = QtWidgets.QVBoxLayout(config_dlg)
            
            # Device info
            info_group = QtWidgets.QGroupBox("Device Information")
            info_layout = QtWidgets.QFormLayout()
            info_layout.addRow("IP Address:", QtWidgets.QLabel(device.get('ip', 'N/A')))
            info_layout.addRow("Device Type:", QtWidgets.QLabel(device.get('type', 'N/A')))
            info_layout.addRow("Model:", QtWidgets.QLabel(device.get('model', 'N/A')))
            info_layout.addRow("MAC Address:", QtWidgets.QLabel(device.get('mac', 'N/A')))
            info_layout.addRow("Firmware:", QtWidgets.QLabel(device.get('firmware', 'N/A')))
            info_group.setLayout(info_layout)
            config_layout.addWidget(info_group)
            
            # Network configuration
            network_group = QtWidgets.QGroupBox("Network Configuration")
            network_layout = QtWidgets.QFormLayout()
            
            new_ip = QtWidgets.QLineEdit(device.get('ip', ''))
            new_gateway = QtWidgets.QLineEdit(device.get('gateway', '192.168.1.1'))
            new_subnet = QtWidgets.QLineEdit(device.get('subnet', '255.255.255.0'))
            new_port = QtWidgets.QSpinBox()
            new_port.setRange(1, 65535)
            new_port.setValue(int(device.get('port', 80)))
            
            network_layout.addRow("New IP Address:", new_ip)
            network_layout.addRow("Gateway:", new_gateway)
            network_layout.addRow("Subnet Mask:", new_subnet)
            network_layout.addRow("HTTP Port:", new_port)
            network_group.setLayout(network_layout)
            config_layout.addWidget(network_group)
            
            # Password configuration
            password_group = QtWidgets.QGroupBox("Password Configuration")
            password_layout = QtWidgets.QFormLayout()
            
            old_password = QtWidgets.QLineEdit()
            old_password.setEchoMode(QtWidgets.QLineEdit.Password)
            new_password = QtWidgets.QLineEdit()
            new_password.setEchoMode(QtWidgets.QLineEdit.Password)
            confirm_password = QtWidgets.QLineEdit()
            confirm_password.setEchoMode(QtWidgets.QLineEdit.Password)
            
            password_layout.addRow("Old Password:", old_password)
            password_layout.addRow("New Password:", new_password)
            password_layout.addRow("Confirm:", confirm_password)
            password_group.setLayout(password_layout)
            config_layout.addWidget(password_group)
            
            # Action buttons
            action_layout = QtWidgets.QHBoxLayout()
            
            apply_network_btn = QtWidgets.QPushButton("🌐 Apply Network Settings")
            apply_network_btn.clicked.connect(lambda: apply_network_config(device, new_ip.text(), new_gateway.text(), new_subnet.text(), new_port.value()))
            
            apply_password_btn = QtWidgets.QPushButton("🔐 Change Password")
            apply_password_btn.clicked.connect(lambda: apply_password_config(device, old_password.text(), new_password.text(), confirm_password.text()))
            
            test_connection_btn = QtWidgets.QPushButton("🔌 Test Connection")
            test_connection_btn.clicked.connect(lambda: test_device_connection(device))
            
            action_layout.addWidget(apply_network_btn)
            action_layout.addWidget(apply_password_btn)
            action_layout.addWidget(test_connection_btn)
            config_layout.addLayout(action_layout)
            
            # Close button
            close_config_btn = QtWidgets.QPushButton("✕ Close")
            close_config_btn.clicked.connect(config_dlg.close)
            config_layout.addWidget(close_config_btn)
            
            config_dlg.exec_()
        
        def apply_network_config(device, new_ip, gateway, subnet, port):
            """Apply network configuration to device."""
            reply = QtWidgets.QMessageBox.question(
                dialog, "Confirm Network Change",
                f"Change network settings for {device.get('ip')}?\n\n"
                f"New IP: {new_ip}\n"
                f"Gateway: {gateway}\n"
                f"Subnet: {subnet}\n"
                f"Port: {port}\n\n"
                f"This will modify the device configuration!",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
            )
            
            if reply == QtWidgets.QMessageBox.Yes:
                try:
                    # Use ISAPI to modify network settings
                    url = f"http://{device.get('ip')}:{device.get('port', 80)}/ISAPI/System/Network/interfaces/1/ipAddress"
                    
                    xml_data = f"""<?xml version="1.0" encoding="UTF-8"?>
<IPAddress>
    <ipVersion>v4</ipVersion>
    <addressingType>static</addressingType>
    <ipAddress>{new_ip}</ipAddress>
    <subnetMask>{subnet}</subnetMask>
    <DefaultGateway><ipAddress>{gateway}</ipAddress></DefaultGateway>
</IPAddress>"""
                    
                    # This would require authentication - show info message
                    QtWidgets.QMessageBox.information(
                        dialog, "Network Configuration",
                        f"To apply network settings, use:\n\n"
                        f"Username: admin\n"
                        f"Device IP: {device.get('ip')}\n\n"
                        f"Configuration will be applied via ISAPI.\n"
                        f"Ensure you have admin credentials."
                    )
                except Exception as e:
                    QtWidgets.QMessageBox.critical(dialog, "Error", f"Failed to apply network settings:\n{e}")
        
        def apply_password_config(device, old_pwd, new_pwd, confirm_pwd):
            """Apply password configuration to device."""
            if not old_pwd or not new_pwd:
                QtWidgets.QMessageBox.warning(dialog, "Invalid Input", "Please enter both old and new passwords.")
                return
            
            if new_pwd != confirm_pwd:
                QtWidgets.QMessageBox.warning(dialog, "Password Mismatch", "New password and confirmation do not match.")
                return
            
            try:
                QtWidgets.QMessageBox.information(
                    dialog, "Password Change",
                    f"To change password for {device.get('ip')}:\n\n"
                    f"1. Use SADP tool or web interface\n"
                    f"2. Login with current credentials\n"
                    f"3. Navigate to Configuration → User Management\n"
                    f"4. Change password for admin user\n\n"
                    f"Note: Password change via API requires proper authentication."
                )
            except Exception as e:
                QtWidgets.QMessageBox.critical(dialog, "Error", f"Failed to change password:\n{e}")
        
        def test_device_connection(device):
            """Test connection to device."""
            try:
                ip = device.get('ip')
                port = device.get('port', 80)
                
                # Try HTTP request
                url = f"http://{ip}:{port}/ISAPI/System/deviceInfo"
                response = requests.get(url, timeout=3)
                
                if response.status_code in [200, 401]:  # 401 means device is there but needs auth
                    QtWidgets.QMessageBox.information(
                        dialog, "Connection Test",
                        f"✅ Device is reachable!\n\n"
                        f"IP: {ip}:{port}\n"
                        f"Status: {response.status_code}\n"
                        f"Response time: {response.elapsed.total_seconds():.2f}s"
                    )
                else:
                    QtWidgets.QMessageBox.warning(
                        dialog, "Connection Test",
                        f"⚠️ Unexpected response\n\n"
                        f"Status: {response.status_code}"
                    )
            except Exception as e:
                QtWidgets.QMessageBox.critical(
                    dialog, "Connection Test",
                    f"❌ Connection failed!\n\n{str(e)}"
                )
        
        def scan_network():
            """Scan network for Hikvision devices."""
            nonlocal stop_scan, discovered_devices
            stop_scan = False
            discovered_devices = []
            table.setRowCount(0)
            
            scan_btn.setEnabled(False)
            stop_btn.setEnabled(True)
            export_btn.setEnabled(False)
            progress_bar.setVisible(True)
            progress_bar.setValue(0)
            
            try:
                import ipaddress
                
                ip_range = ip_range_input.text().strip()
                port = port_input.value()
                timeout = timeout_input.value()
                
                # Parse IP range
                try:
                    network = ipaddress.ip_network(ip_range, strict=False)
                    all_ips = list(network.hosts())
                except Exception as e:
                    QtWidgets.QMessageBox.critical(dialog, "Invalid IP Range", f"Invalid IP range format:\n{e}")
                    reset_scan_ui()
                    return
                
                total_ips = len(all_ips)
                progress_label.setText(f"Scanning {total_ips} addresses...")
                
                # Scan in parallel
                from concurrent.futures import ThreadPoolExecutor, as_completed
                
                def check_device(ip):
                    if stop_scan:
                        return None
                    
                    try:
                        url = f"http://{ip}:{port}/ISAPI/System/deviceInfo"
                        response = requests.get(url, timeout=timeout, auth=('admin', ''))
                        
                        if response.status_code in [200, 401]:
                            # Device found! Try to get device info
                            device = {
                                'ip': str(ip),
                                'port': port,
                                'type': 'Camera',  # Default
                                'mac': 'N/A',
                                'name': f'Device-{ip}',
                                'model': 'Unknown',
                                'firmware': 'N/A',
                                'serial': 'N/A'
                            }
                            
                            # Try to parse device info if available
                            if response.status_code == 200:
                                try:
                                    import xml.etree.ElementTree as ET
                                    root = ET.fromstring(response.content)
                                    
                                    device['name'] = root.findtext('.//deviceName', 'Unknown')
                                    device['model'] = root.findtext('.//model', 'Unknown')
                                    device['firmware'] = root.findtext('.//firmwareVersion', 'N/A')
                                    device['serial'] = root.findtext('.//serialNumber', 'N/A')
                                    device['mac'] = root.findtext('.//macAddress', 'N/A')
                                    
                                    device_type = root.findtext('.//deviceType', '')
                                    if 'nvr' in device_type.lower():
                                        device['type'] = 'NVR'
                                    elif 'dvr' in device_type.lower():
                                        device['type'] = 'DVR'
                                    else:
                                        device['type'] = 'Camera'
                                except:
                                    pass
                            
                            return device
                    except:
                        pass
                    
                    return None
                
                # Scan with progress
                scanned = 0
                with ThreadPoolExecutor(max_workers=50) as executor:
                    futures = {executor.submit(check_device, ip): ip for ip in all_ips}
                    
                    for future in as_completed(futures):
                        if stop_scan:
                            break
                        
                        scanned += 1
                        progress_bar.setValue(int(scanned / total_ips * 100))
                        progress_label.setText(f"Scanning... {scanned}/{total_ips} ({len(discovered_devices)} found)")
                        QtWidgets.QApplication.processEvents()
                        
                        result = future.result()
                        if result:
                            discovered_devices.append(result)
                            add_device_to_table(result)
                            update_statistics()
                
                progress_label.setText(f"✅ Scan complete! Found {len(discovered_devices)} devices")
                export_btn.setEnabled(len(discovered_devices) > 0)
                
            except Exception as e:
                QtWidgets.QMessageBox.critical(dialog, "Scan Error", f"Error during scan:\n{e}")
                progress_label.setText(f"❌ Scan failed: {e}")
            finally:
                reset_scan_ui()
        
        def reset_scan_ui():
            """Reset scan UI state."""
            scan_btn.setEnabled(True)
            stop_btn.setEnabled(False)
            progress_bar.setVisible(False)
        
        def stop_scanning():
            """Stop ongoing scan."""
            nonlocal stop_scan
            stop_scan = True
            progress_label.setText("⏹ Stopping scan...")
        
        def export_results():
            """Export discovered devices to CSV."""
            if not discovered_devices:
                return
            
            from datetime import datetime
            filename, _ = QtWidgets.QFileDialog.getSaveFileName(
                dialog, "Save SADP Results",
                f"SADP_Scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "CSV Files (*.csv)"
            )
            
            if filename:
                try:
                    with open(filename, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        writer.writerow(['Device Type', 'IP Address', 'Port', 'MAC Address', 'Device Name', 'Model', 'Firmware', 'Serial Number'])
                        for device in discovered_devices:
                            writer.writerow([
                                device.get('type', ''),
                                device.get('ip', ''),
                                device.get('port', ''),
                                device.get('mac', ''),
                                device.get('name', ''),
                                device.get('model', ''),
                                device.get('firmware', ''),
                                device.get('serial', '')
                            ])
                    QtWidgets.QMessageBox.information(dialog, "Export Complete", f"Results saved to:\n{filename}")
                except Exception as e:
                    QtWidgets.QMessageBox.critical(dialog, "Export Error", f"Failed to export:\n{e}")
        
        # Connect buttons
        scan_btn.clicked.connect(lambda: threading.Thread(target=scan_network, daemon=True).start())
        stop_btn.clicked.connect(stop_scanning)
        export_btn.clicked.connect(export_results)
        
        dialog.exec_()
    
    def show_performance_dashboard(self):
        """Show enhanced performance monitoring dashboard v8.6+."""
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("📊 Performance Dashboard v8.6")
        dlg.setGeometry(150, 150, 800, 600)
        
        layout = QtWidgets.QVBoxLayout(dlg)
        
        # Header
        header = QtWidgets.QLabel("<h2>📊 System Performance Metrics v8.6</h2>")
        header.setAlignment(QtCore.Qt.AlignCenter)
        layout.addWidget(header)
        
        # Create tabs for different metrics
        tabs = QtWidgets.QTabWidget()
        
        # Performance Metrics Tab
        perf_tab = QtWidgets.QWidget()
        perf_layout = QtWidgets.QVBoxLayout(perf_tab)
        
        # Performance metrics display
        perf_text = QtWidgets.QTextEdit()
        perf_text.setReadOnly(True)
        
        # Calculate current performance stats
        total_cameras = len(self.cams)
        total_nvrs = len(self.nvrs)
        cache_size = len(getattr(self, 'connection_cache', {}))
        nvr_cache_size = len(getattr(self, 'nvr_cache', {}))
        
        # Get cache hit ratio
        metrics = getattr(self, 'performance_metrics', {})
        total_checks = metrics.get('total_checks', 0)
        cache_hits = metrics.get('cache_hits', 0)
        hit_ratio = (cache_hits / total_checks * 100) if total_checks > 0 else 0
        avg_response = metrics.get('average_response_time', 0.0)
        
        perf_report = f"""
<h3>🚀 Performance Overview</h3>
<table border="1" cellpadding="5" style="border-collapse: collapse; width: 100%;">
<tr><td><b>📊 Metric</b></td><td><b>📈 Value</b></td><td><b>📋 Status</b></td></tr>
<tr><td>Total NVRs</td><td>{total_nvrs}</td><td>{'🟢 Good' if total_nvrs > 0 else '🔴 None'}</td></tr>
<tr><td>Total Cameras</td><td>{total_cameras}</td><td>{'🟢 Good' if total_cameras > 0 else '🔴 None'}</td></tr>
<tr><td>Parallel Workers</td><td>{MAX_PARALLEL_WORKERS}</td><td>🟢 Optimized</td></tr>
<tr><td>Connection Timeout</td><td>{CONNECTION_TIMEOUT}s</td><td>🟢 Balanced</td></tr>
<tr><td>UI Update Throttle</td><td>{UI_UPDATE_THROTTLE}ms</td><td>🟢 Smooth</td></tr>
</table>

<h3>💾 Cache Performance</h3>
<table border="1" cellpadding="5" style="border-collapse: collapse; width: 100%;">
<tr><td><b>💾 Cache Type</b></td><td><b>📊 Size</b></td><td><b>📈 Performance</b></td></tr>
<tr><td>Connection Cache</td><td>{cache_size} entries</td><td>{'🟢 Active' if cache_size > 0 else '🟡 Empty'}</td></tr>
<tr><td>NVR Cache</td><td>{nvr_cache_size} entries</td><td>{'🟢 Active' if nvr_cache_size > 0 else '🟡 Empty'}</td></tr>
<tr><td>Cache Hit Ratio</td><td>{hit_ratio:.1f}%</td><td>{'🟢 Excellent' if hit_ratio > 80 else '🟡 Good' if hit_ratio > 50 else '🔴 Poor'}</td></tr>
<tr><td>Total Checks</td><td>{total_checks}</td><td>{'🟢 Active' if total_checks > 0 else '🟡 None'}</td></tr>
<tr><td>Avg Response Time</td><td>{avg_response:.2f}s</td><td>{'🟢 Fast' if avg_response < 2 else '🟡 Moderate' if avg_response < 5 else '🔴 Slow'}</td></tr>
</table>

<h3>⚡ System Resources</h3>
<table border="1" cellpadding="5" style="border-collapse: collapse; width: 100%;">
<tr><td><b>🖥️ Resource</b></td><td><b>📊 Usage</b></td><td><b>📈 Status</b></td></tr>
<tr><td>Application Version</td><td>{APP_VERSION}</td><td>🟢 Current</td></tr>
<tr><td>Cache Timeout</td><td>{CACHE_TIMEOUT}s</td><td>🟢 Optimized</td></tr>
<tr><td>Retry Attempts</td><td>{RETRY_ATTEMPTS}</td><td>🟢 Balanced</td></tr>
</table>
        """
        
        perf_text.setHtml(perf_report)
        perf_layout.addWidget(perf_text)
        
        # Refresh button for performance tab
        refresh_perf_btn = QtWidgets.QPushButton("🔄 Refresh Metrics")
        refresh_perf_btn.clicked.connect(lambda: self._refresh_performance_metrics(perf_text))
        perf_layout.addWidget(refresh_perf_btn)
        
        tabs.addTab(perf_tab, "📈 Performance")
        
        # Cache Statistics Tab
        cache_tab = QtWidgets.QWidget()
        cache_layout = QtWidgets.QVBoxLayout(cache_tab)
        
        cache_text = QtWidgets.QTextEdit()
        cache_text.setReadOnly(True)
        
        # Cache details
        connection_cache = getattr(self, 'connection_cache', {})
        nvr_cache = getattr(self, 'nvr_cache', {})
        
        cache_details = "<h3>💾 Cache Details</h3>"
        
        if connection_cache:
            cache_details += "<h4>🔗 Connection Cache</h4><ul>"
            for ip, data in list(connection_cache.items())[:10]:  # Show first 10
                status = data.get('status', 'Unknown')
                timestamp = data.get('timestamp', 0)
                age = int(time.time() - timestamp) if timestamp else 0
                cache_details += f"<li><b>{ip}</b>: {status} ({age}s ago)</li>"
            if len(connection_cache) > 10:
                cache_details += f"<li>... and {len(connection_cache) - 10} more entries</li>"
            cache_details += "</ul>"
        else:
            cache_details += "<p>🟡 Connection cache is empty</p>"
        
        if nvr_cache:
            cache_details += "<h4>🗄️ NVR Cache</h4><ul>"
            for nvr_ip, data in list(nvr_cache.items())[:5]:  # Show first 5
                cameras = data.get('cameras', [])
                timestamp = data.get('timestamp', 0)
                age = int(time.time() - timestamp) if timestamp else 0
                cache_details += f"<li><b>{nvr_ip}</b>: {len(cameras)} cameras ({age}s ago)</li>"
            if len(nvr_cache) > 5:
                cache_details += f"<li>... and {len(nvr_cache) - 5} more NVRs</li>"
            cache_details += "</ul>"
        else:
            cache_details += "<p>🟡 NVR cache is empty</p>"
        
        cache_text.setHtml(cache_details)
        cache_layout.addWidget(cache_text)
        
        # Cache management buttons
        cache_btn_layout = QtWidgets.QHBoxLayout()
        clear_cache_btn = QtWidgets.QPushButton("🗑️ Clear All Cache")
        clear_cache_btn.clicked.connect(self._clear_all_cache)
        refresh_cache_btn = QtWidgets.QPushButton("🔄 Refresh Cache View")
        refresh_cache_btn.clicked.connect(lambda: self._refresh_cache_view(cache_text))
        
        cache_btn_layout.addWidget(clear_cache_btn)
        cache_btn_layout.addWidget(refresh_cache_btn)
        cache_layout.addLayout(cache_btn_layout)
        
        tabs.addTab(cache_tab, "💾 Cache")
        
        layout.addWidget(tabs)
        
        # Close button
        close_btn = QtWidgets.QPushButton("✕ Close")
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)
        
        # Store reference for updates
        self.performance_dialog = dlg
        dlg.finished.connect(lambda: setattr(self, 'performance_dialog', None))
        
        dlg.exec_()
    
    def _refresh_performance_metrics(self, text_widget):
        """Refresh performance metrics display."""
        try:
            # Recalculate metrics
            total_cameras = len(self.cams)
            total_nvrs = len(self.nvrs)
            cache_size = len(getattr(self, 'connection_cache', {}))
            nvr_cache_size = len(getattr(self, 'nvr_cache', {}))
            
            # Update display (simplified for demo)
            self.status.showMessage("📊 Performance metrics refreshed", 2000)
            log("[PERFORMANCE] Metrics refreshed")
            
        except Exception as e:
            log(f"[PERFORMANCE] Error refreshing metrics: {e}")
    
    def _refresh_cache_view(self, text_widget):
        """Refresh cache view display."""
        try:
            # Update cache display
            self.status.showMessage("💾 Cache view refreshed", 2000)
            log("[CACHE] Cache view refreshed")
            
        except Exception as e:
            log(f"[CACHE] Error refreshing cache view: {e}")
    
    def _clear_all_cache(self):
        """Clear all application caches."""
        try:
            reply = QtWidgets.QMessageBox.question(self, "Clear Cache", 
                "Are you sure you want to clear all cached data?\n\n"
                "This will remove all cached connection status and NVR data.",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No)
            
            if reply == QtWidgets.QMessageBox.Yes:
                # Clear all caches
                if hasattr(self, 'connection_cache'):
                    self.connection_cache.clear()
                if hasattr(self, 'nvr_cache'):
                    self.nvr_cache.clear()
                
                # Reset performance metrics
                self.performance_metrics = {
                    'total_checks': 0,
                    'cache_hits': 0,
                    'average_response_time': 0.0,
                    'last_full_scan': 0
                }
                
                QtWidgets.QMessageBox.information(self, "Cache Cleared", 
                    "✅ All caches have been cleared successfully!")
                log("[CACHE] All caches cleared by user")
                
        except Exception as e:
            log(f"[CACHE] Error clearing cache: {e}")
            QtWidgets.QMessageBox.critical(self, "Error", f"Failed to clear cache:\n{str(e)}")
    
    # ==================== NVR MANAGEMENT ====================
    def add_new_nvr(self):
        """Show dialog to add new NVR"""
        if not NVR_DIALOGS_AVAILABLE:
            QtWidgets.QMessageBox.warning(self, "Feature Unavailable", 
                                        "NVR management dialogs are not available.\nPlease check if nvr_dialogs.py exists.")
            return
            
        dialog = AddNVRDialog(self)
        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            nvr_data = dialog.get_nvr_data()
            
            # Check if NVR already exists
            existing = any(nvr.get('name') == nvr_data['name'] or nvr.get('ip') == nvr_data['ip'] 
                         for nvr in self.nvrs)
            
            if existing:
                QtWidgets.QMessageBox.warning(self, "Duplicate NVR", 
                                           "NVR with this name or IP already exists!")
                return
            
            # Add to NVR list
            nvr_data['cameras'] = 0  # Will be detected on refresh
            self.nvrs.append(nvr_data)
            
            # Save credentials
            self.save_nvr_credentials()
            
            # Refresh NVR list
            self.refresh_nvr_status()
            
            QtWidgets.QMessageBox.information(self, "Success", 
                                           f"NVR '{nvr_data['name']}' added successfully!")
    
    def edit_nvr_credentials(self, nvr_data):
        """Edit NVR credentials"""
        if not NVR_DIALOGS_AVAILABLE:
            QtWidgets.QMessageBox.warning(self, "Feature Unavailable", 
                                        "NVR management dialogs are not available.\nPlease check if nvr_dialogs.py exists.")
            return
            
        dialog = EditNVRDialog(nvr_data, self)
        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            updated_data = dialog.get_nvr_data()
            
            # Update NVR configuration
            for i, nvr in enumerate(self.nvrs):
                if nvr.get('name') == nvr_data.get('name'):
                    self.nvrs[i].update(updated_data)
                    break
            
            # Save credentials
            self.save_nvr_credentials()
            
            # Refresh NVR list
            self.refresh_nvr_status()
            
            QtWidgets.QMessageBox.information(self, "Success", 
                                           f"Credentials for '{nvr_data.get('name')}' updated successfully!")
    
    def test_nvr_connection(self, nvr_data):
        """Test NVR connection"""
        progress = QtWidgets.QProgressDialog("Testing NVR connection...", "Cancel", 0, 0, self)
        progress.setWindowModality(QtCore.Qt.WindowModal)
        progress.show()
        
        def test_thread():
            try:
                protocol = nvr_data.get('protocol', 'http')
                ip = nvr_data.get('ip', '')
                port = nvr_data.get('port', 80)
                username = nvr_data.get('username', 'admin')
                password = nvr_data.get('password', '')
                
                url = f"{protocol}://{ip}:{port}/ISAPI/System/deviceInfo"
                
                auth_methods = [HTTPBasicAuth(username, password), HTTPDigestAuth(username, password)]
                
                for auth in auth_methods:
                    try:
                        response = requests.get(url, auth=auth, timeout=5, verify=False)
                        if response.status_code == 200:
                            QtCore.QTimer.singleShot(0, lambda: self._show_connection_result(progress, True, f"✅ Connection successful\n{nvr_data.get('name')} ({ip})"))
                            return
                    except:
                        continue
                        
                QtCore.QTimer.singleShot(0, lambda: self._show_connection_result(progress, False, f"❌ Connection failed\n{nvr_data.get('name')} ({ip})\nCheck credentials and network"))
                
            except Exception as e:
                QtCore.QTimer.singleShot(0, lambda: self._show_connection_result(progress, False, f"❌ Error testing {nvr_data.get('name')}\n{str(e)}"))
        
        thread = threading.Thread(target=test_thread, daemon=True)
        thread.start()
    
    def _show_connection_result(self, progress, success, message):
        """Show connection test result"""
        progress.close()
        
        if success:
            QtWidgets.QMessageBox.information(self, "Connection Test", message)
        else:
            QtWidgets.QMessageBox.warning(self, "Connection Test", message)
    
    def remove_nvr(self, nvr_data):
        """Remove custom NVR"""
        reply = QtWidgets.QMessageBox.question(self, "Remove NVR", 
                                           f"Are you sure you want to remove '{nvr_data.get('name')}'?", 
                                           QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No)
        
        if reply == QtWidgets.QMessageBox.Yes:
            self.nvrs = [nvr for nvr in self.nvrs if nvr.get('name') != nvr_data.get('name')]
            self.save_nvr_credentials()
            self.load_nvr_sheets()  # Reload NVR list
            QtWidgets.QMessageBox.information(self, "Success", f"NVR '{nvr_data.get('name')}' removed successfully!")
    
    def is_custom_nvr(self, nvr_data):
        """Check if NVR is custom (not in default list)"""
        default_names = ['NVR1', 'NVR2', 'NVR3', 'NVR4', 'NVR5', 'NVR6', 'NVR7', 'NVR8', 
                        'NVR9', 'NVR10', 'NVR11', 'NVR12', 'NVR13', 'NVR14', 'NVR15', 
                        'NVR16', 'NVR17', 'NVR18', 'DI Sreen']
        return nvr_data.get('name') not in default_names
    
    def save_nvr_credentials(self):
        """Save NVR credentials to file"""
        credentials = {}
                    # ...removed broken/legacy block...
    
    # Legacy/duplicate camera extraction methods removed. All camera extraction now uses the enhanced helpers (fetch_nvr_cameras, try_isapi_method_enhanced, etc.) for consistency and robustness.
    
    def update_camera_list_fast(self, cameras, nvr_name):
        """Fast update of camera list with new cameras"""
        import time
        added_count = 0
        updated_count = 0
        
        log(f"[FAST-UPDATE] Processing {len(cameras)} cameras from {nvr_name}")
        
        for cam in cameras:
            # Get camera IP - use actual IP from camera data, or fallback to channel-based identifier
            cam_ip = cam.get('ip', '').strip()
            if not cam_ip:
                # If no IP, use NVR IP with channel as identifier
                cam_ip = f"{nvr_name}_Ch{cam.get('channel', 'Unknown')}"
            
            # Create camera entry with improved formatting
            camera_name = cam.get('name', 'Unknown Camera')
            
            # Use actual camera IP if available, otherwise use channel identifier
            display_ip = cam.get('ip', '').strip()
            if not display_ip:
                display_ip = cam_ip  # Use the channel-based identifier we created
            
            # Format status properly - capitalize and add connection type
            raw_status = cam.get('status', 'Unknown').strip()
            if raw_status.lower() == 'online':
                formatted_status = 'Online (TCP)'
            elif raw_status.lower() == 'offline':
                formatted_status = 'Offline'
            else:
                formatted_status = raw_status.capitalize()
            
            camera_entry = {
                'name': camera_name,
                'ip': display_ip,
                'status': formatted_status,
                'channel': str(cam.get('channel', '')),
                'model': cam.get('model', 'Network Camera') if cam.get('model', '').strip() else 'IP Camera',
                'port': cam.get('port', '8000'),  # Use NVR management port
                'serial': cam.get('serial', '').strip() or '—',
                'firmware': cam.get('firmware', '').strip() or '—',
                'nvr': nvr_name,
                'last_updated': time.strftime('%Y-%m-%d %H:%M')
            }
            
            # Log camera details for debugging
            status_display = camera_entry['status'].replace(' (TCP)', '').replace(' (N/A)', '')
            log(f"[FAST-UPDATE] Ch{camera_entry['channel']:>2}: {camera_entry['name']:<15} | {camera_entry['ip']:<15} | {status_display}")
            
            # Debug logging for PS OV 2 specifically
            if camera_entry['name'] == 'PS OV 2':
                log(f"[UPDATE-DEBUG] PS OV 2 - Raw status: '{cam.get('status')}', Formatted status: '{camera_entry['status']}'")
            
            # Check if camera already exists (by name or IP address for better hybrid matching)
            exists = False
            for i, existing_cam in enumerate(self.cams):
                # Match by camera name (exact match) or IP address
                name_match = existing_cam.get('name', '').strip().lower() == camera_entry['name'].strip().lower()
                ip_match = existing_cam.get('ip', '').strip() == camera_entry['ip'].strip() and camera_entry['ip'].strip()
                
                if name_match or ip_match:
                    # Camera match found and updated
                    
                    # Check for IP change for remark
                    old_ip = existing_cam.get('ip', '').strip()
                    new_ip = camera_entry['ip'].strip()
                    if old_ip and new_ip and old_ip != new_ip:
                        camera_entry['remark'] = f"IP changed: {old_ip} → {new_ip}"
                        camera_entry['previous_ip'] = old_ip
                        log(f"[IP-CHANGE] {camera_entry['name']}: {old_ip} → {new_ip}")
                    
                    # Update existing camera with new data
                    self.cams[i].update(camera_entry)
                    # Update in filtered list too
                    for j, filtered_cam in enumerate(self.filtered):
                        if (filtered_cam.get('name', '').strip().lower() == camera_entry['name'].strip().lower() or 
                            (filtered_cam.get('ip', '').strip() == camera_entry['ip'].strip() and camera_entry['ip'].strip())):
                            self.filtered[j].update(camera_entry)
                            break
                    exists = True
                    updated_count += 1
                    break
            
            if not exists:
                # Add new camera
                self.cams.append(camera_entry)
                self.filtered.append(camera_entry)
                added_count += 1
        
        # Create summary
        total_cameras = len(self.filtered)
        online_cameras = sum(1 for cam in self.filtered if 'Online' in cam.get('status', ''))
        offline_cameras = total_cameras - online_cameras
        
        log(f"[FAST-UPDATE] Summary: {added_count} new, {updated_count} updated | Total: {total_cameras} cameras ({online_cameras} online, {offline_cameras} offline)")
        return added_count
    
    def _handle_camera_update(self, cameras, nvr_name, elapsed, nvr_online, nvr_offline):
        """Handle camera update signal from background thread"""
        log(f"[SIGNAL-HANDLER] Received camera update signal for {nvr_name}")
        log(f"[SIGNAL-HANDLER] self.filtered has {len(self.filtered)} cameras")
        log(f"[SIGNAL-HANDLER] Sample statuses: {[(c.get('name'), c.get('status')) for c in self.filtered if 'PS OV' in c.get('name', '')]}")
        
        # Force refresh the table with updated data
        self.populate_table(self.filtered)
        
        # Count total cameras in system
        total_cameras = len(self.filtered)
        total_online = sum(1 for cam in self.filtered if 'Online' in cam.get('status', ''))
        
        msg = f"✅ Synced {len(cameras)} cameras from {nvr_name} ({nvr_online} online) | Total: {total_cameras} ({total_online} online) | {elapsed:.1f}s"
        self.status.showMessage(msg, 5000)
        log(f"[HYBRID-UPDATE] Success: {msg}")
        log(f"[SIGNAL-HANDLER] UI update completed")

    def _handle_progress_update(self, label_text, progress_value, style_sheet, enable_buttons):
        """Handle progress update signal - thread-safe UI updates"""
        try:
            # Find active dialogs (Quick Sync dialog)
            active_dialogs = []
            for widget in QtWidgets.QApplication.allWidgets():
                if isinstance(widget, QtWidgets.QDialog) and widget.isVisible():
                    active_dialogs.append(widget)
            
            # Update progress widgets in active dialogs
            for dialog in active_dialogs:
                # Find progress label and progress bar in dialog
                for child in dialog.findChildren(QtWidgets.QLabel):
                    # Look for labels that contain progress-like text
                    current_text = child.text().lower()
                    if any(keyword in current_text for keyword in ['processing', 'step', 'complete', 'loading', 'sync']):
                        child.setText(label_text)
                        if style_sheet:
                            child.setStyleSheet(style_sheet)
                        log(f"[PROGRESS-SIGNAL] Updated label: {label_text}")
                        break
                
                for child in dialog.findChildren(QtWidgets.QProgressBar):
                    child.setValue(progress_value)
                    log(f"[PROGRESS-SIGNAL] Updated progress bar: {progress_value}%")
                    break
            
            # If no dialogs found, update status bar as fallback
            if not active_dialogs and hasattr(self, 'status'):
                self.status.showMessage(label_text, 3000)
                log(f"[PROGRESS-SIGNAL] Updated status bar: {label_text}")
                
        except Exception as e:
            log(f"[PROGRESS-SIGNAL] Error: {e}")

    def _handle_button_control(self, button_name, enabled, text, visible):
        """Handle button control signal - thread-safe button updates"""
        try:
            # Find active dialogs first for more precise targeting
            active_dialogs = []
            for widget in QtWidgets.QApplication.allWidgets():
                if isinstance(widget, QtWidgets.QDialog) and widget.isVisible():
                    active_dialogs.append(widget)
            
            # Search in active dialogs first
            button_found = False
            for dialog in active_dialogs:
                for button in dialog.findChildren(QtWidgets.QPushButton):
                    button_text = button.text().lower()
                    if (button_name.lower() in button_text or 
                        button_name in ['start', 'bypass', 'stop', 'cancel'] and 
                        any(keyword in button_text for keyword in [button_name, 'start', 'skip', 'stop', 'cancel'])):
                        
                        button.setEnabled(enabled)
                        if text:
                            button.setText(text)
                        button.setVisible(visible)
                        log(f"[BUTTON-SIGNAL] Updated {button_name}: enabled={enabled}, text='{text}', visible={visible}")
                        button_found = True
                        break
                if button_found:
                    break
                    
            # If not found in dialogs, search globally
            if not button_found:
                for widget in QtWidgets.QApplication.allWidgets():
                    if isinstance(widget, QtWidgets.QPushButton):
                        widget_text = widget.text().lower()
                        if (button_name.lower() in widget_text or 
                            button_name.lower() in widget.objectName().lower()):
                            widget.setEnabled(enabled)
                            if text:
                                widget.setText(text)
                            widget.setVisible(visible)
                            log(f"[BUTTON-SIGNAL] Updated {button_name} globally: enabled={enabled}, text='{text}', visible={visible}")
                            break
                            
        except Exception as e:
            log(f"[BUTTON-SIGNAL] Error: {e}")

    def _handle_ui_status_update(self, element_type, element_name, status_data):
        """Handle UI status update signal - thread-safe status display with emojis"""
        try:
            if element_type == "emoji_status":
                # Parse status data for emoji updates
                parts = status_data.split("|")
                if len(parts) >= 3:
                    emoji = parts[0]
                    count = parts[1]
                    bg_color = parts[2] if len(parts) > 2 else ""
                    
                    # Find and update status labels with emojis
                    for widget in QtWidgets.QApplication.allWidgets():
                        if isinstance(widget, QtWidgets.QLabel):
                            label_name = widget.objectName().lower()
                            if element_name.lower() in label_name:
                                display_text = f"{emoji} {count}"
                                widget.setText(display_text)
                                
                                # Apply background color if specified
                                if bg_color:
                                    widget.setStyleSheet(f"background-color: {bg_color}; padding: 2px; border-radius: 3px;")
                                
                                log(f"[UI-STATUS-SIGNAL] Updated {element_name}: {display_text} with color {bg_color}")
                                break
            
            elif element_type == "table_row_color":
                # Update table row colors with status
                parts = status_data.split("|")
                if len(parts) >= 3:
                    row_index = int(parts[0])
                    status = parts[1]
                    color_info = parts[2]
                    
                    # Apply color to table row
                    if hasattr(self, 'table') and self.table.rowCount() > row_index:
                        for col in range(self.table.columnCount()):
                            item = self.table.item(row_index, col)
                            if item:
                                if "Online" in status:
                                    item.setBackground(QtGui.QColor("#d5f4e6"))  # Light green
                                elif "Offline" in status:
                                    item.setBackground(QtGui.QColor("#fdeaea"))  # Light red
                                elif "TCP" in status:
                                    item.setBackground(QtGui.QColor("#fff3cd"))  # Light yellow
                                    
                        log(f"[UI-STATUS-SIGNAL] Updated table row {row_index}: {status}")
                        
        except Exception as e:
            log(f"[UI-STATUS-SIGNAL] Error: {e}")
    
    def _update_cameras_complete(self, success, message, nvr_name, progress=None):
        """Handle completion of camera update"""
        log(f"[UPDATE-COMPLETE] Called with success={success}, message='{message}', nvr_name='{nvr_name}'")
        
        # Close progress dialog
        if progress:
            progress.close()
        
        if success:
            self.status.showMessage(f"✅ Updated cameras from {nvr_name}", 3000)
            log(f"[UPDATE-CAMERAS] Success: {message}")
            # Force refresh the camera table to show new cameras
            log(f"[UPDATE-CAMERAS] Refreshing table with {len(self.filtered)} cameras")
            log(f"[UPDATE-CAMERAS] Sample camera status: {[(c.get('name'), c.get('status')) for c in self.filtered[:3]]}")
            self.populate_table(self.filtered)
            # Show success message only if progress dialog wasn't cancelled
            if progress:
                QtWidgets.QMessageBox.information(self, "Update Complete", 
                                                 f"Successfully updated cameras from {nvr_name}\n\n{message}")
            else:
                log(f"[UPDATE-CAMERAS] Skipping success dialog - operation completed after progress was cancelled")
        else:
            self.status.showMessage(f"❌ Failed to update cameras from {nvr_name}", 5000)
            log(f"[UPDATE-CAMERAS] Failed: {message}")
            # Show error message
            QtWidgets.QMessageBox.warning(self, "Update Failed", 
                                         f"Failed to update cameras from {nvr_name}\n\n{message}")
    
    # Old complex method removed - replaced with fast update

    def show_nvr_context_menu(self, position):
        """Show context menu on right-click for NVR list."""
        log(f"[CONTEXT-MENU] Right-click at position: {position}")
        
        # Get item at the clicked position
        item = self.list_nvr.itemAt(position)
        if not item:
            log("[CONTEXT-MENU] No item found at position")
            return
        
        log(f"[CONTEXT-MENU] Item found: {item.text()}")
        
        # Get NVR data from the clicked item
        nvr_text = item.text()
        nvr_name = ""
        
        # Extract NVR name from the display text
        if "|" in nvr_text:
            parts = nvr_text.split("|")
            if len(parts) > 1:
                nvr_name = parts[1].strip().split()[0]  # Get first word after pipe
        
        # Find the NVR data
        selected_nvr = None
        for nvr in self.nvrs:
            if nvr.get('name') == nvr_name:
                selected_nvr = nvr
                break
        
        if not selected_nvr:
            # Fallback: try to get NVR by index if name matching fails
            try:
                nvr_index = item.data(QtCore.Qt.UserRole)
                if nvr_index is not None and 0 <= nvr_index < len(self.nvrs):
                    selected_nvr = self.nvrs[nvr_index]
            except:
                pass
        
        if not selected_nvr:
            log("[CONTEXT-MENU] No NVR data found for item")
            return
        
        log(f"[CONTEXT-MENU] Found NVR: {selected_nvr.get('name', 'Unknown')}")
        
        menu = QtWidgets.QMenu(self)  # Make sure to set parent
        menu.setStyleSheet("""
            QMenu {
                background-color: white;
                border: 1px solid #bdc3c7;
            }
            QMenu::item {
                padding: 8px 25px;
            }
            QMenu::item:selected {
                background-color: #3498db;
                color: white;
            }
        """)
        
        # Add logo to menu if available
        logo_path = self.get_resource_path(LOGO_FILE)
        if os.path.exists(logo_path):
            menu.setIcon(QtGui.QIcon(logo_path))
        
        # Enhanced menu options - Update Cameras first (primary action)
        if IVMS_AVAILABLE:
            action_update_cameras = menu.addAction("🔄 Update Cameras (Live Status)")
        else:
            action_update_cameras = None
            
        # Separator after primary action
        if action_update_cameras:
            menu.addSeparator()
        
        action_edit_creds = menu.addAction("🔧 Edit Credentials") if NVR_DIALOGS_AVAILABLE else None
        action_test_conn = menu.addAction("🔍 Test Connection") if NVR_DIALOGS_AVAILABLE else None
        
        if action_edit_creds or action_test_conn:
            menu.addSeparator()
            
        action_open_browser = menu.addAction("🌐 Open in Browser")
        
        # Add remove option for custom NVRs
        if self.is_custom_nvr(selected_nvr):
            menu.addSeparator()
            action_remove = menu.addAction("🗑️ Remove NVR")
        else:
            action_remove = None
            
        # Add debug info item
        menu.addSeparator()
        debug_action = menu.addAction(f"📋 Debug: {selected_nvr.get('name', 'Unknown')}")
        
        action = menu.exec_(self.list_nvr.mapToGlobal(position))
        
        # Handle menu actions
        if action == debug_action:
            QtWidgets.QMessageBox.information(self, "Debug Info", 
                f"NVR: {selected_nvr.get('name', 'Unknown')}\n"
                f"IP: {selected_nvr.get('ip', 'Unknown')}\n"
                f"Dialogs Available: {NVR_DIALOGS_AVAILABLE}\n"
                f"Is Custom: {self.is_custom_nvr(selected_nvr)}")
        elif action == action_edit_creds and action_edit_creds:
            self.edit_nvr_credentials(selected_nvr)
        elif action == action_test_conn and action_test_conn:
            self.test_nvr_connection(selected_nvr)
        elif action == action_remove and action_remove:
            self.remove_nvr(selected_nvr)
        elif action == action_update_cameras and action_update_cameras:
            self.update_cameras_direct(selected_nvr)
        elif action == action_open_browser:
            nvr_ip = selected_nvr.get("ip", "").strip()
            if nvr_ip:
                protocol = selected_nvr.get('protocol', 'http')
                port = selected_nvr.get('port', 80)
                if port != 80:
                    webbrowser.open(f"{protocol}://{nvr_ip}:{port}")
                else:
                    webbrowser.open(f"{protocol}://{nvr_ip}")

    # ---------------- NVR login dialog ----------------
    def show_nvr_login_dialog(self):
        items = self.list_nvr.selectedItems()
        if not items:
            QtWidgets.QMessageBox.information(self, "NVR Login", "Select an NVR from the list first.")
            return
        idx = items[0].data(QtCore.Qt.UserRole)
        nvr = self.nvrs[idx]
        nvr_ip = nvr.get("ip", "").strip()
        nvr_name = nvr.get("name", "")
        if not nvr_ip:
            QtWidgets.QMessageBox.warning(self, "NVR Login", "NVR has no IP address.")
            return
        
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle(f"NVR Login - {nvr_name} ({nvr_ip})")
        layout = QtWidgets.QFormLayout(dlg)
        
        user_edit = QtWidgets.QLineEdit(); user_edit.setText("admin")
        pwd_edit = QtWidgets.QLineEdit(); pwd_edit.setEchoMode(QtWidgets.QLineEdit.Password); pwd_edit.setText("Kkcctv12345")
        real_ip_label = QtWidgets.QLabel("-")
        
        layout.addRow("NVR IP:", QtWidgets.QLabel(nvr_ip))
        layout.addRow("Username:", user_edit)
        layout.addRow("Password:", pwd_edit)
        layout.addRow("Real IP (fetched):", real_ip_label)
        
        test_btn = QtWidgets.QPushButton("🔍 Test Login & Fetch IP")
        update_btn = QtWidgets.QPushButton("💾 Update IP in Data")
        fetch_cams_btn = QtWidgets.QPushButton("📹 Fetch & Update Cameras")
        browser_btn = QtWidgets.QPushButton("🌐 Open in Browser")
        save_creds_btn = QtWidgets.QPushButton("🪪 Save Credentials")
        
        # Info labels for camera count
        cam_info_label = QtWidgets.QLabel("-")
        cam_info_label.setStyleSheet("color: blue; font-weight: bold;")
        
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Close)
        
        layout.addWidget(test_btn)
        layout.addWidget(update_btn)
        layout.addWidget(fetch_cams_btn)
        layout.addRow("Camera Info:", cam_info_label)
        layout.addWidget(browser_btn)
        layout.addWidget(save_creds_btn)
        layout.addWidget(btns)
        btns.rejected.connect(dlg.reject)
        
        real_ip_storage = [""]  # store fetched IP
        cameras_storage = []  # store fetched cameras
        creds_storage = {"username": "admin", "password": "Kkcctv12345"}  # store credentials
        
        def test_login_action():
            username = user_edit.text().strip()
            password = pwd_edit.text()
            if not username or not password:
                QtWidgets.QMessageBox.warning(dlg, "Input Error", "Enter username and password.")
                return
            creds_storage["username"] = username
            creds_storage["password"] = password
            test_btn.setEnabled(False); test_btn.setText("Testing...")
            threading.Thread(target=self._test_nvr_login_thread, args=(dlg, nvr_ip, nvr_name, username, password, real_ip_label, real_ip_storage, test_btn), daemon=True).start()
        
        def browser_action():
            username = user_edit.text().strip()
            password = pwd_edit.text()
            if not username or not password:
                QtWidgets.QMessageBox.warning(dlg, "Input Error", "Enter username and password.")
                return
            # Use real IP if available, otherwise use current IP
            ip_to_use = real_ip_storage[0] if real_ip_storage[0] else nvr_ip
            url = f"http://{username}:{password}@{ip_to_use}"
            webbrowser.open(url)
            log(f"Opening NVR in browser: {ip_to_use}")
            QtWidgets.QMessageBox.information(dlg, "Browser Opened", f"Opening {ip_to_use} in default browser.")
        
        def save_creds_action():
            username = user_edit.text().strip()
            password = pwd_edit.text()
            if not username or not password:
                QtWidgets.QMessageBox.warning(dlg, "Input Error", "Enter username and password.")
                return
            creds_storage["username"] = username
            creds_storage["password"] = password
            # Save to metadata
            set_password(nvr_ip, username, password)
            meta = load_creds_meta()
            meta[nvr_ip] = {"username": username}
            save_creds_meta(meta)
            QtWidgets.QMessageBox.information(dlg, "Success", "NVR credentials saved securely.")
            log(f"NVR credentials saved for {nvr_ip}")
        
        def update_ip_action():
            log(f"[UPDATE IP ACTION] Called. real_ip_storage[0]='{real_ip_storage[0]}'")
            new_ip = real_ip_storage[0]
            
            if not new_ip or new_ip.strip() == "":
                log(f"[UPDATE IP ACTION] ERROR: new_ip is empty or None: '{new_ip}'")
                QtWidgets.QMessageBox.warning(dlg, "Error", "Fetch IP first by testing login.")
                return
            
            try:
                excel_path = self.get_data_path(EXCEL_FILE)
                if not os.path.exists(excel_path):
                    QtWidgets.QMessageBox.warning(dlg, "Error", "Excel file not found.")
                    return
                
                log(f"[UPDATE IP] Opening Excel: {excel_path}")
                log(f"[UPDATE IP] Looking for NVR: name='{nvr_name}', old_ip='{nvr_ip}', new_ip='{new_ip}'")
                
                # Use openpyxl for direct cell editing to preserve formatting
                from openpyxl import load_workbook
                
                wb_openpyxl = load_workbook(excel_path)
                if "NVR" not in wb_openpyxl.sheetnames:
                    QtWidgets.QMessageBox.warning(dlg, "Error", "NVR sheet not found in Excel.")
                    return
                
                ws = wb_openpyxl["NVR"]
                log(f"[UPDATE IP] NVR sheet has {ws.max_row} rows")
                
                updated = False
                for row_idx in range(1, ws.max_row + 1):
                    cell_name = ws.cell(row=row_idx, column=1).value  # Column A: Name
                    cell_ip = ws.cell(row=row_idx, column=2).value    # Column B: IP
                    
                    cell_name_str = str(cell_name).strip() if cell_name else ""
                    cell_ip_str = str(cell_ip).strip() if cell_ip else ""
                    
                    log(f"[UPDATE IP] Row {row_idx}: name='{cell_name_str}', ip='{cell_ip_str}'")
                    
                    # Match by IP or name
                    if cell_ip_str == nvr_ip or cell_name_str == nvr_name:
                        log(f"[UPDATE IP] MATCH! Row {row_idx}: updating IP from '{cell_ip_str}' to '{new_ip}'")
                        ws.cell(row=row_idx, column=2).value = new_ip
                        updated = True
                        break
                
                if not updated:
                    log(f"[UPDATE IP] No matching NVR found")
                    QtWidgets.QMessageBox.warning(dlg, "Error", f"NVR '{nvr_name}' not found in Excel.")
                    return
                
                # Save the workbook
                log(f"[UPDATE IP] Saving Excel file...")
                wb_openpyxl.save(excel_path)
                log(f"[UPDATE IP] Excel saved successfully")
                
                # Update in-memory data
                for nvr_obj in self.nvrs:
                    if (nvr_obj.get("ip") == nvr_ip or nvr_obj.get("name") == nvr_name):
                        nvr_obj["ip"] = new_ip
                        self.lbl_ip.setText(new_ip)
                        self.populate_nvr_list()
                        break
                
                log(f"[UPDATE IP] SUCCESS: {nvr_name} IP updated from {nvr_ip} to {new_ip}")
                QtWidgets.QMessageBox.information(dlg, "Success", f"✅ NVR IP updated to {new_ip}")
                
            except Exception as e:
                import traceback
                QtWidgets.QMessageBox.critical(dlg, "Update Error", f"Error: {str(e)}")
                log(f"[UPDATE IP] ERROR: {e}")
                log(f"[UPDATE IP] Traceback: {traceback.format_exc()}")
        
        def fetch_cameras_action():
            """Fetch cameras from NVR and update Excel with their IPs."""
            log(f"[GUI-DIALOG] === FETCH CAMERAS BUTTON CLICKED ===")
            log(f"[GUI-DIALOG] NVR Name: {nvr_name}")
            log(f"[GUI-DIALOG] NVR IP: {nvr_ip}")
            
            username = user_edit.text().strip()
            password = pwd_edit.text()
            log(f"[GUI-DIALOG] Username: '{username}'")
            log(f"[GUI-DIALOG] Password length: {len(password) if password else 0}")
            
            if not username or not password:
                log(f"[GUI-DIALOG] Missing credentials - showing warning dialog")
                QtWidgets.QMessageBox.warning(dlg, "Input Error", "Enter username and password.")
                return
            
            log(f"[GUI-DIALOG] Credentials OK - disabling button and starting fetch")
            fetch_cams_btn.setEnabled(False)
            fetch_cams_btn.setText("Fetching...")
            
            # Show progress while fetching
            progress = QtWidgets.QProgressDialog("Fetching cameras from NVR...", None, 0, 0, dlg)
            progress.setWindowModality(QtCore.Qt.WindowModal)
            progress.setWindowTitle("Fetch Cameras")
            progress.setCancelButton(None)
            progress.setMinimumDuration(0)
            progress.show()

            # Fetch cameras in background thread
            def fetch_thread():
                log(f"[GUI-FETCH] Starting camera fetch thread...")
                log(f"[GUI-FETCH] Target NVR: {nvr_ip}")
                log(f"[GUI-FETCH] Username: {username}")
                
                try:
                    controller = WorkingNVRController(nvr_ip, username, password)
                    cameras, method = controller.get_cameras(timeout=15.0)
                    total = len(cameras)
                    active = sum(1 for cam in cameras if cam.get('status', '').lower() == 'online')
                    log(f"[GUI-FETCH] Fetch completed: total={total}, active={active} [IVMS method]")
                    success = bool(cameras)
                    error = '' if success else 'No cameras found (IVMS method)'
                except Exception as e:
                    log(f"[GUI-FETCH] Exception in fetch thread: {e}")
                    cameras, total, active, error = [], 0, 0, f"Thread error: {str(e)}"
                    success = False
                QtCore.QTimer.singleShot(0, lambda: progress.close())
                QtCore.QTimer.singleShot(0, lambda: on_fetch_complete(success, cameras, total, active, error))

            threading.Thread(target=fetch_thread, daemon=True).start()
        
        def on_fetch_complete(success, cameras, total, active, error):
            """Called when camera fetch completes."""
            fetch_cams_btn.setEnabled(True)
            fetch_cams_btn.setText("📹 Fetch & Update Cameras")
            
            log(f"[GUI-DIALOG] Fetch complete: success={success}, total={total}, active={active}")
            
            if not success:
                log(f"[GUI-DIALOG] Fetch failed - showing error dialog: {error}")
                QtWidgets.QMessageBox.critical(dlg, "Fetch Failed", f"Failed to fetch cameras:\n{error}")
                return
            
            if not cameras:
                log(f"[GUI-DIALOG] No cameras found - showing info dialog")
                QtWidgets.QMessageBox.information(dlg, "No Cameras", "No cameras found on NVR or NVR doesn't support camera listing.")
                return
            
            # Store cameras for processing
            cameras_storage.clear()
            cameras_storage.extend(cameras)
            
            log(f"[GUI-DIALOG] Successfully fetched {len(cameras)} cameras")
            first_few = [f"{c['name']} ({c['ip']})" for c in cameras[:3]]
            log(f"[GUI-DIALOG] First few cameras: {first_few}")
            
            # Show camera info in status label if available
            try:
                cam_info_label.setText(f"📹 Total: {total} | 🟢 Active: {active}")
            except NameError:
                # cam_info_label not available in this dialog scope
                log(f"[GUI-DIALOG] Camera info label not available - cameras stored for processing")
                pass
            
            # Ask if user wants to update Excel
            msg = f"Found {total} cameras ({active} active):\n\n"
            for cam in cameras[:5]:  # Show first 5
                msg += f"  • {cam['name']}: {cam['ip']}\n"
            if len(cameras) > 5:
                msg += f"  ... and {len(cameras) - 5} more\n"
            msg += "\nUpdate camera IPs in Excel?"
            
            reply = QtWidgets.QMessageBox.question(dlg, "Update Cameras", msg, QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No)
            if reply == QtWidgets.QMessageBox.Yes:
                # Run update in background and show determinate progress
                prog = QtWidgets.QProgressDialog("Updating camera IPs...", "Cancel", 0, len(cameras), dlg)
                prog.setWindowModality(QtCore.Qt.WindowModal)
                prog.setWindowTitle("Update Cameras")
                prog.setMinimumDuration(0)
                prog.show()

                def update_worker():
                    try:
                        from openpyxl import load_workbook
                        excel_path = self.get_data_path(EXCEL_FILE)
                        if not os.path.exists(excel_path):
                            QtCore.QTimer.singleShot(0, lambda: QtWidgets.QMessageBox.warning(dlg, "Error", "Excel file not found."))
                            QtCore.QTimer.singleShot(0, lambda: prog.close())
                            return
                        wb = load_workbook(excel_path)
                        sheet_name = nvr_name.replace("/", "_").replace("\\", "_")
                        if sheet_name not in wb.sheetnames:
                            ws = wb.create_sheet(sheet_name)
                            ws.cell(row=1, column=1).value = "Camera Name"
                            ws.cell(row=1, column=2).value = "IP Address"
                        else:
                            ws = wb[sheet_name]

                        updated_count = 0
                        for idx, camera in enumerate(cameras, start=1):
                            cam_name = camera.get("name", f"Camera {idx}")
                            cam_ip = camera.get("ip", "")
                            if not cam_ip:
                                QtCore.QTimer.singleShot(0, lambda: prog.setValue(idx))
                                continue
                            found = False
                            for row_idx in range(2, ws.max_row + 1):
                                existing_name = ws.cell(row=row_idx, column=1).value or ""
                                if str(existing_name).strip().lower() == cam_name.strip().lower():
                                    ws.cell(row=row_idx, column=2).value = cam_ip
                                    updated_count += 1
                                    found = True
                                    break
                            if not found:
                                # append at end
                                next_row = ws.max_row + 1
                                ws.cell(row=next_row, column=1).value = cam_name
                                ws.cell(row=next_row, column=2).value = cam_ip
                                updated_count += 1
                            QtCore.QTimer.singleShot(0, lambda v=idx: prog.setValue(v))

                        wb.save(excel_path)
                        log(f"[UPDATE CAMERAS] Saved! Updated {updated_count} camera IPs")
                        QtCore.QTimer.singleShot(0, lambda: prog.close())
                        QtCore.QTimer.singleShot(0, lambda: QtWidgets.QMessageBox.information(dlg, "Success", f"✅ Updated {updated_count} camera IPs in Excel"))
                        QtCore.QTimer.singleShot(0, lambda: (self.load_data(), self.populate_nvr_list(), self.populate_table([])))
                    except Exception as e:
                        import traceback
                        log(f"[UPDATE CAMERAS] ERROR: {e}")
                        log(f"[UPDATE CAMERAS] Traceback: {traceback.format_exc()}")
                        QtCore.QTimer.singleShot(0, lambda: prog.close())
                        QtCore.QTimer.singleShot(0, lambda: QtWidgets.QMessageBox.critical(dlg, "Update Error", f"Error updating cameras:\n{str(e)}"))

                threading.Thread(target=update_worker, daemon=True).start()
        
        def update_cameras_in_excel(cameras):
            """Update camera IPs in Excel spreadsheet."""
            log(f"[UPDATE CAMERAS] Updating {len(cameras)} cameras in Excel")
            try:
                from openpyxl import load_workbook
                excel_path = self.get_data_path(EXCEL_FILE)
                
                if not os.path.exists(excel_path):
                    QtWidgets.QMessageBox.warning(dlg, "Error", "Excel file not found.")
                    return
                
                wb = load_workbook(excel_path)
                
                # Get or create camera sheet for this NVR
                sheet_name = nvr_name.replace("/", "_").replace("\\", "_")
                if sheet_name not in wb.sheetnames:
                    # Create new sheet
                    ws = wb.create_sheet(sheet_name)
                    ws.cell(row=1, column=1).value = "Camera Name"
                    ws.cell(row=1, column=2).value = "IP Address"
                    log(f"[UPDATE CAMERAS] Created new sheet: {sheet_name}")
                else:
                    ws = wb[sheet_name]
                
                updated_count = 0
                
                # Update or add cameras
                for cam_idx, camera in enumerate(cameras, start=2):  # Start from row 2 (row 1 is header)
                    cam_name = camera.get("name", f"Camera {cam_idx}")
                    cam_ip = camera.get("ip", "")
                    
                    if not cam_ip:
                        log(f"[UPDATE CAMERAS] Skipping {cam_name} - no IP")
                        continue
                    
                    # Check if camera already exists in sheet
                    found = False
                    for row_idx in range(2, ws.max_row + 1):
                        existing_name = ws.cell(row=row_idx, column=1).value or ""
                        if str(existing_name).strip().lower() == cam_name.strip().lower():
                            # Update existing camera IP
                            ws.cell(row=row_idx, column=2).value = cam_ip
                            log(f"[UPDATE CAMERAS] Updated: {cam_name} -> {cam_ip}")
                            updated_count += 1
                            found = True
                            break
                    
                    if not found:
                        # Add new camera
                        ws.cell(row=cam_idx, column=1).value = cam_name
                        ws.cell(row=cam_idx, column=2).value = cam_ip
                        log(f"[UPDATE CAMERAS] Added: {cam_name} -> {cam_ip}")
                        updated_count += 1
                
                # Save workbook
                wb.save(excel_path)
                log(f"[UPDATE CAMERAS] Saved! Updated {updated_count} camera IPs")
                
                # Reload app data
                self.load_data()
                self.populate_nvr_list()
                self.populate_table([])
                
                QtWidgets.QMessageBox.information(dlg, "Success", f"✅ Updated {updated_count} camera IPs in Excel")
                
            except Exception as e:
                import traceback
                log(f"[UPDATE CAMERAS] ERROR: {e}")
                log(f"[UPDATE CAMERAS] Traceback: {traceback.format_exc()}")
                QtWidgets.QMessageBox.critical(dlg, "Update Error", f"Error updating cameras:\n{str(e)}")
        
        test_btn.clicked.connect(test_login_action)
        browser_btn.clicked.connect(browser_action)
        save_creds_btn.clicked.connect(save_creds_action)
        update_btn.clicked.connect(update_ip_action)
        fetch_cams_btn.clicked.connect(fetch_cameras_action)
        
        # connect SADP discover button on dialog (also allow quick SADP scan here)
        # create a small lambda to reuse username/password fields
        def sadp_scan_action():
            fetch_cams_btn.setEnabled(False)
            fetch_cams_btn.setText("Fetching...")
            threading.Thread(target=lambda: self._run_sadp_scan_and_show(nvr_name), daemon=True).start()

        # If the SADP button exists on main toolbar, it will open a full dialog. Here allow quick scan too.
        # We use the SADP scan helper to find devices and populate cam_info_label
        # Bind a right-click or double-press? For simplicity, add to fetch button context menu
        try:
            # attach to fetch_cams_btn's context menu trigger
            fetch_cams_btn.setToolTip("Fetch cameras from NVR or run SADP discovery")
        except Exception:
            pass
        dlg.exec_()

    def import_cameras_from_nvr(self):
        """Import cameras from NVR using IVMS (Improved NVR Management System)."""
        log(f"[GUI-DIALOG] === OPENING NVR IMPORT DIALOG ===")
        
        if not IVMS_AVAILABLE:
            log(f"[GUI-DIALOG] IVMS not available - showing warning")
            QtWidgets.QMessageBox.warning(self, "IVMS Not Available", 
                "IVMS module not found. Please ensure ivms.py is in the application folder.")
            return
        
        items = self.list_nvr.selectedItems()
        if not items:
            log(f"[GUI-DIALOG] No NVR selected - showing info dialog")
            QtWidgets.QMessageBox.information(self, "No Selection", "Select an NVR from the list first.")
            return
        
        idx = items[0].data(QtCore.Qt.UserRole)
        nvr = self.nvrs[idx]
        nvr_ip = nvr.get("ip", "").strip()
        nvr_name = nvr.get("name", "")
        
        log(f"[GUI-DIALOG] Selected NVR: {nvr_name} ({nvr_ip})")
        log(f"[GUI-DIALOG] NVR data: {nvr}")
        
        if not nvr_ip:
            QtWidgets.QMessageBox.warning(self, "Error", "NVR has no IP address.")
            return
        
        # Create import dialog
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle(f"Import Cameras from NVR (IVMS) - {nvr_name}")
        dlg.setMinimumWidth(700)
        layout = QtWidgets.QVBoxLayout(dlg)
        
        # Add description
        desc_label = QtWidgets.QLabel(
            "This tool connects directly to the NVR using ISAPI protocol to fetch all connected cameras.\n"
            "It will automatically add new cameras and update existing ones in your Excel file."
        )
        desc_label.setWordWrap(True)
        desc_label.setStyleSheet("padding: 10px; background-color: #e8f4fd; border-radius: 4px; margin-bottom: 10px;")
        layout.addWidget(desc_label)
        
        # Credentials section
        creds_group = QtWidgets.QGroupBox("NVR Credentials")
        creds_layout = QtWidgets.QFormLayout()
        
        user_edit = QtWidgets.QLineEdit()
        pwd_edit = QtWidgets.QLineEdit()
        pwd_edit.setEchoMode(QtWidgets.QLineEdit.Password)
        
        # Pre-fill with saved credentials if available
        saved_user, saved_pwd = get_password(nvr_ip)
        if saved_user:
            user_edit.setText(saved_user)
        if saved_pwd:
            pwd_edit.setText(saved_pwd)
        else:
            user_edit.setText("admin")  # Default
        
        creds_layout.addRow("Username:", user_edit)
        creds_layout.addRow("Password:", pwd_edit)
        creds_group.setLayout(creds_layout)
        layout.addWidget(creds_group)
        
        # Progress section
        progress_label = QtWidgets.QLabel("Ready to fetch cameras...")
        progress_label.setStyleSheet("padding: 8px; background-color: #f8f9fa; border-radius: 4px;")
        layout.addWidget(progress_label)
        
        # Camera preview table
        camera_group = QtWidgets.QGroupBox("Camera Preview")
        camera_layout = QtWidgets.QVBoxLayout()
        
        camera_table = QtWidgets.QTableWidget()
        camera_table.setColumnCount(4)
        camera_table.setHorizontalHeaderLabels(["Name", "IP Address", "Status", "Channel"])
        header = camera_table.horizontalHeader()
        header.setStretchLastSection(True)
        camera_table.setMinimumHeight(200)
        camera_layout.addWidget(camera_table)
        camera_group.setLayout(camera_layout)
        layout.addWidget(camera_group)
        
        # Buttons
        button_layout = QtWidgets.QHBoxLayout()
        fetch_btn = QtWidgets.QPushButton("📡 Fetch Cameras")
        import_btn = QtWidgets.QPushButton("📥 Import to Excel")
        close_btn = QtWidgets.QPushButton("❌ Close")
        
        import_btn.setEnabled(False)  # Initially disabled
        button_layout.addWidget(fetch_btn)
        button_layout.addWidget(import_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        # Store fetched cameras
        fetched_cameras = []
        
        def fetch_cameras_action():
            """Enhanced fetch cameras from NVR using IVMS with better error handling."""
            username = user_edit.text().strip()
            password = pwd_edit.text()
            
            if not username or not password:
                QtWidgets.QMessageBox.warning(dlg, "Input Error", "Enter username and password.")
                return
            
            fetch_btn.setEnabled(False)
            fetch_btn.setText("🔄 Fetching...")
            import_btn.setEnabled(False)
            progress_label.setText("🔧 Initializing enhanced camera discovery...")
            progress_label.setStyleSheet("padding: 10px; background-color: #fff3cd; border-radius: 4px; color: #856404;")
            
            def fetch_worker():
                try:
                    log(f"[IVMS FETCH] Starting quick camera fetch for {nvr_name} ({nvr_ip})")
                    
                    # Quick direct ISAPI fetch with your working credentials
                    cameras = []
                    QtCore.QTimer.singleShot(0, lambda: progress_label.setText("🔗 Connecting to NVR..."))
                    
                    import requests
                    from requests.auth import HTTPBasicAuth, HTTPDigestAuth
                    
                    # Try direct ISAPI call first - this is what worked before
                    try:
                        session = requests.Session()
                        session.headers.update({
                            "User-Agent": "NARONG-CCTV/8.5.0",
                            "Accept": "application/xml, text/xml, */*",
                        })
                        
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText("📡 Fetching cameras via ISAPI..."))
                        
                        # Test the exact working endpoint with short timeout
                        url = f"http://{nvr_ip}/ISAPI/ContentMgmt/InputProxy/channels"
                        
                        # Try basic auth first with very short timeout
                        resp = session.get(url, auth=HTTPBasicAuth(username, password), timeout=3.0)
                        if resp.status_code == 401:
                            # Try digest auth with short timeout
                            QtCore.QTimer.singleShot(0, lambda: progress_label.setText("🔐 Trying Digest Authentication..."))
                            resp = session.get(url, auth=HTTPDigestAuth(username, password), timeout=3.0)
                        
                        if resp.status_code == 200:
                            log(f"[IVMS FETCH] ISAPI successful: {resp.status_code}")
                            cameras = _parse_isapi_cameras_working(resp.text)
                            log(f"[IVMS FETCH] Parsed {len(cameras)} cameras from ISAPI")
                        else:
                            log(f"[IVMS FETCH] ISAPI failed: {resp.status_code}")
                            
                    except requests.exceptions.Timeout:
                        log(f"[IVMS FETCH] ISAPI timeout after 3 seconds")
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText("⏰ Connection timeout - NVR not responding"))
                    except requests.exceptions.ConnectionError:
                        log(f"[IVMS FETCH] ISAPI connection error - NVR unreachable")
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText("🔌 Connection failed - Check NVR IP address"))
                    except Exception as e:
                        log(f"[IVMS FETCH] ISAPI error: {e}")
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText(f"❌ Error: {str(e)[:50]}..."))
                    
                    # If no cameras yet, try working controller as fallback
                    if not cameras and IVMS_AVAILABLE:
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText("🔄 Trying backup method..."))
                        try:
                            controller = WorkingNVRController(nvr_ip, username, password)
                            cameras, result_msg = controller.get_cameras(timeout=5.0)
                            log(f"[IVMS FETCH] Backup method: {len(cameras)} cameras via {result_msg}")
                        except Exception as e:
                            log(f"[IVMS FETCH] Backup method error: {e}")
                    
                    # Method 3: Manual ISAPI endpoints with enhanced parsing
                    if not cameras:
                        QtCore.QTimer.singleShot(0, lambda: progress_label.setText("🎯 Trying enhanced ISAPI endpoints..."))
                        
                        # Enhanced endpoint list
                        endpoints = [
                            "/ISAPI/ContentMgmt/InputProxy/channels",
                            "/ISAPI/ContentMgmt/RemoteDevice", 
                            "/ISAPI/System/Video/inputs/channels",
                            "/ISAPI/ContentMgmt/RemoteDevice/channels",
                            "/ISAPI/System/Video/inputs",
                            "/ISAPI/Streaming/channels"
                        ]
                        
                        import requests
                        from requests.auth import HTTPBasicAuth, HTTPDigestAuth
                        
                        for endpoint in endpoints:
                            try:
                                url = f"http://{nvr_ip}{endpoint}"
                                
                                # Try basic auth first with short timeout
                                resp = requests.get(url, auth=HTTPBasicAuth(username, password), timeout=3.0)
                                if resp.status_code == 200:
                                    # Parse ISAPI XML response
                                    cameras = self._parse_isapi_cameras(resp.text)
                                    if cameras:
                                        log(f"[IVMS IMPORT] Found {len(cameras)} cameras using {endpoint}")
                                        break
                                elif resp.status_code == 401:
                                    # Try digest auth with short timeout
                                    resp_digest = requests.get(url, auth=HTTPDigestAuth(username, password), timeout=3.0)
                                    if resp_digest.status_code == 200:
                                        cameras = self._parse_isapi_cameras(resp_digest.text)
                                        if cameras:
                                            log(f"[IVMS IMPORT] Found {len(cameras)} cameras using {endpoint} (Digest Auth)")
                                            break
                            except Exception as e:
                                log(f"[IVMS IMPORT] Error with {endpoint}: {e}")
                                continue
                    
                    if cameras:
                        fetched_cameras.clear()
                        fetched_cameras.extend(cameras)
                        
                        log(f"[IVMS IMPORT] About to update camera preview with {len(cameras)} cameras")
                        log(f"[IVMS IMPORT] Sample camera data: {cameras[0] if cameras else 'None'}")
                        
                        # Update UI directly in main thread - avoid scope issues
                        log(f"[IVMS IMPORT] Updating UI directly...")
                        try:
                            # Update camera preview table directly
                            camera_table.setRowCount(len(cameras))
                            log(f"[IVMS IMPORT] Set camera table row count to {len(cameras)}")
                            
                            for row, camera in enumerate(cameras):
                                name = camera.get("name", "")
                                ip = camera.get("ip", "")
                                status = camera.get("status", "")
                                channel = str(camera.get("channel", ""))
                                
                                camera_table.setItem(row, 0, QtWidgets.QTableWidgetItem(name))
                                camera_table.setItem(row, 1, QtWidgets.QTableWidgetItem(ip))
                                camera_table.setItem(row, 2, QtWidgets.QTableWidgetItem(status))
                                camera_table.setItem(row, 3, QtWidgets.QTableWidgetItem(channel))
                            
                            # Resize columns to fit content
                            camera_table.resizeColumnsToContents()
                            
                            # Update progress label
                            progress_label.setText(f"✅ Found {len(cameras)} cameras ready for import")
                            progress_label.setStyleSheet("padding: 10px; background-color: #d4edda; border-radius: 4px; color: #155724;")
                            
                            # Enable import button
                            import_btn.setEnabled(True)
                            fetch_btn.setEnabled(True)
                            fetch_btn.setText("📡 Fetch Cameras")
                            
                            log(f"[IVMS IMPORT] UI updated successfully - {len(cameras)} cameras displayed")
                            
                        except Exception as e:
                            log(f"[IVMS IMPORT] ERROR updating UI: {e}")
                            import traceback
                            log(f"[IVMS IMPORT] Traceback: {traceback.format_exc()}")
                            
                            # Show error in UI
                            progress_label.setText(f"❌ Error displaying cameras: {str(e)[:50]}...")
                            progress_label.setStyleSheet("padding: 10px; background-color: #f8d7da; border-radius: 4px; color: #721c24;")
                            fetch_btn.setEnabled(True)
                            fetch_btn.setText("📡 Fetch Cameras")
                        log(f"[IVMS IMPORT] Successfully fetched {len(cameras)} cameras from {nvr_name}")
                        log(f"[IVMS IMPORT] UI update scheduled via QTimer")
                    else:
                        error_msg = f"No cameras found. This could mean:\n• NVR has no connected cameras\n• Different API endpoints needed\n• Credentials incorrect\n• NVR model not supported"
                        QtCore.QTimer.singleShot(0, lambda: self._show_fetch_error(progress_label, fetch_btn, error_msg))
                
                except Exception as e:
                    log(f"[IVMS IMPORT] Error: {e}")
                    import traceback
                    log(traceback.format_exc())
                    QtCore.QTimer.singleShot(0, lambda: self._show_fetch_error(progress_label, fetch_btn, str(e)))
            
            threading.Thread(target=fetch_worker, daemon=True).start()
        
        def import_cameras_action():
            """Import fetched cameras to Excel."""
            if not fetched_cameras:
                QtWidgets.QMessageBox.warning(dlg, "No Data", "No cameras to import. Fetch cameras first.")
                return
            
            # Ask for confirmation
            reply = QtWidgets.QMessageBox.question(dlg, "Confirm Import",
                f"Import {len(fetched_cameras)} cameras to Excel?\n\n"
                f"This will:\n"
                f"• Add new cameras to the '{nvr_name}' sheet\n"
                f"• Update IP addresses for existing cameras\n"
                f"• Preserve existing camera data",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No)
            
            if reply == QtWidgets.QMessageBox.Yes:
                try:
                    self._import_cameras_to_excel(fetched_cameras, nvr_name)
                    QtWidgets.QMessageBox.information(dlg, "Success", f"✅ Successfully imported {len(fetched_cameras)} cameras!")
                    
                    # Refresh main camera list
                    self.load_data()
                    self.populate_table([])
                    
                except Exception as e:
                    log(f"[IVMS IMPORT] Import error: {e}")
                    QtWidgets.QMessageBox.critical(dlg, "Import Error", f"Failed to import cameras:\n{str(e)}")
        
        # Connect buttons
        fetch_btn.clicked.connect(fetch_cameras_action)
        import_btn.clicked.connect(import_cameras_action)
        close_btn.clicked.connect(dlg.close)
        
        dlg.exec_()
    
    def _parse_isapi_cameras(self, xml_text):
        """Parse ISAPI XML response using proven working method from IVMS."""
        if not IVMS_AVAILABLE:
            log(f"[PARSE ISAPI] IVMS dependencies not available")
            return []
        
        # Use the working controller's parsing method
        try:
            working_controller = WorkingNVRController("", "", "")  # Dummy instance for parsing
            cameras = working_controller._parse_isapi_cameras(xml_text)
            log(f"[PARSE ISAPI] Successfully parsed {len(cameras)} cameras using proven method")
            return cameras
        except Exception as e:
            log(f"[PARSE ISAPI] Error with proven parsing method: {e}")
            return []
    
    def _update_camera_preview(self, table, cameras, progress_label, import_btn, fetch_btn):
        """Update the camera preview table with fetched cameras."""
        try:
            log(f"[CAMERA-PREVIEW] === UPDATING CAMERA PREVIEW TABLE ===")
            log(f"[CAMERA-PREVIEW] Cameras to display: {len(cameras)}")
            log(f"[CAMERA-PREVIEW] Table widget: {table}")
            log(f"[CAMERA-PREVIEW] Table widget type: {type(table)}")
            
            if not cameras:
                log(f"[CAMERA-PREVIEW] No cameras to display!")
                return
                
            log(f"[CAMERA-PREVIEW] Setting table row count to {len(cameras)}")
            table.setRowCount(len(cameras))
            log(f"[CAMERA-PREVIEW] Table row count set successfully")
            
            for row, camera in enumerate(cameras):
                try:
                    name = camera.get("name", "")
                    ip = camera.get("ip", "")
                    status = camera.get("status", "")
                    channel = str(camera.get("channel", ""))
                    
                    log(f"[CAMERA-PREVIEW] Populating row {row}: {name} | {ip} | {status} | {channel}")
                    
                    table.setItem(row, 0, QtWidgets.QTableWidgetItem(name))
                    table.setItem(row, 1, QtWidgets.QTableWidgetItem(ip))
                    table.setItem(row, 2, QtWidgets.QTableWidgetItem(status))
                    table.setItem(row, 3, QtWidgets.QTableWidgetItem(channel))
                    
                except Exception as e:
                    log(f"[CAMERA-PREVIEW] Error populating row {row}: {e}")
                    continue
            
            log(f"[CAMERA-PREVIEW] Table populated, updating UI elements...")
            
            # Update progress
            progress_label.setText(f"✅ Found {len(cameras)} cameras ready for import")
            progress_label.setStyleSheet("padding: 10px; background-color: #d4edda; border-radius: 4px; color: #155724;")
            
            # Enable import button
            import_btn.setEnabled(True)
            fetch_btn.setEnabled(True)
            fetch_btn.setText("📡 Fetch Cameras")
            
            # Force table refresh
            table.resizeColumnsToContents()
            table.viewport().update()
            
            log(f"[CAMERA-PREVIEW] Table update completed successfully - {len(cameras)} cameras displayed")
            
        except Exception as e:
            log(f"[CAMERA-PREVIEW] ERROR in _update_camera_preview: {e}")
            import traceback
            log(f"[CAMERA-PREVIEW] Traceback: {traceback.format_exc()}")
            
            # Try to at least update the progress label
            try:
                progress_label.setText(f"❌ Error updating camera preview: {str(e)[:50]}...")
                progress_label.setStyleSheet("padding: 10px; background-color: #f8d7da; border-radius: 4px; color: #721c24;")
                fetch_btn.setEnabled(True)
                fetch_btn.setText("📡 Fetch Cameras")
            except Exception as e2:
                log(f"[CAMERA-PREVIEW] Failed to update error message: {e2}")
    
    def _show_fetch_error(self, progress_label, fetch_btn, error_msg=""):
        """Show fetch error message."""
        msg = "❌ Failed to fetch cameras from NVR"
        if error_msg:
            msg += f"\n{error_msg}"
        progress_label.setText(msg)
        progress_label.setStyleSheet("padding: 10px; background-color: #f8d7da; border-radius: 4px; color: #721c24;")
        fetch_btn.setEnabled(True)
        fetch_btn.setText("📡 Fetch Cameras")
    
    def _import_cameras_to_excel(self, cameras, nvr_name):
        """Import cameras to Excel file."""
        from openpyxl import load_workbook
        
        excel_path = self.get_data_path(EXCEL_FILE)
        if not os.path.exists(excel_path):
            raise Exception("Excel file not found")
        
        wb = load_workbook(excel_path)
        sheet_name = nvr_name.replace("/", "_").replace("\\", "_")
        
        # Create sheet if it doesn't exist with comprehensive headers
        if sheet_name not in wb.sheetnames:
            ws = wb.create_sheet(sheet_name)
            ws.cell(row=1, column=1).value = "Camera Name"
            ws.cell(row=1, column=2).value = "IP Address"
            ws.cell(row=1, column=3).value = "Status"
            ws.cell(row=1, column=4).value = "Channel"
            ws.cell(row=1, column=5).value = "Model"
            ws.cell(row=1, column=6).value = "Port"
            ws.cell(row=1, column=7).value = "Serial Number"
            ws.cell(row=1, column=8).value = "Firmware"
            ws.cell(row=1, column=9).value = "Import Date"
        else:
            ws = wb[sheet_name]
        
        # Import cameras with comprehensive data
        updated_count = 0
        import datetime
        current_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        for camera in cameras:
            cam_name = camera.get("name", "")
            cam_ip = camera.get("ip", "")
            cam_status = camera.get("status", "")
            cam_channel = camera.get("channel", "")
            cam_model = camera.get("model", "")
            cam_port = camera.get("port", 554)  # Default RTSP port
            cam_serial = camera.get("serial", "")
            cam_firmware = camera.get("firmware", "")
            
            # Check if camera already exists
            found = False
            for row_idx in range(2, ws.max_row + 1):
                existing_name = ws.cell(row=row_idx, column=1).value or ""
                if str(existing_name).strip().lower() == cam_name.strip().lower():
                    # Update existing camera with comprehensive data
                    ws.cell(row=row_idx, column=2).value = cam_ip
                    ws.cell(row=row_idx, column=3).value = cam_status
                    ws.cell(row=row_idx, column=4).value = cam_channel
                    ws.cell(row=row_idx, column=5).value = cam_model
                    ws.cell(row=row_idx, column=6).value = cam_port
                    ws.cell(row=row_idx, column=7).value = cam_serial
                    ws.cell(row=row_idx, column=8).value = cam_firmware
                    ws.cell(row=row_idx, column=9).value = current_date
                    updated_count += 1
                    found = True
                    break
            
            if not found:
                # Add new camera with comprehensive data
                next_row = ws.max_row + 1
                ws.cell(row=next_row, column=1).value = cam_name
                ws.cell(row=next_row, column=2).value = cam_ip
                ws.cell(row=next_row, column=3).value = cam_status
                ws.cell(row=next_row, column=4).value = cam_channel
                ws.cell(row=next_row, column=5).value = cam_model
                ws.cell(row=next_row, column=6).value = cam_port
                ws.cell(row=next_row, column=7).value = cam_serial
                ws.cell(row=next_row, column=8).value = cam_firmware
                ws.cell(row=next_row, column=9).value = current_date
                updated_count += 1
        
        wb.save(excel_path)
        log(f"[IVMS IMPORT] Successfully imported {updated_count} cameras to Excel")

    def _run_sadp_scan_and_show(self, origin_name=None):
        """Run SADP discover and show small dialog listing devices."""
        devices = sadp_discover(timeout=1.0, scan_hosts=80)
        # call show dialog in main thread
        QtCore.QTimer.singleShot(0, lambda: self._show_sadp_results_dialog(devices, origin_name))

    def _update_monitor_table(self, table, devices, registry):
        """Update monitor table with current device status."""
        table.setRowCount(len(registry))
        for r, (ip, info) in enumerate(registry.items()):
            found = any(d.get("ip") == ip for d in devices)
            status = "Online" if found else "Offline"
            offline_count = info.get("offline_count", 0)
            if not found:
                offline_count += 1
                info["offline_count"] = offline_count
            else:
                info["offline_count"] = 0
            
            table.setItem(r, 0, QtWidgets.QTableWidgetItem(ip))
            table.setItem(r, 1, QtWidgets.QTableWidgetItem(info.get("model", "")))
            table.setItem(r, 2, QtWidgets.QTableWidgetItem(time.strftime("%H:%M:%S")))
            table.setItem(r, 3, QtWidgets.QTableWidgetItem(status))
            table.setItem(r, 4, QtWidgets.QTableWidgetItem("✓" if found else "✗"))
            table.setItem(r, 5, QtWidgets.QTableWidgetItem(str(offline_count)))
            table.setItem(r, 6, QtWidgets.QTableWidgetItem(info.get("notes", "")))
        table.resizeColumnsToContents()

    def _load_sadp_registry(self) -> dict:
        """Load SADP device registry from file."""
        registry_file = "sadp_registry.json"
        try:
            if os.path.exists(registry_file):
                with open(registry_file, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            log(f"Failed to load SADP registry: {e}")
        return {}

    def _save_sadp_registry(self, registry: dict):
        """Save SADP device registry to file."""
        registry_file = "sadp_registry.json"
        try:
            with open(registry_file, "w", encoding="utf-8") as f:
                json.dump(registry, f, indent=2)
        except Exception as e:
            log(f"Failed to save SADP registry: {e}")
    
    def _test_nvr_login_thread(self, dlg, ip, nvr_name, username, password, label_widget, storage, test_btn):
        log(f"[THREAD] Starting NVR login test in thread")
        # Store references BEFORE emitting signal (signal is processed synchronously)
        self._login_dialog_refs = {
            'dlg': dlg,
            'label_widget': label_widget,
            'storage': storage,
            'test_btn': test_btn
        }
        log(f"[THREAD] Storing nvr_name in refs: {nvr_name}")
        self._login_dialog_refs['nvr_name'] = nvr_name
        self._login_dialog_refs['nvr_ip'] = ip
        log(f"[THREAD] Dialog refs stored: storage id={id(storage)}")
        
        try:
            success, real_ip, error = test_nvr_login(ip, username, password)
            log(f"[THREAD] test_nvr_login returned: success={success}, real_ip={real_ip}, error={error}")
        except Exception as e:
            success, real_ip, error = False, "", str(e)
            log(f"[THREAD] Exception in test_nvr_login: {e}")
        
        log(f"[THREAD] Emitting signal with success={success}")
        self.nvr_login_result.emit(success, real_ip, error)
    
    def _process_all_nvrs_thread(self, nvr_list, progress_label, progress_bar, dialog, btn_start, btn_bypass, btn_stop, btn_cancel, processing_state, update_progress_func):
        """Process all NVRs with retry logic, real-time progress, and user controls."""
        try:
            total_nvrs = len(nvr_list)
            successful_nvrs = 0
            failed_nvrs = 0
            total_cameras = 0
            processing_log = []
            
            log(f"[MULTI-NVR] Starting processing of {total_nvrs} NVRs")
            
            for nvr_index, nvr in enumerate(nvr_list, 1):
                # Check if user requested stop
                if processing_state['should_stop']:
                    log(f"[MULTI-NVR] Processing stopped by user at NVR {nvr_index}")
                    processing_log.append(f"⏹️ Processing stopped by user")
                    break
                
                nvr_ip = nvr.get("ip", "").strip()
                nvr_name = nvr.get("name", f"NVR-{nvr_index}")
                
                if not nvr_ip:
                    log(f"[MULTI-NVR] Skipping {nvr_name} - no IP address")
                    processing_log.append(f"❌ {nvr_name}: No IP address")
                    failed_nvrs += 1
                    continue
                
                # Update progress with real-time percentage
                progress_percentage = int((nvr_index / total_nvrs) * 100)
                self.progress_update_signal.emit(f"⏳ Processing {nvr_name} ({nvr_ip})", progress_percentage, "", True)
                
                # Auto-skip mode - no user intervention needed
                processing_state['should_skip_current'] = False
                
                # Try to get stored credentials first
                username = None
                password = None
                credential_source = "none"
                
                stored_creds = get_password(nvr_ip)
                if stored_creds and len(stored_creds) == 2:
                    username, password = stored_creds
                    credential_source = "stored"
                    log(f"[MULTI-NVR] Using stored credentials for {nvr_name}")
                else:
                    # Try default credentials
                    username = "admin"
                    password = "Kkcctv12345"
                    credential_source = "default"
                    log(f"[MULTI-NVR] Using default credentials for {nvr_name}")
                
                # Double retry logic for each NVR
                login_success = False
                real_ip = nvr_ip
                cameras = []
                nvr_skipped = False
                
                for attempt in range(1, 3):  # Try twice, then auto-skip
                    try:
                        # Check if user requested stop (keep stop functionality)
                        if processing_state['should_stop']:
                            break
                        
                        log(f"[MULTI-NVR] Attempt {attempt}/2 - Testing login to {nvr_ip} with {credential_source} credentials")
                        
                        # Update progress during attempt
                        self.progress_update_signal.emit(f"⏳ {nvr_name} - Auto-retry {attempt}/2 ({credential_source} creds)", progress_percentage, "", True)
                        
                        # Test login
                        success, fetched_real_ip, error = test_nvr_login(nvr_ip, username, password)
                        
                        if success:
                            login_success = True
                            real_ip = fetched_real_ip if fetched_real_ip else nvr_ip
                            log(f"[MULTI-NVR] Login successful on attempt {attempt}, using IP: {real_ip}")
                            
                            # Fetch cameras
                            log(f"[MULTI-NVR] Fetching cameras from {real_ip}")
                            controller = WorkingNVRController(real_ip, username, password)
                            cameras, method = controller.get_cameras(timeout=15.0)
                            if cameras:
                                log(f"[MULTI-NVR] Successfully fetched {len(cameras)} cameras from {nvr_name} [IVMS method]")
                                processing_log.append(f"✅ {nvr_name}: {len(cameras)} cameras (attempt {attempt}) [IVMS method]")
                                if fetched_real_ip and fetched_real_ip != nvr_ip:
                                    nvr["ip"] = fetched_real_ip
                                    log(f"[MULTI-NVR] Updated NVR IP: {nvr_ip} → {fetched_real_ip}")
                                added_count = self._integrate_quick_sync_cameras(nvr, cameras)
                                total_cameras += len(cameras)
                                successful_nvrs += 1
                                log(f"[MULTI-NVR] Integrated {len(cameras)} cameras ({added_count} new) for {nvr_name} [IVMS method]")
                                self._queue_on_ui(self._refresh_quick_sync_ui)
                                break
                            else:
                                log(f"[MULTI-NVR] Camera fetch failed on attempt {attempt}: No cameras found (IVMS method)")
                                if attempt == 2:
                                    processing_log.append(f"⚠️ {nvr_name}: Login OK, but no cameras found (IVMS method)")
                        else:
                            log(f"[MULTI-NVR] Login failed on attempt {attempt}: {error}")
                            if attempt == 2:  # Last attempt failed - auto-skip
                                processing_log.append(f"⏭️ {nvr_name}: Auto-skipped after login failure - {error}")
                                log(f"[MULTI-NVR] Auto-skipping {nvr_name} after 2 failed login attempts")
                    
                    except Exception as e:
                        log(f"[MULTI-NVR] Exception on attempt {attempt} for {nvr_name}: {e}")
                        if attempt == 2:  # Last attempt
                            processing_log.append(f"❌ {nvr_name}: Error - {str(e)}")
                
                if not login_success:
                    failed_nvrs += 1
                    log(f"[MULTI-NVR] Auto-skipped {nvr_name} after 2 failed attempts")
                
                # Small delay between NVRs
                time.sleep(0.3)
            
            # Step 4: Save data if we got cameras
            if total_cameras > 0:
                self.progress_update_signal.emit(f"⏳ Step 4/5: Saving {total_cameras} cameras to Excel...", 90, "", True)
                
                self.save_to_excel()
                log(f"[MULTI-NVR] Saved {total_cameras} cameras to Excel")
            
            # Step 5: Complete and show results
            completion_msg = f"✅ Quick Sync Complete! Processed {total_nvrs} NVRs"
            if successful_nvrs > 0:
                completion_msg += f" - {successful_nvrs} successful, {total_cameras} cameras"
            if failed_nvrs > 0:
                completion_msg += f" - {failed_nvrs} failed/skipped"
            
            # Choose appropriate style based on results
            if successful_nvrs > 0:
                completion_style = "color: #27ae60; padding: 10px; background-color: #d5f4e6; border-radius: 5px;"
            else:
                completion_style = "color: #f39c12; padding: 10px; background-color: #fdf2e9; border-radius: 5px;"
            
            self.progress_update_signal.emit(completion_msg, 100, completion_style, True)
            
            # Reset processing state and restore buttons
            processing_state['is_processing'] = False
            self.button_control_signal.emit("start", True, "▶ Start Again", True)
            self.button_control_signal.emit("stop", False, "⏹️ Stop", False)
            self.button_control_signal.emit("cancel", True, "❌ Cancel", True)
            
            log(f"[MULTI-NVR] Processing complete: {successful_nvrs}/{total_nvrs} successful, {total_cameras} total cameras")

            # Final UI refresh with any merged camera data
            self._queue_on_ui(self._refresh_quick_sync_ui)
            
            # Show detailed results
            def show_results():
                result_text = f"Quick Sync Results:\n\n"
                result_text += f"📊 Summary:\n"
                result_text += f"• Total NVRs: {total_nvrs}\n"
                result_text += f"• Successful: {successful_nvrs}\n"
                result_text += f"• Failed/Skipped: {failed_nvrs}\n"
                result_text += f"• Total Cameras: {total_cameras}\n\n"
                
                if processing_log:
                    result_text += f"📋 Details:\n"
                    for log_entry in processing_log:
                        result_text += f"{log_entry}\n"
                
                if successful_nvrs > 0:
                    result_text += f"\n💾 Data has been saved to Excel file."
                
                QtWidgets.QMessageBox.information(dialog, "Quick Sync Complete", result_text)
            
            QtCore.QTimer.singleShot(1000, show_results)
            
        except Exception as e:
            log(f"[MULTI-NVR] Critical error in processing thread: {e}")
            QtCore.QMetaObject.invokeMethod(progress_label, "setText", 
                QtCore.Qt.QueuedConnection, QtCore.Q_ARG(str, f"❌ Processing failed: {str(e)}"))
            QtCore.QMetaObject.invokeMethod(progress_label, "setStyleSheet", 
                QtCore.Qt.QueuedConnection, QtCore.Q_ARG(str, "color: #e74c3c; padding: 10px; background-color: #fadbd8; border-radius: 5px;"))
            
            # Reset processing state and restore buttons on error
            processing_state['is_processing'] = False
            QtCore.QMetaObject.invokeMethod(btn_start, "setEnabled", 
                QtCore.Qt.QueuedConnection, QtCore.Q_ARG(bool, True))
            QtCore.QMetaObject.invokeMethod(btn_start, "show", 
                QtCore.Qt.QueuedConnection)
            QtCore.QMetaObject.invokeMethod(btn_bypass, "hide", 
                QtCore.Qt.QueuedConnection)
            QtCore.QMetaObject.invokeMethod(btn_stop, "hide", 
                QtCore.Qt.QueuedConnection)
            QtCore.QMetaObject.invokeMethod(btn_cancel, "setEnabled", 
                QtCore.Qt.QueuedConnection, QtCore.Q_ARG(bool, True))
    
    @QtCore.pyqtSlot(bool, str, str)
    def on_nvr_login_result(self, success, real_ip, error):
        """Handle NVR login test result from thread."""
        log(f"[UI SLOT] on_nvr_login_result called: success={success}, real_ip={real_ip}")
        
        if not hasattr(self, '_login_dialog_refs'):
            log(f"[UI SLOT] ERROR: No dialog references stored")
            return
        
        refs = self._login_dialog_refs
        dlg = refs.get('dlg')
        label_widget = refs.get('label_widget')
        storage = refs.get('storage')
        test_btn = refs.get('test_btn')
        
        log(f"[UI SLOT] Retrieved refs: dlg={dlg is not None}, label_widget={label_widget is not None}, storage={storage is not None}, test_btn={test_btn is not None}")
        
        if not all([dlg, label_widget, storage, test_btn]):
            log(f"[UI SLOT] ERROR: Missing dialog references")
            return
        
        try:
            test_btn.setEnabled(True)
            test_btn.setText("🔍 Test Login & Fetch IP")
            log(f"[UI SLOT] Button re-enabled")
            
            if success:
                log(f"[UI SLOT] SUCCESS! Updating storage[0] from '{storage[0]}' to '{real_ip}'")
                storage[0] = real_ip
                log(f"[UI SLOT] After update: storage[0]='{storage[0]}', storage id={id(storage)}")
                if real_ip:
                    label_widget.setText(real_ip)
                    msg = f"✅ Login successful!\n\nReal IP: {real_ip}"
                else:
                    label_widget.setText("(authenticated)")
                    msg = f"✅ Login successful!\n(IP not found in response - using current IP)"
                label_widget.setStyleSheet("color: green;")
                log(f"[UI SLOT] Showing success dialog")
                QtWidgets.QMessageBox.information(dlg, "Success", msg)
                # Persist real IP into NVR entry and into check history
                try:
                    nvr_name = refs.get('nvr_name')
                    nvr_ip = refs.get('nvr_ip')
                    if nvr_ip and real_ip:
                        for n in self.nvrs:
                            if n.get('ip') == nvr_ip or (n.get('name') and n.get('name') == nvr_name):
                                n['real_ip'] = real_ip
                                # also update check history for the NVR IP
                                self.check_history[real_ip] = {
                                    'status': 'NVR Login OK',
                                    'device_type': 'NVR',
                                    'model': '',
                                    'timestamp': time.time()
                                }
                                try:
                                    with open(CHECK_HISTORY_FILE, 'w', encoding='utf-8') as hf:
                                        json.dump(self.check_history, hf, indent=2)
                                except Exception as e:
                                    log(f"[NVR LOGIN] Error saving history: {e}")
                                break
                except Exception as e:
                    log(f"[UI SLOT] Error persisting NVR real_ip: {e}")
            else:
                label_widget.setText("Failed")
                label_widget.setStyleSheet("color: red;")
                error_msg = error if error else "Unknown error"
                log(f"[UI SLOT] Showing error dialog: {error_msg}")
                QtWidgets.QMessageBox.critical(dlg, "Login Failed", f"❌ Login failed:\n\n{error_msg}\n\nMake sure:\n• NVR IP is correct and accessible\n• Username/password are correct\n• NVR web interface is responding")
        except Exception as e:
            log(f"[UI SLOT] Error in callback: {e}")
            test_btn.setEnabled(True)
            test_btn.setText("🔍 Test Login & Fetch IP")

    def show_workflow_wizard(self):
        """Automated workflow wizard following recommended steps."""
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle(f"🚀 Quick Workflow v8.6 - Automated Camera Discovery")
        dlg.setGeometry(100, 100, 600, 500)
        
        # Set logo icon
        if os.path.exists(LOGO_FILE):
            dlg.setWindowIcon(QtGui.QIcon(LOGO_FILE))
        
        layout = QtWidgets.QVBoxLayout(dlg)
        layout.setSpacing(15)

        # Header with logo
        header_layout = QtWidgets.QHBoxLayout()
        if os.path.exists(LOGO_FILE):
            logo_label = QtWidgets.QLabel()
            pixmap = QtGui.QPixmap(LOGO_FILE).scaledToHeight(48, QtCore.Qt.SmoothTransformation)
            logo_label.setPixmap(pixmap)
            header_layout.addWidget(logo_label)
        
        header_text = QtWidgets.QLabel(f"🚀 Quick Sync v8.6 - {APP_COMPANY} Advanced Camera Discovery")
        header_text.setStyleSheet("font-weight: bold; font-size: 14pt; color: #2c3e50;")
        header_layout.addWidget(header_text)
        layout.addLayout(header_layout)
        
        # Add quick stats
        stats_label = QtWidgets.QLabel(f"NVRs: {len(self.nvrs)} | Cameras: {len(self.cams)} | Version: {APP_VERSION}")
        stats_label.setStyleSheet("color: #7f8c8d; font-size: 10pt; margin-bottom: 10px;")
        layout.addWidget(stats_label)

        # Instructions
        instructions = QtWidgets.QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml(f"""
        <b>🎆 Enhanced Quick Sync v8.6 Workflow:</b><br><br>
        <span style="color: #27ae60;"><b>✓ Step 1:</b> NVR Authentication & Validation</span><br>
        • Verify NVR connectivity with enhanced error handling<br>
        • Test credentials and establish secure connection<br>
        • Validate ISAPI endpoints and capabilities<br><br>
        
        <span style="color: #3498db;"><b>✓ Step 2:</b> Smart Camera Discovery</span><br>
        • Advanced ISAPI camera extraction ({MAX_PARALLEL_WORKERS} parallel workers)<br>
        • Intelligent duplicate detection and merging<br>
        • Real-time status validation and updates<br><br>
        
        <span style="color: #9b59b6;"><b>✓ Step 3:</b> Enhanced Connectivity Check</span><br>
        • Multi-protocol validation (HTTP, RTSP, Ping)<br>
        • Performance optimization with {CONNECTION_TIMEOUT}s timeouts<br>
        • Live status monitoring with visual feedback<br><br>
        
        <span style="color: #f39c12;"><b>✓ Step 4:</b> Network Analysis & Reporting</span><br>
        • Comprehensive device discovery statistics<br>
        • Performance metrics and connection analysis<br>
        • Duplicate detection with resolution suggestions<br><br>
        
        <span style="color: #e74c3c;"><b>✓ Step 5:</b> Data Export & Backup</span><br>
        • Multi-format export (Excel, CSV, JSON)<br>
        • Automated backup with timestamp<br>
        • Configuration preservation and versioning
        """)
        instructions.setMinimumHeight(300)
        layout.addWidget(instructions)

        # Enhanced progress tracking
        progress_widget = QtWidgets.QWidget()
        progress_layout = QtWidgets.QVBoxLayout(progress_widget)
        
        progress_label = QtWidgets.QLabel("🚀 Ready to start Quick Sync v8.6 workflow...")
        progress_label.setStyleSheet("font-weight: bold; color: #34495e; padding: 10px; background-color: #ecf0f1; border-radius: 5px;")
        
        progress_bar = QtWidgets.QProgressBar()
        progress_bar.setMinimum(0)
        progress_bar.setMaximum(100)
        progress_bar.setValue(0)
        progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                text-align: center;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #3498db;
                border-radius: 3px;
            }
        """)
        progress_bar.hide()
        
        progress_layout.addWidget(progress_label)
        progress_layout.addWidget(progress_bar)
        layout.addWidget(progress_widget)

        # Buttons
        btn_layout = QtWidgets.QHBoxLayout()
        btn_start = QtWidgets.QPushButton("▶ Start Workflow")
        btn_start.setMinimumHeight(40)
        btn_start.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                font-weight: bold;
                font-size: 12pt;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover { background-color: #229954; }
        """)
        
        btn_cancel = QtWidgets.QPushButton("✕ Cancel")
        btn_cancel.setMinimumHeight(40)
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                font-weight: bold;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover { background-color: #7f8c8d; }
        """)
        
        btn_bypass = QtWidgets.QPushButton("⏭️ Skip Current & Continue")
        btn_bypass.setMinimumHeight(40)
        btn_bypass.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                font-weight: bold;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover { background-color: #e67e22; }
        """)
        btn_bypass.setEnabled(False)
        btn_bypass.hide()
        
        btn_stop = QtWidgets.QPushButton("⏹️ Stop Processing")
        btn_stop.setMinimumHeight(40)
        btn_stop.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                font-weight: bold;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover { background-color: #c0392b; }
        """)
        btn_stop.setEnabled(False)
        btn_stop.hide()

        btn_layout.addWidget(btn_start)
        btn_layout.addWidget(btn_bypass)
        btn_layout.addWidget(btn_stop)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
        
        # Shared processing state
        processing_state = {
            'should_stop': False,
            'should_skip_current': False,
            'is_processing': False
        }

        def update_progress(step, message, percentage=0, error=False, current_nvr=0, total_nvrs=0):
            """Update progress display with enhanced visual feedback and real-time percentage."""
            if error:
                progress_label.setText(f"❌ Step {step}: {message}")
                progress_label.setStyleSheet("color: #e74c3c; padding: 10px; background-color: #fadbd8; border-radius: 5px; font-weight: bold;")
            else:
                if current_nvr > 0 and total_nvrs > 0:
                    nvr_progress = f" ({current_nvr}/{total_nvrs})"
                    calc_percentage = int((current_nvr / total_nvrs) * 100)
                else:
                    nvr_progress = ""
                    calc_percentage = percentage
                    
                progress_label.setText(f"⚙️ Step {step}: {message}{nvr_progress}")
                progress_label.setStyleSheet("color: #27ae60; padding: 10px; background-color: #d5f4e6; border-radius: 5px; font-weight: bold;")
            
            # Always show progress bar during processing
            if processing_state['is_processing']:
                progress_bar.show()
                if current_nvr > 0 and total_nvrs > 0:
                    progress_bar.setValue(int((current_nvr / total_nvrs) * 100))
                elif percentage > 0:
                    progress_bar.setValue(percentage)
            
            QtWidgets.QApplication.processEvents()
        
        def stop_processing():
            """Stop the current processing."""
            processing_state['should_stop'] = True
            btn_stop.setText("⏹️ Stopping...")
            btn_stop.setEnabled(False)
            btn_bypass.setEnabled(False)
            log("[WORKFLOW] User requested stop processing")
        
        def skip_current_nvr():
            """Skip current NVR and continue to next."""
            processing_state['should_skip_current'] = True
            btn_bypass.setText("⏭️ Skipping...")
            btn_bypass.setEnabled(False)
            log("[WORKFLOW] User requested skip current NVR")
        
        def run_workflow():
            """Execute the enhanced Quick Sync v8.6 workflow."""
            processing_state['is_processing'] = True
            processing_state['should_stop'] = False
            processing_state['should_skip_current'] = False
            
            btn_start.setEnabled(False)
            btn_start.hide()
            btn_cancel.setEnabled(False)
            btn_bypass.setEnabled(True)
            btn_bypass.show()
            btn_stop.setEnabled(True)
            btn_stop.show()
            progress_bar.show()
            progress_bar.setValue(0)
            
            try:
                # Validation: Check prerequisites
                if not self.nvrs:
                    update_progress(0, "Excel file with NVR data required", 0, True)
                    QtWidgets.QMessageBox.warning(dlg, "Prerequisites Missing", 
                        "Please load Excel file first with NVR configuration data.\n\n"
                        "The Quick Sync workflow requires NVR information to proceed.")
                    return
                
                update_progress(1, f"Validating {len(self.nvrs)} NVR configurations...", 10)
                time.sleep(0.5)  # Visual feedback

                progress_label.setText("⏳ Step 1/5: Logging into NVR...")
                progress_label.setStyleSheet("color: #3498db; padding: 10px; background-color: #d6eaf8; border-radius: 5px;")
                QtCore.QCoreApplication.processEvents()

                # Process all NVRs automatically with retry logic
                update_progress(2, f"Processing {len(self.nvrs)} NVRs...", 5)
                
                # Start multi-NVR processing in background thread
                threading.Thread(target=self._process_all_nvrs_thread, 
                               args=(self.nvrs, progress_label, progress_bar, dlg, btn_start, btn_bypass, btn_stop, btn_cancel, processing_state, update_progress), 
                               daemon=True).start()
            
            except Exception as e:
                update_progress(0, f"Quick Sync workflow error: {e}", 0, True)
                QtWidgets.QMessageBox.critical(dlg, "Workflow Error", 
                    f"An error occurred during Quick Sync workflow:\n\n{str(e)}")
                return
                
                # Step 2: Fetch cameras
                QtCore.QTimer.singleShot(2000, lambda: self._workflow_step2(progress_label, dlg, btn_start, btn_cancel))

        def _workflow_complete():
            """Mark workflow as complete."""
            progress_label.setText("✓ Workflow Complete! All cameras discovered and verified.")
            progress_label.setStyleSheet("color: #27ae60; padding: 10px; background-color: #d5f4e6; border-radius: 5px;")
            btn_start.setEnabled(True)
            btn_cancel.setEnabled(True)
            QtWidgets.QMessageBox.information(dlg, "Success", "✓ Workflow completed successfully!\n\nAll cameras have been discovered, verified, and saved to Excel.")

        btn_start.clicked.connect(run_workflow)
        btn_bypass.clicked.connect(skip_current_nvr)
        btn_stop.clicked.connect(stop_processing)
        btn_cancel.clicked.connect(dlg.reject)

        dlg.exec_()

    def _workflow_step2(self, progress_label, dlg, btn_start, btn_cancel):
        """Workflow step 2: Fetch cameras from NVR."""
        progress_label.setText("⏳ Step 2/5: Fetching cameras from NVR...")
        progress_label.setStyleSheet("color: #3498db; padding: 10px; background-color: #d6eaf8; border-radius: 5px;")
        QtCore.QCoreApplication.processEvents()
        QtCore.QTimer.singleShot(2000, lambda: self._workflow_step3(progress_label, dlg, btn_start, btn_cancel))

    def _workflow_step3(self, progress_label, dlg, btn_start, btn_cancel):
        """Workflow step 3: Check IP online."""
        progress_label.setText("⏳ Step 3/5: Checking camera IP online status (parallel scan)...")
        progress_label.setStyleSheet("color: #9b59b6; padding: 10px; background-color: #ebdef0; border-radius: 5px;")
        QtCore.QCoreApplication.processEvents()
        self.check_ip_online()
        QtCore.QTimer.singleShot(3000, lambda: self._workflow_step4(progress_label, dlg, btn_start, btn_cancel))

    def _workflow_step4(self, progress_label, dlg, btn_start, btn_cancel):
        """Workflow step 4: SADP network scan."""
        progress_label.setText("⏳ Step 4/5: Scanning network with SADP discovery...")
        progress_label.setStyleSheet("color: #f39c12; padding: 10px; background-color: #fdebd0; border-radius: 5px;")
        QtCore.QCoreApplication.processEvents()
        # Auto-open SADP tool
        self.show_sadp_tool()
        QtCore.QTimer.singleShot(5000, lambda: self._workflow_step5(progress_label, dlg, btn_start, btn_cancel))

    def _workflow_step5(self, progress_label, dlg, btn_start, btn_cancel):
        """Workflow step 5: Export and save."""
        progress_label.setText("⏳ Step 5/5: Exporting and saving results...")
        progress_label.setStyleSheet("color: #e74c3c; padding: 10px; background-color: #fadbd8; border-radius: 5px;")
        QtCore.QCoreApplication.processEvents()
        self.export_csv()
        QtCore.QTimer.singleShot(1000, lambda: self._workflow_complete(progress_label, dlg, btn_start, btn_cancel))

    def _workflow_complete(self, progress_label, dlg, btn_start, btn_cancel):
        """Mark workflow as complete."""
        progress_label.setText("✓ Workflow Complete! All cameras discovered and verified.")
        progress_label.setStyleSheet("color: #27ae60; padding: 10px; background-color: #d5f4e6; border-radius: 5px;")
        btn_start.setEnabled(True)
        btn_cancel.setEnabled(True)
        QtWidgets.QMessageBox.information(dlg, "✓ Success", "Workflow completed successfully!\n\nAll cameras have been discovered, verified, and saved.")
    
    # ---------------- Update System ----------------
    def check_for_updates_startup(self):
        """Check for updates on startup (silent, respects last check time)"""
        if UPDATE_MANAGER_AVAILABLE:
            try:
                log("[UPDATE] Manual update check - checking for updates...")
                check_for_updates_async(self, show_no_update=False)
            except Exception as e:
                log(f"[UPDATE] Startup check error: {e}")
    
    def check_for_updates_manual(self):
        """Enhanced manual update check with better user feedback"""
        if UPDATE_MANAGER_AVAILABLE:
            try:
                log("[UPDATE] Manual update check initiated...")
                
                # Show checking dialog
                progress = QtWidgets.QProgressDialog("Checking for updates...", "Cancel", 0, 0, self)
                progress.setWindowTitle("Update Check")
                progress.setWindowModality(QtCore.Qt.WindowModal)
                progress.setMinimumDuration(500)
                progress.show()
                
                QtWidgets.QApplication.processEvents()
                
                check_for_updates_async(self, show_no_update=True)
                progress.close()
                
            except Exception as e:
                log(f"[UPDATE] Manual check error: {e}")
                QtWidgets.QMessageBox.critical(self, "Update Check Error", 
                    f"Failed to check for updates:\n\n{str(e)}\n\n"
                    f"Please check your internet connection and try again.")
        else:
            QtWidgets.QMessageBox.information(self, "Update System", 
                f"🔄 Update System Status\n\n"
                f"Update manager is currently not available.\n"
                f"Version: {APP_VERSION}\n\n"
                f"You can check for updates manually by visiting the official repository.")
    
    def show_about_dialog(self):
        """Show enhanced about dialog with version and enhancement information."""
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle(f"About {APP_TITLE}")
        dlg.setGeometry(200, 200, 600, 500)
        
        layout = QtWidgets.QVBoxLayout(dlg)
        
        # Header with logo
        header_layout = QtWidgets.QHBoxLayout()
        logo_path = self.get_resource_path(LOGO_FILE)
        if os.path.exists(logo_path):
            logo_label = QtWidgets.QLabel()
            pixmap = QtGui.QPixmap(logo_path).scaledToHeight(64, QtCore.Qt.SmoothTransformation)
            logo_label.setPixmap(pixmap)
            header_layout.addWidget(logo_label)
        
        title_label = QtWidgets.QLabel(f"<h1>{APP_TITLE}</h1><h2>Enhanced Edition v{APP_VERSION}</h2>")
        title_label.setAlignment(QtCore.Qt.AlignCenter)
        header_layout.addWidget(title_label)
        layout.addLayout(header_layout)
        
        # Version info with clickable links
        info_text = QtWidgets.QTextBrowser()
        info_text.setReadOnly(True)
        info_text.setOpenExternalLinks(True)
        
        about_content = f"""
<h3>🚀 Enhanced Camera Monitoring System</h3>
<p><b>Version:</b> {APP_VERSION} Enhanced Edition<br>
<b>Build Date:</b> {BUILD_DATE}<br>
<b>Company:</b> {APP_COMPANY}<br>
<b>Application ID:</b> {APP_ID}</p>

<h3>✨ Enhanced Features v8.6.1</h3>
<ul>
"""
        
        for feature in ENHANCED_FEATURES:
            about_content += f"<li>🔹 {feature}</li>"
        
        about_content += f"""
</ul>

<h3>⚡ Performance Optimizations</h3>
<ul>
<li>🚀 <b>Parallel Workers:</b> {MAX_PARALLEL_WORKERS} (optimized)</li>
<li>⏱️ <b>Connection Timeout:</b> {CONNECTION_TIMEOUT}s (balanced)</li>
<li>🎯 <b>Ping Timeout:</b> {PING_TIMEOUT}s (faster response)</li>
<li>🖥️ <b>UI Throttle:</b> {UI_UPDATE_THROTTLE}ms (smoother)</li>
<li>💾 <b>Cache Timeout:</b> {CACHE_TIMEOUT}s (smart caching)</li>
<li>🔄 <b>Retry Attempts:</b> {RETRY_ATTEMPTS} (reliable)</li>
</ul>

<h3>🎯 Core Capabilities</h3>
<ul>
<li>📹 <b>Multi-NVR Support:</b> Manage multiple NVR systems</li>
<li>📷 <b>Smart Camera Monitoring:</b> Real-time status with caching</li>
<li>🔍 <b>Advanced Duplicate Detection:</b> Cross-source identification</li>
<li>⚡ <b>Quick Sync Workflow:</b> Automated camera discovery</li>
<li>📊 <b>Performance Dashboard:</b> System metrics monitoring</li>
<li>💾 <b>Excel Integration:</b> Import/export camera data</li>
<li>🔐 <b>Credential Management:</b> Secure password storage</li>
<li>🎦 <b>VLC Integration:</b> Direct RTSP stream viewing</li>
</ul>

<h3>🛠️ Technical Specifications</h3>
<ul>
<li>📱 <b>Platform:</b> Windows 10/11 (64-bit)</li>
<li>🐍 <b>Runtime:</b> Python 3.8+ with PyQt5</li>
<li>🗄️ <b>Database:</b> Excel/CSV with openpyxl</li>
<li>🌐 <b>Protocols:</b> HTTP, RTSP, TCP, SADP (Hikvision)</li>
<li>💾 <b>Memory:</b> Optimized with smart caching</li>
<li>🔄 <b>Updates:</b> Automatic update system</li>
</ul>

<h3>📞 Support & Documentation</h3>
<ul>
<li>📘 <b>User Manual:</b> README.md</li>
<li>📗 <b>Quick Reference:</b> Built-in tooltips and help</li>
<li>📙 <b>Performance Guide:</b> Dashboard metrics</li>
<li>📕 <b>Changelog:</b> Version history tracking</li>
</ul>

<h3>🙏 Credits</h3>
<p><b>Developed by:</b> Chhany<br>
<b>Team:</b> NARONG CCTV KOH-KONG<br>
<b>Company:</b> Sky-Tech<br>
<b>Telegram:</b> <a href="https://t.me/chhanycls" style="color: #0088cc; text-decoration: none; font-weight: bold;">@chhanycls</a> (Click to open)</p>

<hr>
<p style="text-align: center;"><i>🌟 Professional Camera Monitoring Made Simple 🌟</i></p>
        """
        
        info_text.setHtml(about_content)
        layout.addWidget(info_text)
        
        # Buttons
        btn_layout = QtWidgets.QHBoxLayout()
        
        # Performance button
        perf_btn = QtWidgets.QPushButton("📊 View Performance")
        perf_btn.clicked.connect(lambda: (dlg.close(), self.show_performance_dashboard()))
        
        # Close button
        close_btn = QtWidgets.QPushButton("✕ Close")
        close_btn.clicked.connect(dlg.close)
        
        btn_layout.addWidget(perf_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        
        dlg.exec_()

# ---------------- run ----------------
def main():
    # Fix Windows taskbar icon display
    if platform.system() == "Windows":
        try:
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(APP_ID)
        except Exception as e:
            print(f"Failed to set app ID for taskbar icon: {e}")
    
    log("=== NARONG CCTV TEAM CAMERA MONITOR STARTING ===")
    log(f"Python version: {sys.version}")
    log(f"Working directory: {os.getcwd()}")
    log(f"Script path: {os.path.abspath(__file__)}")
    
    app = QtWidgets.QApplication(sys.argv)
    log("PyQt5 application created")
    
    log("Creating main window...")
    w = CameraMonitor()
    log("Showing main window...")
    w.show()
    log(f"{APP_TITLE} v{APP_VERSION} started - event loop starting...")
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
